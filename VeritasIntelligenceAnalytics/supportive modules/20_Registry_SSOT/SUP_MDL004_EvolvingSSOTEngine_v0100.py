# -*- coding: utf-8 -*-
from __future__ import annotations

# =============================================================================
# def VIA EVOLVING SSOT ENGINE v0100
# def Purpose: read many document/code/data/UI formats and integrate them into ONE governed SSOT matrix.
# def Policy : read-only, append-only output, no delete, no subprocess, no network, no canonical mutation.
# def Output : one matrix schema exported as CSV / JSON / optional Parquet / HTML.
# =============================================================================

# =============================================================================
# def 00 PARAMETERS · put every adjustable parameter here
# =============================================================================

PARAM_RUN_TAG = "VIA_EVOLVING_SSOT_ENGINE_v0100"
PARAM_DEFAULT_PATH_LIST = r".\VIA_Evolving_SSOT_Input_Paths_20260702.txt"
PARAM_DEFAULT_SCAN_ROOTS = [
    r"C:\Users\tonyk\Downloads",
    r"C:\Users\tonyk\Downloads\VeritasIntelligenceAnalytics",
]
PARAM_OUTPUT_ROOT = r"C:\Users\tonyk\Downloads\VIA_Evolving_SSOT_Output"
PARAM_ENABLE_SCAN_ROOTS_WHEN_PATH_LIST_MISSING = True
PARAM_RECURSIVE_SCAN = True
PARAM_MAX_SCAN_FILES = 20000
PARAM_MAX_FILE_BYTES_FOR_FULL_TEXT = 20 * 1024 * 1024
PARAM_MAX_TEXT_CHARS_PER_FILE = 12000
PARAM_MAX_ZIP_MEMBERS = 2000
PARAM_CSV_ENCODING_OUT = "utf-8-sig"
PARAM_JSON_ENCODING_OUT = "utf-8"
PARAM_READ_ENCODINGS = ["utf-8-sig", "utf-8", "cp950", "big5", "latin-1"]
PARAM_ALLOWED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".json", ".csv", ".tsv", ".xlsx", ".xls",
    ".html", ".htm", ".py", ".ps1", ".bat", ".cmd", ".yaml", ".yml",
    ".pdf", ".docx", ".zip", ".png", ".jpg", ".jpeg", ".webp", ".svg",
}
PARAM_TEXT_LIKE_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".html", ".htm", ".py", ".ps1",
    ".bat", ".cmd", ".json", ".csv", ".tsv", ".yaml", ".yml", ".svg",
}
PARAM_SSOT_KEYWORDS = [
    "SSOT", "single source", "registry", "matrix", "manifest", "schema", "codex",
    "governance", "canonical", "pointer", "ingest", "superbom", "BOM",
    "PLM", "PMBOK", "SAP", "MSProject", "VDF", "VRN", "VIA", "VIS", "VPNS",
]
PARAM_RISK_KEYWORDS = [
    "delete", "remove-item", "stop-process", "invoke-expression", "iex",
    "start-process", "subprocess", "shell=true", "rm -rf", "format-volume",
    "canonical mutation", "writeback", "global ", "exec(", "eval(",
]
PARAM_DOMAIN_RULES = [
    ("SSOT/Registry/Matrix", ["ssot", "registry", "matrix", "manifest", "schema", "codex", "unified"]),
    ("BOM/SAP/PLM", ["bom", "superbom", "sap", "plm", "pmbok", "mro", "component", "supplier", "thermal", "psu"]),
    ("PMIS/ProcessMining", ["pmis", "project", "msproject", "process", "vpns", "mining"]),
    ("UI/Dashboard", ["ui", "html", "console", "dashboard", "launcher", "visual", "lock"]),
    ("Engine/Automation", ["engine", "fetch", "builder", "ingest", "validate", "launch", ".py", ".ps1", ".bat"]),
    ("Financial/MarketData", ["financial", "eps", "stock", "macro", "market", "risk", "valuation", "industry"]),
]
PARAM_MATRIX_COLUMNS = [
    "def_matrix_id",
    "def_run_id",
    "def_ingested_at",
    "def_source_path",
    "def_parent_path",
    "def_file_name",
    "def_ext",
    "def_format_family",
    "def_domain",
    "def_role",
    "def_exists",
    "def_status",
    "def_risk",
    "def_size_bytes",
    "def_modified_time",
    "def_sha12",
    "def_reader",
    "def_text_chars",
    "def_table_count",
    "def_record_count",
    "def_schema_keys",
    "def_keyword_hits",
    "def_risk_hits",
    "def_canonical_area",
    "def_ingest_action",
    "def_validation_gate",
    "def_evolution_score",
    "def_message",
]
PARAM_HTML_TITLE = "VIA · Evolving SSOT Matrix"
PARAM_HTML_DENSE_FONT_SIZE = "11px"

# =============================================================================
# def 01 imports · standard libs first, optional libs are lazy-loaded
# =============================================================================

import argparse
import csv
import datetime as _dt
import hashlib
import html
import io
import json
import os
import re
import sys
import traceback
import zipfile
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple


# =============================================================================
# def 02 optional dependency loader
# =============================================================================

def def_try_import(module_name: str):
    try:
        return __import__(module_name)
    except Exception:
        return None


def def_get_pandas():
    return def_try_import("pandas")


def def_get_openpyxl():
    return def_try_import("openpyxl")


def def_get_bs4():
    try:
        from bs4 import BeautifulSoup
        return BeautifulSoup
    except Exception:
        return None


def def_get_docx():
    try:
        import docx
        return docx
    except Exception:
        return None


def def_get_pil_image():
    try:
        from PIL import Image
        return Image
    except Exception:
        return None


# =============================================================================
# def 03 basic utilities
# =============================================================================

def def_now_iso() -> str:
    return _dt.datetime.now().replace(microsecond=0).isoformat()


def def_safe_str(value: Any) -> str:
    if value is None:
        return ""
    try:
        return str(value)
    except Exception:
        return repr(value)


def def_clean_path_line(line: str) -> str:
    line = (line or "").strip()
    line = line.strip("\ufeff").strip()
    if not line:
        return ""
    if line.startswith('"') and line.endswith('"'):
        line = line[1:-1]
    if line.startswith("'") and line.endswith("'"):
        line = line[1:-1]
    return line.strip()


def def_windows_safe_path(raw_path: str) -> Path:
    return Path(def_clean_path_line(raw_path)).expanduser()


def def_read_small_text(path: Path, max_chars: int = PARAM_MAX_TEXT_CHARS_PER_FILE) -> Tuple[str, str]:
    if not path.exists() or not path.is_file():
        return "", "missing"
    if path.stat().st_size > PARAM_MAX_FILE_BYTES_FOR_FULL_TEXT:
        return "", "skipped_large"
    raw = path.read_bytes()
    for enc in PARAM_READ_ENCODINGS:
        try:
            return raw.decode(enc, errors="replace")[:max_chars], enc
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace")[:max_chars], "utf-8-replace"


def def_sha12(path: Path) -> str:
    try:
        if not path.exists() or not path.is_file():
            return ""
        h = hashlib.sha256()
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()[:12]
    except Exception:
        return ""


def def_rel_parent(path: Path) -> str:
    try:
        return str(path.parent)
    except Exception:
        return ""


def def_modified_iso(path: Path) -> str:
    try:
        return _dt.datetime.fromtimestamp(path.stat().st_mtime).replace(microsecond=0).isoformat()
    except Exception:
        return ""


def def_format_family(ext: str) -> str:
    ext = ext.lower()
    if ext in {".csv", ".tsv", ".xlsx", ".xls"}:
        return "table"
    if ext in {".json", ".yaml", ".yml"}:
        return "structured"
    if ext in {".html", ".htm"}:
        return "html_ui"
    if ext in {".md", ".markdown", ".txt"}:
        return "document_text"
    if ext in {".py", ".ps1", ".bat", ".cmd"}:
        return "code_or_launcher"
    if ext in {".pdf", ".docx"}:
        return "document_binary"
    if ext in {".zip"}:
        return "archive"
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".svg"}:
        return "image_or_vector"
    if ext == "":
        return "folder"
    return "other"


def def_count_tables_in_text(text: str) -> int:
    if not text:
        return 0
    html_tables = len(re.findall(r"<\s*table\b", text, flags=re.I))
    md_tables = len(re.findall(r"\n\s*\|.+\|\s*\n\s*\|[\s:\-\|]+\|", text))
    csv_like = 1 if ("," in text[:3000] and "\n" in text[:3000]) else 0
    return html_tables + md_tables + csv_like


def def_keyword_hits(text: str, name: str, keywords: List[str]) -> str:
    hay = f"{name}\n{text[:5000]}".lower()
    hits = []
    for kw in keywords:
        if kw.lower() in hay:
            hits.append(kw)
    return ";".join(sorted(set(hits), key=lambda x: x.lower()))


def def_classify_domain(path: Path, text: str = "") -> str:
    hay = f"{path.name} {path.suffix} {text[:3000]}".lower()
    for domain, keys in PARAM_DOMAIN_RULES:
        for key in keys:
            if key.lower() in hay:
                return domain
    return "Other"


def def_classify_role(path: Path, text: str = "") -> str:
    hay = f"{path.name} {text[:3000]}".lower()
    ext = path.suffix.lower()
    if "ssot" in hay:
        return "CanonicalSSOT"
    if "registry" in hay:
        return "Registry"
    if "schema" in hay:
        return "Schema"
    if "matrix" in hay:
        return "Matrix"
    if "manifest" in hay:
        return "Manifest"
    if "codex" in hay:
        return "Codex"
    if ext in {".py"}:
        return "PythonEngine"
    if ext in {".ps1", ".bat", ".cmd"}:
        return "LauncherOrchestrator"
    if ext in {".html", ".htm"}:
        return "HTMLUIReport"
    if ext in {".csv", ".xlsx", ".xls"}:
        return "TabularData"
    if ext in {".json", ".yaml", ".yml"}:
        return "ConfigOrData"
    if ext in {".pdf", ".docx", ".md", ".txt"}:
        return "Document"
    if ext in {".zip"}:
        return "ArchivePackage"
    return "Asset"


def def_canonical_area(domain: str, role: str) -> str:
    if role in {"CanonicalSSOT", "Registry", "Schema", "Matrix", "Manifest", "Codex"}:
        return f"SSOT::{role}"
    if domain == "BOM/SAP/PLM":
        return "SSOT::BOM_SAP_PLM"
    if domain == "PMIS/ProcessMining":
        return "SSOT::PMIS_PROCESS"
    if domain == "UI/Dashboard":
        return "SSOT::UI_VISUAL_LOCK"
    if domain == "Engine/Automation":
        return "SSOT::ENGINE_AUTOMATION"
    if domain == "Financial/MarketData":
        return "SSOT::MARKET_FINANCIAL"
    return "SSOT::REFERENCE"


def def_assess_risk(ext: str, role: str, text: str) -> Tuple[str, str]:
    hits = def_keyword_hits(text, role, PARAM_RISK_KEYWORDS)
    lower = text.lower()
    if any(x in lower for x in ["remove-item", "stop-process", "rm -rf", "format-volume"]):
        return "HIGH_REVIEW", hits
    if hits:
        return "MEDIUM_REVIEW", hits
    if ext.lower() in {".ps1", ".bat", ".cmd", ".py"}:
        return "LOW_STATIC_REVIEW"
    return "LOW"


def def_ingest_action(status: str, role: str, risk: str) -> str:
    if status.startswith("MISSING"):
        return "REGISTER_PATH_ONLY"
    if risk.startswith("HIGH"):
        return "REGISTER_REVIEW_ONLY"
    if role in {"CanonicalSSOT", "Registry", "Schema", "Matrix", "Manifest", "Codex"}:
        return "PROMOTE_TO_SSOT_CANDIDATE"
    if role in {"PythonEngine", "LauncherOrchestrator"}:
        return "REGISTER_AS_ENGINE_ASSET_STATIC_ONLY"
    if role in {"HTMLUIReport"}:
        return "LINK_TO_UI_PORTAL"
    return "REGISTER_AS_REFERENCE_ASSET"


def def_validation_gate(status: str, risk: str, role: str) -> str:
    if status.startswith("ERROR"):
        return "BLOCKED_BY_READ_ERROR"
    if status.startswith("MISSING"):
        return "PATH_ONLY_WAITING_SOURCE"
    if risk.startswith("HIGH"):
        return "HARD_REVIEW_BEFORE_USE"
    if role in {"PythonEngine", "LauncherOrchestrator"}:
        return "AST_OR_SCRIPT_REVIEW_BEFORE_RUN"
    return "STATIC_READABLE_PASS"


def def_evolution_score(role: str, domain: str, status: str, risk: str, keyword_hits: str) -> int:
    score = 10
    if status == "READ_OK":
        score += 20
    if role in {"CanonicalSSOT", "Registry", "Schema", "Matrix", "Manifest", "Codex"}:
        score += 35
    if domain in {"SSOT/Registry/Matrix", "BOM/SAP/PLM", "PMIS/ProcessMining"}:
        score += 15
    if keyword_hits:
        score += min(15, 3 * len(keyword_hits.split(";")))
    if risk.startswith("MEDIUM"):
        score -= 10
    if risk.startswith("HIGH"):
        score -= 25
    return max(0, min(100, score))


# =============================================================================
# def 04 path discovery
# =============================================================================

def def_read_path_list(path_list: Path) -> List[Path]:
    if not path_list.exists():
        return []
    text, _ = def_read_small_text(path_list, max_chars=10_000_000)
    out: List[Path] = []
    for line in text.splitlines():
        cleaned = def_clean_path_line(line)
        if cleaned:
            out.append(Path(cleaned))
    return out


def def_scan_roots(roots: List[str], recursive: bool) -> List[Path]:
    out: List[Path] = []
    for root_raw in roots:
        root = Path(root_raw)
        if not root.exists():
            continue
        iterator = root.rglob("*") if recursive else root.glob("*")
        for p in iterator:
            if len(out) >= PARAM_MAX_SCAN_FILES:
                break
            try:
                if p.is_file() and p.suffix.lower() in PARAM_ALLOWED_EXTENSIONS:
                    out.append(p)
                elif p.is_dir():
                    out.append(p)
            except Exception:
                continue
    return out


def def_collect_candidates(path_list: Optional[Path], scan_roots: List[str], recursive: bool) -> List[Path]:
    candidates: List[Path] = []
    if path_list:
        candidates.extend(def_read_path_list(path_list))
    if not candidates and PARAM_ENABLE_SCAN_ROOTS_WHEN_PATH_LIST_MISSING:
        candidates.extend(def_scan_roots(scan_roots, recursive))
    seen = set()
    unique = []
    for p in candidates:
        key = str(p).lower()
        if key not in seen:
            seen.add(key)
            unique.append(p)
    return unique


# =============================================================================
# def 05 readers by format
# =============================================================================

def def_read_json(path: Path) -> Dict[str, Any]:
    text, enc = def_read_small_text(path, max_chars=PARAM_MAX_TEXT_CHARS_PER_FILE)
    info = {"reader": f"json:{enc}", "text": text, "record_count": 0, "schema_keys": "", "table_count": 0}
    try:
        data = json.loads(text)
        if isinstance(data, dict):
            info["schema_keys"] = ",".join(list(map(str, data.keys()))[:80])
            info["record_count"] = len(data)
        elif isinstance(data, list):
            info["record_count"] = len(data)
            if data and isinstance(data[0], dict):
                info["schema_keys"] = ",".join(list(map(str, data[0].keys()))[:80])
        info["table_count"] = 1 if info["record_count"] else 0
    except Exception:
        pass
    return info


def def_read_csv_like(path: Path) -> Dict[str, Any]:
    text, enc = def_read_small_text(path, max_chars=PARAM_MAX_TEXT_CHARS_PER_FILE)
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    info = {"reader": f"csv:{enc}", "text": text, "record_count": 0, "schema_keys": "", "table_count": 1}
    try:
        sample = io.StringIO(text)
        reader = csv.reader(sample, delimiter=delimiter)
        rows = list(reader)
        if rows:
            info["schema_keys"] = ",".join(rows[0][:80])
            info["record_count"] = max(0, len(rows) - 1)
    except Exception:
        pass
    return info


def def_read_excel(path: Path) -> Dict[str, Any]:
    info = {"reader": "xlsx:unavailable", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    pd = def_get_pandas()
    if pd is not None:
        try:
            xl = pd.ExcelFile(path)
            texts = []
            total_rows = 0
            keys = []
            for sheet in xl.sheet_names[:20]:
                df = xl.parse(sheet, nrows=200)
                total_rows += int(len(df))
                keys.extend([f"{sheet}.{c}" for c in list(df.columns)[:20]])
                texts.append(f"[SHEET] {sheet}\n" + df.head(20).to_csv(index=False))
            info.update({
                "reader": "pandas.ExcelFile",
                "text": "\n".join(texts)[:PARAM_MAX_TEXT_CHARS_PER_FILE],
                "record_count": total_rows,
                "schema_keys": ",".join(map(str, keys[:100])),
                "table_count": len(xl.sheet_names),
            })
            return info
        except Exception:
            pass
    opx = def_get_openpyxl()
    if opx is not None:
        try:
            wb = opx.load_workbook(path, read_only=True, data_only=True)
            texts = []
            keys = []
            rows_total = 0
            for ws in wb.worksheets[:20]:
                table_rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i == 0:
                        keys.extend([f"{ws.title}.{def_safe_str(c)}" for c in row[:20]])
                    if i < 30:
                        table_rows.append(",".join(def_safe_str(c) for c in row))
                    rows_total += 1
                    if i >= 200:
                        break
                texts.append(f"[SHEET] {ws.title}\n" + "\n".join(table_rows))
            info.update({
                "reader": "openpyxl",
                "text": "\n".join(texts)[:PARAM_MAX_TEXT_CHARS_PER_FILE],
                "record_count": rows_total,
                "schema_keys": ",".join(keys[:100]),
                "table_count": len(wb.worksheets),
            })
        except Exception:
            pass
    return info


def def_read_html(path: Path) -> Dict[str, Any]:
    text, enc = def_read_small_text(path, max_chars=PARAM_MAX_TEXT_CHARS_PER_FILE)
    info = {
        "reader": f"html:{enc}",
        "text": text,
        "record_count": 0,
        "schema_keys": "",
        "table_count": def_count_tables_in_text(text),
    }
    soup_cls = def_get_bs4()
    if soup_cls:
        try:
            soup = soup_cls(text, "html.parser")
            title = soup.title.get_text(" ", strip=True) if soup.title else ""
            headings = [h.get_text(" ", strip=True) for h in soup.find_all(["h1", "h2", "h3"])[:30]]
            info["schema_keys"] = ",".join([title] + headings)[:1200]
            info["record_count"] = len(soup.find_all("tr"))
            visible_text = soup.get_text("\n", strip=True)
            info["text"] = visible_text[:PARAM_MAX_TEXT_CHARS_PER_FILE]
        except Exception:
            pass
    return info


def def_read_docx(path: Path) -> Dict[str, Any]:
    info = {"reader": "docx:unavailable", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    docx = def_get_docx()
    if not docx:
        return info
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        table_count = len(d.tables)
        rows = sum(len(t.rows) for t in d.tables)
        info.update({
            "reader": "python-docx",
            "text": "\n".join(paras)[:PARAM_MAX_TEXT_CHARS_PER_FILE],
            "record_count": rows,
            "schema_keys": ",".join(paras[:10])[:1000],
            "table_count": table_count,
        })
    except Exception:
        info["reader"] = "docx:error"
    return info


def def_read_pdf(path: Path) -> Dict[str, Any]:
    info = {"reader": "pdf:unavailable", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        parts = []
        for page in doc[:20]:
            parts.append(page.get_text("text"))
        text = "\n".join(parts)
        info.update({
            "reader": "PyMuPDF",
            "text": text[:PARAM_MAX_TEXT_CHARS_PER_FILE],
            "record_count": len(doc),
            "schema_keys": "pages,text_blocks",
            "table_count": def_count_tables_in_text(text),
        })
        return info
    except Exception:
        pass
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        parts = []
        for page in reader.pages[:20]:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts)
        info.update({
            "reader": "pypdf",
            "text": text[:PARAM_MAX_TEXT_CHARS_PER_FILE],
            "record_count": len(reader.pages),
            "schema_keys": "pages,text",
            "table_count": def_count_tables_in_text(text),
        })
    except Exception:
        info["reader"] = "pdf:error_or_dependency_missing"
    return info


def def_read_zip(path: Path) -> Dict[str, Any]:
    info = {"reader": "zipfile", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    try:
        with zipfile.ZipFile(path, "r") as z:
            names = z.namelist()[:PARAM_MAX_ZIP_MEMBERS]
            ext_counts: Dict[str, int] = {}
            sample_lines: List[str] = []
            for name in names:
                ext = Path(name).suffix.lower() or "[folder]"
                ext_counts[ext] = ext_counts.get(ext, 0) + 1
                if len(sample_lines) < 80:
                    sample_lines.append(name)
            info["record_count"] = len(names)
            info["schema_keys"] = json.dumps(ext_counts, ensure_ascii=False)
            info["text"] = "\n".join(sample_lines)[:PARAM_MAX_TEXT_CHARS_PER_FILE]
            info["table_count"] = 1
    except Exception:
        info["reader"] = "zip:error"
    return info


def def_read_image(path: Path) -> Dict[str, Any]:
    info = {"reader": "image:metadata", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    if path.suffix.lower() == ".svg":
        text, enc = def_read_small_text(path, max_chars=PARAM_MAX_TEXT_CHARS_PER_FILE)
        info.update({"reader": f"svg:{enc}", "text": text, "schema_keys": "svg_xml"})
        return info
    Image = def_get_pil_image()
    if Image is None:
        return info
    try:
        with Image.open(path) as img:
            info["schema_keys"] = f"format={img.format},mode={img.mode},size={img.size[0]}x{img.size[1]}"
            info["record_count"] = 1
    except Exception:
        info["reader"] = "image:error"
    return info


def def_read_text_or_code(path: Path) -> Dict[str, Any]:
    text, enc = def_read_small_text(path, max_chars=PARAM_MAX_TEXT_CHARS_PER_FILE)
    ext = path.suffix.lower()
    info = {
        "reader": f"text:{enc}",
        "text": text,
        "record_count": max(0, text.count("\n")),
        "schema_keys": "",
        "table_count": def_count_tables_in_text(text),
    }
    if ext == ".py":
        funcs = re.findall(r"^\s*def\s+([A-Za-z_]\w*)", text, flags=re.M)
        classes = re.findall(r"^\s*class\s+([A-Za-z_]\w*)", text, flags=re.M)
        imports = re.findall(r"^\s*(?:import|from)\s+([A-Za-z0-9_\.]+)", text, flags=re.M)
        info["schema_keys"] = "functions=" + "|".join(funcs[:30]) + ";classes=" + "|".join(classes[:20]) + ";imports=" + "|".join(imports[:30])
    elif ext == ".ps1":
        funcs = re.findall(r"function\s+([A-Za-z0-9_\-]+)", text, flags=re.I)
        info["schema_keys"] = "functions=" + "|".join(funcs[:50])
    elif ext in {".md", ".markdown"}:
        heads = re.findall(r"^\s{0,3}#{1,6}\s+(.+)", text, flags=re.M)
        info["schema_keys"] = ",".join(heads[:50])
    return info


def def_read_any(path: Path) -> Dict[str, Any]:
    ext = path.suffix.lower()
    if not path.exists():
        return {"reader": "missing", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}
    if path.is_dir():
        try:
            child_count = sum(1 for _ in path.iterdir())
        except Exception:
            child_count = 0
        return {"reader": "folder", "text": "", "record_count": child_count, "schema_keys": "folder_children", "table_count": 0}
    if ext in {".json"}:
        return def_read_json(path)
    if ext in {".csv", ".tsv"}:
        return def_read_csv_like(path)
    if ext in {".xlsx", ".xls"}:
        return def_read_excel(path)
    if ext in {".html", ".htm"}:
        return def_read_html(path)
    if ext in {".docx"}:
        return def_read_docx(path)
    if ext in {".pdf"}:
        return def_read_pdf(path)
    if ext in {".zip"}:
        return def_read_zip(path)
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".svg"}:
        return def_read_image(path)
    if ext in PARAM_TEXT_LIKE_EXTENSIONS:
        return def_read_text_or_code(path)
    return {"reader": "unsupported_static_metadata", "text": "", "record_count": 0, "schema_keys": "", "table_count": 0}


# =============================================================================
# def 06 matrix construction
# =============================================================================

def def_build_record(path: Path, run_id: str, seq: int, ingested_at: str) -> Dict[str, Any]:
    try:
        exists = path.exists()
        ext = path.suffix.lower()
        reader_info = def_read_any(path)
        text = reader_info.get("text", "") or ""
        status = "READ_OK" if exists else "MISSING_SOURCE"
        if exists and reader_info.get("reader", "").endswith("error"):
            status = "ERROR_READ_PARTIAL"
        domain = def_classify_domain(path, text)
        role = def_classify_role(path, text)
        risk, risk_hits = def_assess_risk(ext, role, text)
        kw_hits = def_keyword_hits(text, path.name, PARAM_SSOT_KEYWORDS)
        action = def_ingest_action(status, role, risk)
        gate = def_validation_gate(status, risk, role)
        score = def_evolution_score(role, domain, status, risk, kw_hits)
        size = path.stat().st_size if exists and path.is_file() else 0
        record = {
            "def_matrix_id": f"{run_id}-{seq:06d}",
            "def_run_id": run_id,
            "def_ingested_at": ingested_at,
            "def_source_path": str(path),
            "def_parent_path": def_rel_parent(path),
            "def_file_name": path.name,
            "def_ext": ext,
            "def_format_family": def_format_family(ext),
            "def_domain": domain,
            "def_role": role,
            "def_exists": bool(exists),
            "def_status": status,
            "def_risk": risk,
            "def_size_bytes": size,
            "def_modified_time": def_modified_iso(path),
            "def_sha12": def_sha12(path),
            "def_reader": reader_info.get("reader", ""),
            "def_text_chars": len(text),
            "def_table_count": int(reader_info.get("table_count", 0) or 0),
            "def_record_count": int(reader_info.get("record_count", 0) or 0),
            "def_schema_keys": def_safe_str(reader_info.get("schema_keys", ""))[:4000],
            "def_keyword_hits": kw_hits,
            "def_risk_hits": risk_hits,
            "def_canonical_area": def_canonical_area(domain, role),
            "def_ingest_action": action,
            "def_validation_gate": gate,
            "def_evolution_score": score,
            "def_message": "Static readable and registered." if exists else "Path listed but source is unavailable on this machine.",
        }
        return record
    except Exception as e:
        return {
            "def_matrix_id": f"{run_id}-{seq:06d}",
            "def_run_id": run_id,
            "def_ingested_at": ingested_at,
            "def_source_path": str(path),
            "def_parent_path": def_rel_parent(path),
            "def_file_name": path.name,
            "def_ext": path.suffix.lower(),
            "def_format_family": def_format_family(path.suffix.lower()),
            "def_domain": "ERROR",
            "def_role": "ERROR",
            "def_exists": path.exists(),
            "def_status": "ERROR_EXCEPTION",
            "def_risk": "HIGH_REVIEW",
            "def_size_bytes": 0,
            "def_modified_time": "",
            "def_sha12": "",
            "def_reader": "exception",
            "def_text_chars": 0,
            "def_table_count": 0,
            "def_record_count": 0,
            "def_schema_keys": "",
            "def_keyword_hits": "",
            "def_risk_hits": "",
            "def_canonical_area": "SSOT::ERROR_REVIEW",
            "def_ingest_action": "REGISTER_REVIEW_ONLY",
            "def_validation_gate": "BLOCKED_BY_EXCEPTION",
            "def_evolution_score": 0,
            "def_message": f"{type(e).__name__}: {e}",
        }


def def_build_matrix(paths: List[Path], run_id: str) -> List[Dict[str, Any]]:
    ingested_at = def_now_iso()
    records: List[Dict[str, Any]] = []
    total = len(paths)
    for i, path in enumerate(paths, 1):
        if i == 1 or i == total or i % 100 == 0:
            print(f"[{i}/{total}] {path}")
        records.append(def_build_record(path, run_id, i, ingested_at))
    return records


# =============================================================================
# def 07 outputs
# =============================================================================

def def_ensure_out_dir(output_root: Path, run_id: str) -> Path:
    out = output_root / run_id
    out.mkdir(parents=True, exist_ok=True)
    return out


def def_write_csv(records: List[Dict[str, Any]], path: Path) -> None:
    with path.open("w", encoding=PARAM_CSV_ENCODING_OUT, newline="") as f:
        writer = csv.DictWriter(f, fieldnames=PARAM_MATRIX_COLUMNS, extrasaction="ignore")
        writer.writeheader()
        for r in records:
            writer.writerow({k: r.get(k, "") for k in PARAM_MATRIX_COLUMNS})


def def_write_json(records: List[Dict[str, Any]], path: Path) -> None:
    path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding=PARAM_JSON_ENCODING_OUT)


def def_write_parquet_optional(records: List[Dict[str, Any]], path: Path) -> str:
    pd = def_get_pandas()
    if pd is None:
        return "SKIP_PANDAS_MISSING"
    try:
        df = pd.DataFrame(records)
        df.to_parquet(path, index=False)
        return "OK"
    except Exception as e:
        return f"SKIP_PARQUET_ERROR:{type(e).__name__}:{e}"


def def_html_escape(v: Any) -> str:
    return html.escape(def_safe_str(v))


def def_write_html(records: List[Dict[str, Any]], path: Path, run_id: str) -> None:
    cols = PARAM_MATRIX_COLUMNS
    kpi = {
        "Rows": len(records),
        "SSOT Candidates": sum(1 for r in records if r.get("def_ingest_action") == "PROMOTE_TO_SSOT_CANDIDATE"),
        "High Review": sum(1 for r in records if def_safe_str(r.get("def_risk")).startswith("HIGH")),
        "Missing": sum(1 for r in records if not r.get("def_exists")),
    }
    rows_html = []
    for r in records:
        tds = "".join(f"<td>{def_html_escape(r.get(c,''))}</td>" for c in cols)
        rows_html.append(f"<tr data-risk='{def_html_escape(r.get('def_risk',''))}'><td>{def_html_escape(r.get('def_evolution_score',''))}</td>{tds}</tr>")
    head = "".join(f"<th>{def_html_escape(c)}</th>" for c in ["score"] + cols)
    kpis = "".join(f"<div class='kpi'><span>{def_html_escape(k)}</span><b>{def_html_escape(v)}</b></div>" for k, v in kpi.items())
    html_text = f"""<!doctype html>
<html lang="zh-Hant"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{def_html_escape(PARAM_HTML_TITLE)}</title>
<style>
:root{{--bg:#F9F9F6;--sf:#fff;--ink:#1F2933;--mut:#6B7C78;--line:#DFECEA;--ok:#166534;--warn:#B7791F;--fail:#B91C1C;--blue:#2563EB;}}
*{{box-sizing:border-box}}body{{margin:0;background:var(--bg);color:var(--ink);font-family:Inter,'Noto Sans TC',system-ui,sans-serif;font-size:{PARAM_HTML_DENSE_FONT_SIZE};}}
.wrap{{max-width:1680px;margin:0 auto;padding:18px}}h1{{font-size:20px;margin:0 0 6px}}.sub{{color:var(--mut);font-family:ui-monospace,Consolas,monospace}}
.kpis{{display:grid;grid-template-columns:repeat(auto-fit,minmax(160px,1fr));gap:10px;margin:14px 0}}.kpi{{background:var(--sf);border:1px solid var(--line);border-radius:10px;padding:10px}}.kpi span{{display:block;color:var(--mut);font-family:ui-monospace,Consolas,monospace;font-size:10px}}.kpi b{{font-size:18px}}
.controls{{display:flex;gap:8px;margin:12px 0;flex-wrap:wrap}}input,button{{font:inherit;padding:7px 9px;border:1px solid var(--line);border-radius:8px;background:white}}
.tablewrap{{overflow:auto;background:var(--sf);border:1px solid var(--line);border-radius:10px;max-height:76vh}}
table{{border-collapse:collapse;width:100%;min-width:1900px}}th,td{{border-bottom:1px solid var(--line);padding:5px 7px;vertical-align:top}}th{{position:sticky;top:0;background:#fff;text-align:left;z-index:1}}td{{font-family:ui-monospace,Consolas,monospace;max-width:320px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
tr:hover td{{background:#f3f7f6}}.foot{{margin-top:10px;color:var(--mut);font-family:ui-monospace,Consolas,monospace}}
</style></head><body><div class="wrap">
<h1>{def_html_escape(PARAM_HTML_TITLE)}</h1>
<div class="sub">RUN_ID={def_html_escape(run_id)} · policy=read-only/no-delete/no-subprocess/no-canonical-mutation</div>
<div class="kpis">{kpis}</div>
<div class="controls"><input id="q" placeholder="Search matrix..." oninput="filterRows()"><button onclick="setRisk('')">ALL</button><button onclick="setRisk('HIGH')">HIGH</button><button onclick="setRisk('MEDIUM')">MEDIUM</button><button onclick="setRisk('LOW')">LOW</button></div>
<div class="tablewrap"><table id="tbl"><thead><tr>{head}</tr></thead><tbody>{''.join(rows_html)}</tbody></table></div>
<div class="foot">def Matrix schema is fixed. Physical exports may be CSV/JSON/Parquet/HTML, but logical SSOT format is one matrix.</div>
</div><script>
let riskFilter='';
function setRisk(r){{riskFilter=r;filterRows();}}
function filterRows(){{
  const q=(document.getElementById('q').value||'').toLowerCase();
  document.querySelectorAll('#tbl tbody tr').forEach(tr=>{{
    const text=tr.innerText.toLowerCase();
    const risk=tr.getAttribute('data-risk')||'';
    tr.style.display = ((!q || text.includes(q)) && (!riskFilter || risk.includes(riskFilter))) ? '' : 'none';
  }});
}}
</script></body></html>"""
    path.write_text(html_text, encoding="utf-8")


def def_write_manifest(records: List[Dict[str, Any]], out_dir: Path, run_id: str, parquet_status: str) -> None:
    counts = {
        "run_id": run_id,
        "generated_at": def_now_iso(),
        "rows": len(records),
        "columns": PARAM_MATRIX_COLUMNS,
        "policy": {
            "read_only": True,
            "delete": False,
            "subprocess": False,
            "network": False,
            "canonical_mutation": False,
            "output_mode": "append_only_run_folder",
        },
        "parquet_status": parquet_status,
        "status_counts": {},
        "domain_counts": {},
        "role_counts": {},
        "risk_counts": {},
    }
    for key, out_key in [
        ("def_status", "status_counts"),
        ("def_domain", "domain_counts"),
        ("def_role", "role_counts"),
        ("def_risk", "risk_counts"),
    ]:
        for r in records:
            v = def_safe_str(r.get(key, ""))
            counts[out_key][v] = counts[out_key].get(v, 0) + 1
    (out_dir / "VIA_Evolving_SSOT_RunManifest.json").write_text(json.dumps(counts, ensure_ascii=False, indent=2), encoding="utf-8")


def def_write_outputs(records: List[Dict[str, Any]], out_dir: Path, run_id: str) -> Dict[str, str]:
    paths = {
        "csv": str(out_dir / "VIA_Evolving_SSOT_Matrix.csv"),
        "json": str(out_dir / "VIA_Evolving_SSOT_Matrix.json"),
        "parquet": str(out_dir / "VIA_Evolving_SSOT_Matrix.parquet"),
        "html": str(out_dir / "VIA_Evolving_SSOT_Matrix.html"),
        "manifest": str(out_dir / "VIA_Evolving_SSOT_RunManifest.json"),
    }
    def_write_csv(records, Path(paths["csv"]))
    def_write_json(records, Path(paths["json"]))
    parquet_status = def_write_parquet_optional(records, Path(paths["parquet"]))
    def_write_html(records, Path(paths["html"]), run_id)
    def_write_manifest(records, out_dir, run_id, parquet_status)
    return paths


# =============================================================================
# def 08 CLI and main
# =============================================================================

def def_make_run_id() -> str:
    stamp = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"RUN_{stamp}_{PARAM_RUN_TAG}"


def def_parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="VIA Evolving SSOT Engine: integrate all formats into one governed SSOT matrix.")
    parser.add_argument("--path-list", default=PARAM_DEFAULT_PATH_LIST, help="Text file containing one source path per line.")
    parser.add_argument("--output-root", default=PARAM_OUTPUT_ROOT, help="Append-only output root.")
    parser.add_argument("--scan-root", action="append", default=None, help="Optional scan root. Can be repeated.")
    parser.add_argument("--no-recursive", action="store_true", help="Disable recursive root scan.")
    return parser.parse_args()


def def_print_summary(records: List[Dict[str, Any]], outputs: Dict[str, str]) -> None:
    def_count = lambda key: sorted({r.get(key, "") for r in records})
    print("")
    print("=" * 88)
    print("def VIA EVOLVING SSOT ENGINE COMPLETE")
    print("=" * 88)
    print(f"def Rows: {len(records)}")
    print(f"def Domains: {', '.join(map(str, def_count('def_domain')))}")
    print(f"def Roles: {', '.join(map(str, def_count('def_role')))}")
    for k, v in outputs.items():
        print(f"def Output {k}: {v}")


def def_main() -> int:
    args = def_parse_args()
    run_id = def_make_run_id()
    path_list = Path(args.path_list) if args.path_list else None
    scan_roots = args.scan_root or PARAM_DEFAULT_SCAN_ROOTS
    recursive = not args.no_recursive and PARAM_RECURSIVE_SCAN
    output_root = Path(args.output_root)
    print("=" * 88)
    print("def VIA · EVOLVING SSOT ENGINE v0100")
    print("def Policy: read-only · append-only · no delete · no subprocess · no network · no canonical mutation")
    print("=" * 88)
    print(f"def RunId: {run_id}")
    print(f"def PathList: {path_list}")
    print(f"def OutputRoot: {output_root}")
    candidates = def_collect_candidates(path_list, scan_roots, recursive)
    print(f"def Candidates: {len(candidates)}")
    out_dir = def_ensure_out_dir(output_root, run_id)
    records = def_build_matrix(candidates, run_id)
    outputs = def_write_outputs(records, out_dir, run_id)
    def_print_summary(records, outputs)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(def_main())
    except KeyboardInterrupt:
        print("def Interrupted by operator. No destructive action was executed.")
        raise
    except Exception as exc:
        print("[FATAL]", type(exc).__name__, exc)
        traceback.print_exc()
        raise SystemExit(2)
