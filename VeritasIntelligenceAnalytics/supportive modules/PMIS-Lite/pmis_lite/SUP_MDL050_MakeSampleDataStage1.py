"""擴充測試資料：MS Project XML 匯出、PLM 匯出、郵件附件。
用來在本機驗證 msproject / plm adapter、附件擷取、完整性驗證、利害關係人、追蹤事項。
"""
# ===== [VIA:ACCEL-BRIDGE:v0100] SuperAccel 加速器橋(批102 全樹導入令;graceful 零行為變更) =====
try:
    import sys as _sa_sys
    from pathlib import Path as _sa_Path
    _sa_p = _sa_Path(__file__).resolve()
    while _sa_p.parent != _sa_p:
        if (_sa_p / "supportive modules" / "VIA_SuperAccel_Module.py").exists():
            _sa_sys.path.insert(0, str(_sa_p / "supportive modules"))
            break
        _sa_p = _sa_p.parent
    import VIA_SuperAccel_Module as VIA_ACCEL  # noqa: N816
except Exception:
    VIA_ACCEL = None  # graceful:加速器缺席零影響
# ===== [VIA:ACCEL-BRIDGE:END] =====
import os
import pandas as pd
from openpyxl import Workbook
import docx

HERE = os.path.dirname(os.path.abspath(__file__))
SD = os.path.join(HERE, "sample_data")
ATT = os.path.join(SD, "attachments")
os.makedirs(ATT, exist_ok=True)

# 1) MS Project 原生 XML 匯出（精簡但符合 schema）
msp_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<Project xmlns="http://schemas.microsoft.com/project">
  <Name>PSU_NewProject.xml</Name>
  <Resources>
    <Resource><UID>1</UID><Name>Alex</Name></Resource>
    <Resource><UID>2</UID><Name>Ben</Name></Resource>
  </Resources>
  <Tasks>
    <Task><UID>1</UID><Name>PSU 電路設計</Name><PercentComplete>100</PercentComplete><Start>2026-06-01T08:00:00</Start><Finish>2026-06-10T17:00:00</Finish></Task>
    <Task><UID>2</UID><Name>PSU EVT 樣品製作</Name><PercentComplete>60</PercentComplete><Start>2026-06-11T08:00:00</Start><Finish>2026-06-30T17:00:00</Finish></Task>
    <Task><UID>3</UID><Name>PSU DVT 驗證</Name><PercentComplete>0</PercentComplete><Start>2026-07-01T08:00:00</Start><Finish>2026-07-20T17:00:00</Finish></Task>
    <Task><UID>4</UID><Name></Name><PercentComplete>0</PercentComplete><Start>2026-07-01T08:00:00</Start><Finish>2026-07-20T17:00:00</Finish></Task>
  </Tasks>
  <Assignments>
    <Assignment><TaskUID>1</TaskUID><ResourceUID>1</ResourceUID></Assignment>
    <Assignment><TaskUID>2</TaskUID><ResourceUID>2</ResourceUID></Assignment>
  </Assignments>
</Project>
'''
with open(os.path.join(SD, "project_export.xml"), "w", encoding="utf-8") as f:
    f.write(msp_xml)

# 2) PLM 匯出（模擬 Windchill 風格欄名）
plm = pd.DataFrame([
    {"Object": "PSU-2026-001 規格書", "Lifecycle State": "In Work", "Owner": "Alex",
     "Product": "電源供應器", "Modified": "2026-06-26", "Description": "PSU 主規格 v3，LLC 拓樸"},
    {"Object": "VFD-2026-014 ECN", "Lifecycle State": "Released", "Owner": "Ben",
     "Product": "變頻器", "Modified": "2026-06-24", "Description": "變頻器散熱片變更，請追蹤量產導入"},
    {"Object": "", "Lifecycle State": "In Work", "Owner": "", "Product": "",
     "Modified": "", "Description": ""},  # 空列 → 應被完整性驗證標記為 failed(no_title_or_body)
])
plm.to_csv(os.path.join(SD, "plm_export.csv"), index=False, encoding="utf-8-sig")

# 3) 郵件附件：一份 docx 規格、一份 xlsx 測試結果
d = docx.Document()
d.add_heading("PSU LLC 效率測試報告", 0)
d.add_paragraph("測試人：Alex　日期：2026-06-27")
d.add_paragraph("結論：輕載效率 78%，未達 82% 目標，請 RD 追蹤改善，2026-07-05 前回覆。")
tbl = d.add_table(rows=2, cols=2)
tbl.rows[0].cells[0].text = "負載"; tbl.rows[0].cells[1].text = "效率"
tbl.rows[1].cells[0].text = "10%"; tbl.rows[1].cells[1].text = "78%"
d.save(os.path.join(ATT, "PSU_LLC_efficiency_report.docx"))

wb = Workbook(); ws = wb.active; ws.title = "EVT結果"
ws.append(["項目", "規格", "實測", "判定"])
ws.append(["輸出電壓", "12V±5%", "12.1V", "PASS"])
ws.append(["輕載效率", "≥82%", "78%", "FAIL"])
wb.save(os.path.join(ATT, "EVT_results.xlsx"))

# 一個壞檔（不支援類型）測完整性驗證會不會誠實標記
with open(os.path.join(ATT, "weird_file.dat"), "wb") as f:
    f.write(b"\x00\x01binary-not-supported")

print("extended sample data written:")
for root, _, files in os.walk(SD):
    for n in sorted(files):
        print("  -", os.path.relpath(os.path.join(root, n), SD))
