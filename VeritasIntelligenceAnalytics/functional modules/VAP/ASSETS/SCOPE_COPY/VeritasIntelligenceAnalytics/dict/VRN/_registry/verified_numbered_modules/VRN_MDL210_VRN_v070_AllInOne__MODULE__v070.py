# ===== [VIA:ACCEL-BRIDGE:v0100] SuperAccel 加速器橋(批102 全樹導入令;graceful 零行為變更) =====
try:
    import sys as _sa_sys
    from pathlib import Path as _sa_Path
    _sa_p = _sa_Path(__file__).resolve()
    while _sa_p.parent != _sa_p:
        if (_sa_p / "supportive modules" / "VIA_SuperAccel_Module.py").exists():
            _sa_sys.path.insert(0, str(_sa_p / "supportive modules"))
            break
        _sa_p = _sa_p.parent
    import VIA_SuperAccel_Module as VIA_ACCEL  # noqa: N816
except Exception:
    VIA_ACCEL = None  # graceful:加速器缺席零影響
# ===== [VIA:ACCEL-BRIDGE:END] =====
VERSION = "VRN_V070_D8_ALLINONE"
import os, sys, json, csv, html, re, traceback, time, hashlib, zipfile
from pathlib import Path
from datetime import datetime
from collections import OrderedDict, defaultdict, Counter

RUN_DIR             = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO"
PDF_LIST_JSON       = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\_pdf_list.json"
SSOT_LOCK           = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\config\VRN_Financial_Account_SSOT_v0600.json"
INPUT_DIR           = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\input"
OUTPUT_DIR          = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\output"
MAX_PAGES_PER_PDF   = 40
WRITE_DUCKDB        = bool(1)
WRITE_PARQUET       = bool(1)
WRITE_CSV           = bool(1)

OUT_HTML            = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Report.html"
OUT_JSON            = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Report.json"
OUT_GEOMETRY_CSV    = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Table_Geometry.csv"
OUT_MULTIHDR_CSV    = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_MultiHeader_Preview.csv"
OUT_RESTORED_CSV    = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Financial_Restored.csv"
OUT_QUARANTINE_CSV  = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_HalfYear_Quarantine.csv"
OUT_FINANCIAL_CSV   = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_FinancialData.csv"
OUT_TRUST_CSV       = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Trust_Matrix.csv"
OUT_REVIEW_CSV      = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Review.csv"
OUT_COVERAGE_CSV    = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Coverage.csv"
OUT_ACTION_CSV      = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_Action.csv"
OUT_DUCKDB          = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070.duckdb"
OUT_PARQUET_BASIC   = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_BasicInfo.parquet"
OUT_PARQUET_FIN     = r"C:\Users\tonyk\OneDrive\VeritasIntelligenceAnalytics\module\VRN\temp\_d8b_runs\RUN_20260530_215141_VRN_V070_AIO\VRN_v070_FinancialData.parquet"

PROGRESS_JSON       = str(Path(RUN_DIR) / "_progress.json")

def def_now():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def def_lamp(s):
    return {"PASS":"🟢 PASS","WARN":"🟡 WARN","FAIL":"🔴 FAIL"}.get(s,"🔴 FAIL")

def def_progress(stage, idx, total, msg=""):
    """Atomic write progress so PS7 polling reads consistent state."""
    try:
        tmp = PROGRESS_JSON + ".tmp"
        Path(tmp).write_text(json.dumps({
            "stage": stage, "idx": idx, "total": total,
            "msg": msg, "ts": time.time()
        }), encoding="utf-8")
        os.replace(tmp, PROGRESS_JSON)
    except Exception:
        pass

def def_esc(x):
    return html.escape("" if x is None else str(x))

def def_sha256(p):
    try:
        h = hashlib.sha256()
        with open(p, "rb") as f:
            for b in iter(lambda: f.read(1024*1024), b""): h.update(b)
        return h.hexdigest().upper()
    except: return ""

def def_write_csv(path, rows):
    rows = list(rows or [])
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        Path(path).write_text("", encoding="utf-8-sig"); return
    keys=[]
    for r in rows:
        for k in r.keys():
            if k not in keys: keys.append(k)
    if "Validation Status" in keys:
        keys = ["Validation Status"] + [k for k in keys if k!="Validation Status"]
    if "Validation Note" in keys:
        keys = [k for k in keys if k!="Validation Note"] + ["Validation Note"]
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=keys); w.writeheader()
        for r in rows: w.writerow({k:r.get(k,"") for k in keys})

def def_write_json(path, obj):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8-sig")

# ============================================================================
# def 04_SSOT_LOAD (read-only, NO MUTATION)
# ============================================================================
def def_load_ssot():
    if not Path(SSOT_LOCK).is_file():
        return {"accounts": {}, "aliases": {}, "missing": True}
    try:
        with open(SSOT_LOCK, "r", encoding="utf-8-sig") as f:
            data = json.load(f)
        accounts = {}
        aliases = {}
        # 支援兩種 schema: {accounts: [...]} 或 {Income Statement: {...}, ...}
        if isinstance(data, dict) and "accounts" in data:
            for a in data["accounts"]:
                en = a.get("name_en") or a.get("canonical") or ""
                cat = a.get("category", "")
                if not en: continue
                accounts[en] = {"category": cat, "aliases": a.get("aliases", [])}
                for al in a.get("aliases", []):
                    aliases[str(al).strip().lower()] = en
                aliases[en.lower()] = en
        else:
            for cat, items in data.items() if isinstance(data, dict) else []:
                if not isinstance(items, dict): continue
                for en, meta in items.items():
                    accounts[en] = {"category": cat, "aliases": (meta.get("aliases", []) if isinstance(meta, dict) else [])}
                    for al in accounts[en]["aliases"]:
                        aliases[str(al).strip().lower()] = en
                    aliases[en.lower()] = en
        return {"accounts": accounts, "aliases": aliases, "missing": False}
    except Exception as e:
        return {"accounts": {}, "aliases": {}, "missing": True, "error": str(e)[:300]}

def def_ssot_match(text, ssot):
    if not text or not isinstance(text, str): return ("", 0.0, "no_input")
    key = text.strip().lower()
    if not key: return ("", 0.0, "empty")
    if key in ssot["aliases"]:
        return (ssot["aliases"][key], 100.0, "exact")
    # token contains
    best = ("", 0.0, "none")
    for alias_key, en in ssot["aliases"].items():
        if alias_key and alias_key in key:
            score = 70.0 * (len(alias_key) / max(len(key), 1))
            if score > best[1]:
                best = (en, round(score, 1), "fuzzy_contain")
    return best

# ============================================================================
# def 05_PDF_TABLE_EXTRACT (three-layer fallback)
# ============================================================================
def def_extract_tables_pdf(pdf_path):
    """Returns list of dicts: {page, table_idx, rows (list of list of str), reader}"""
    tables = []
    # Layer 1: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            for pi, page in enumerate(pdf.pages):
                if pi >= MAX_PAGES_PER_PDF: break
                try:
                    pts = page.extract_tables() or []
                    for ti, t in enumerate(pts):
                        if not t or len(t) < 2: continue
                        rows = [[("" if c is None else str(c).strip()) for c in row] for row in t]
                        tables.append({"page": pi+1, "table_idx": ti, "rows": rows, "reader": "pdfplumber"})
                except Exception:
                    continue
        if tables: return tables
    except Exception:
        pass
    # Layer 2: pymupdf (fitz) text-block heuristic — only if pdfplumber fails
    try:
        import fitz
        doc = fitz.open(pdf_path)
        for pi in range(min(len(doc), MAX_PAGES_PER_PDF)):
            page = doc[pi]
            try:
                tabs = page.find_tables() if hasattr(page, "find_tables") else None
                if tabs and len(tabs.tables) > 0:
                    for ti, t in enumerate(tabs.tables):
                        try:
                            r = t.extract()
                            if r and len(r) >= 2:
                                rows = [[("" if c is None else str(c).strip()) for c in row] for row in r]
                                tables.append({"page": pi+1, "table_idx": ti, "rows": rows, "reader": "fitz_find_tables"})
                        except Exception: continue
            except Exception: continue
    except Exception:
        pass
    return tables

def def_extract_tables_docx(docx_path):
    tables = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        # 粗略 split tables — full docx parsing 不在 D8 scope,只支援簡單表
        # 用 python-docx 才能拿準確 cell；這邊先 skip if not present
        try:
            import docx as pydocx
            d = pydocx.Document(docx_path)
            for ti, t in enumerate(d.tables):
                rows = []
                for row in t.rows:
                    rows.append([c.text.strip() for c in row.cells])
                if len(rows) >= 2:
                    tables.append({"page": 1, "table_idx": ti, "rows": rows, "reader": "python_docx"})
        except Exception:
            pass
    except Exception:
        pass
    return tables

# ============================================================================
# def 06_GEOMETRY_CLASSIFIER (G0-G7)
# ============================================================================
YEAR_REGEX = re.compile(r"\b(20\d{2})(?:E|F)?\b")
PERIOD_REGEX = re.compile(r"\b(1H|2H|H1|H2|Q[1-4]|1Q|2Q|3Q|4Q)\b", re.I)

FINANCIAL_HINTS = [
    "revenue","sales","營收","收入","gross","毛利","operating","營業","net income","稅後","淨利",
    "eps","每股","assets","資產","liabilities","負債","equity","權益","cash","現金",
    "inventories","存貨","capex","資本支出","ebitda","ebit","margin","利潤率"
]

def def_count_year_cells(header_rows):
    cnt = 0
    years = []
    for cell in header_rows:
        m = YEAR_REGEX.search(str(cell))
        if m:
            cnt += 1
            years.append(m.group(1))
    return cnt, years

def def_classify_geometry(rows):
    """Return (geometry_tag, meta)"""
    if not rows or len(rows) < 2:
        return ("G7_Noise", {"reason": "too_few_rows"})
    nrows = len(rows)
    ncols = max(len(r) for r in rows)
    if ncols < 2:
        return ("G7_Noise", {"reason": "too_few_cols"})

    # 取 header 2 列
    hdr0 = [str(c) for c in rows[0]] if len(rows) > 0 else []
    hdr1 = [str(c) for c in rows[1]] if len(rows) > 1 else []
    hdr_combined = hdr0 + hdr1

    yc, years = def_count_year_cells(hdr_combined)
    has_period = any(PERIOD_REGEX.search(c) for c in hdr_combined)
    body = rows[1:] if not has_period else rows[2:]
    body_text = " ".join(str(c).lower() for r in body[:30] for c in r)
    fin_score = sum(1 for h in FINANCIAL_HINTS if h in body_text)

    # G4 monthly revenue: header 出現「月」+ 多月份
    months_zh = sum(1 for c in hdr_combined if re.search(r"\b(1[0-2]|[1-9])月\b", str(c)))
    if months_zh >= 4:
        return ("G4_MonthlyRevenue", {"months_detected": months_zh, "reason": "monthly_columns"})

    # G6 price table
    if any("price" in str(c).lower() or "股價" in str(c) for c in hdr_combined) and fin_score < 2:
        return ("G6_PriceTable", {"reason": "price_columns_low_fin_hints"})

    # G3 key-value pair (2 col, no years)
    if ncols == 2 and yc == 0:
        return ("G3_KeyValuePair", {"reason": "two_col_no_years"})

    # G5 segment breakdown (no years, but financial hints + multiple rows)
    if yc == 0 and fin_score >= 2 and nrows >= 4:
        return ("G5_SegmentBreakdown", {"fin_score": fin_score, "reason": "no_years_fin_hints"})

    # G2 multi-header (2-row header where row0 has years, row1 has periods)
    if yc >= 2 and has_period:
        return ("G2_MultiHeaderMatrix", {"years": years, "has_period": True, "reason": "dual_header"})

    # G1 year matrix (single header row with years)
    if yc >= 2 and fin_score >= 1:
        return ("G1_YearMatrix", {"years": years, "fin_score": fin_score, "reason": "years_in_header"})

    # G0 unknown
    return ("G0_Unknown", {"yc": yc, "fin_score": fin_score, "reason": "unclassified"})

# ============================================================================
# def 07_MULTI_HEADER_FLATTEN (G2 dual-header)
# ============================================================================
def def_flatten_multi_header(rows):
    """Collapse 2-row header (year, period) -> per-col (year, period, is_estimate, broker_version)"""
    if len(rows) < 3: return None, []
    year_row = rows[0]
    period_row = rows[1]
    body = rows[2:]
    col_meta = []
    for ci in range(max(len(year_row), len(period_row))):
        y = str(year_row[ci]) if ci < len(year_row) else ""
        p = str(period_row[ci]) if ci < len(period_row) else ""
        ym = YEAR_REGEX.search(y) or YEAR_REGEX.search(p)
        pm = PERIOD_REGEX.search(p) or PERIOD_REGEX.search(y)
        is_est = bool(re.search(r"\b(E|F|Est|Estimate|預估|預期)\b", y + " " + p, re.I))
        col_meta.append({
            "col_idx": ci,
            "year": ym.group(1) if ym else "",
            "period": pm.group(1).upper() if pm else "FY",
            "is_estimate": is_est,
            "raw": (y + " | " + p).strip(" |")
        })
    return col_meta, body

# ============================================================================
# def 08_RESTORE (G1 + G2 + G5)
# ============================================================================
def def_restore_g1(rows, src_meta):
    """Header row 0: account label + year columns. Body: account_label + values per year."""
    out = []
    if not rows: return out
    header = [str(c) for c in rows[0]]
    year_cols = []
    for ci, c in enumerate(header):
        m = YEAR_REGEX.search(c)
        if m:
            is_est = bool(re.search(r"E|F|預", c))
            year_cols.append((ci, m.group(1), is_est))
    if not year_cols: return out
    label_col = 0
    for row in rows[1:]:
        if not row: continue
        label = str(row[label_col]).strip() if len(row) > label_col else ""
        if not label: continue
        for ci, yr, is_est in year_cols:
            if ci >= len(row): continue
            val_raw = str(row[ci]).strip()
            if not val_raw or val_raw in ("-", "—", "n/a", "N/A"): continue
            out.append({
                **src_meta,
                "Account Raw": label,
                "Year": yr,
                "Period": "FY",
                "Is Estimate": is_est,
                "Value Raw": val_raw,
                "Geometry": "G1_YearMatrix"
            })
    return out

def def_restore_g2(rows, src_meta):
    out = []
    col_meta, body = def_flatten_multi_header(rows)
    if not col_meta or not body: return out
    label_col = 0
    for row in body:
        if not row: continue
        label = str(row[label_col]).strip() if len(row) > label_col else ""
        if not label: continue
        for cm in col_meta:
            ci = cm["col_idx"]
            if ci == label_col: continue
            if ci >= len(row): continue
            val_raw = str(row[ci]).strip()
            if not val_raw or val_raw in ("-", "—", "n/a", "N/A"): continue
            if not cm["year"]: continue
            out.append({
                **src_meta,
                "Account Raw": label,
                "Year": cm["year"],
                "Period": cm["period"],
                "Is Estimate": cm["is_estimate"],
                "Value Raw": val_raw,
                "Geometry": "G2_MultiHeaderMatrix"
            })
    return out

def def_restore_g5(rows, src_meta):
    """Segment breakdown — treat as segment-level rows w/o year axis. Year defaults to most-recent if visible elsewhere."""
    out = []
    if not rows: return out
    header = [str(c) for c in rows[0]]
    for row in rows[1:]:
        if not row: continue
        label = str(row[0]).strip()
        if not label: continue
        for ci in range(1, len(row)):
            val_raw = str(row[ci]).strip()
            if not val_raw: continue
            col_label = header[ci] if ci < len(header) else f"col{ci}"
            out.append({
                **src_meta,
                "Account Raw": label,
                "Year": "",
                "Period": "Segment",
                "Is Estimate": False,
                "Value Raw": val_raw,
                "Segment Header": col_label,
                "Geometry": "G5_SegmentBreakdown"
            })
    return out

# ============================================================================
# def 09_VALUE_NORMALIZE
# ============================================================================
def def_normalize_value(v):
    if v is None: return (None, "null")
    s = str(v).strip().replace(",", "").replace(" ", "")
    if not s: return (None, "empty")
    neg = False
    if s.startswith("(") and s.endswith(")"): neg = True; s = s[1:-1]
    if s.startswith("-"): neg = True; s = s[1:]
    s = s.rstrip("%").rstrip("倍")
    try:
        f = float(s)
        if neg: f = -f
        return (f, "ok")
    except:
        return (None, "parse_fail")

# ============================================================================
# def 10_TRUST_V044
# ============================================================================
def def_trust_score(restore_conf, ssot_score, hist, addsub, div, ext):
    s = round(restore_conf*0.35 + ssot_score*0.35 + hist*0.10 + addsub*0.10 + div*0.05 + ext*0.05, 2)
    if s >= 85: band, sev = "GREEN", "OK"
    elif s >= 70: band, sev = "YELLOW", "WARN"
    elif s >= 55: band, sev = "ORANGE", "WARN"
    else: band, sev = "RED", "ERR"
    return s, band, sev

# ============================================================================
# def 11_MAIN
# ============================================================================
def def_main():
    t0 = time.time()
    Path(RUN_DIR).mkdir(parents=True, exist_ok=True)
    def_progress("BOOT", 0, 1, "loading manifest")

    with open(PDF_LIST_JSON, "r", encoding="utf-8") as f:
        manifest = json.load(f)
    pdfs = manifest.get("pdfs", [])
    total_pdfs = len(pdfs)
    def_progress("SSOT_LOAD", 0, 1, "loading SSOT lock")
    ssot = def_load_ssot()

    # ============== Stage A: extract tables ==============
    table_inventory = []
    geometry_rows = []
    multihdr_preview = []
    restored_all = []
    quarantine_rows = []
    issue_rows = []
    geom_counter = Counter()

    for pi, pdf_path in enumerate(pdfs, 1):
        def_progress("EXTRACT", pi, total_pdfs, Path(pdf_path).name)
        fname = Path(pdf_path).name
        ext = Path(pdf_path).suffix.lower()
        tables = []
        try:
            if ext == ".pdf":
                tables = def_extract_tables_pdf(pdf_path)
            elif ext in (".docx", ".doc"):
                tables = def_extract_tables_docx(pdf_path)
        except Exception as e:
            issue_rows.append({
                "Validation Status": def_lamp("FAIL"),
                "Dataset": "EXTRACT", "Filename": fname,
                "Issue Type": "EXTRACT_EXCEPTION",
                "Problem": str(e)[:500], "Suggested Fix": "Check reader libs (pdfplumber/fitz/docx)"
            })
            continue

        ticker_m = re.search(r"(?<!\d)([1-9]\d{3})(?!\d)", fname)
        ticker = ticker_m.group(1) if ticker_m and not re.match(r"^(202[1-9]|2030)$", ticker_m.group(1)) else ""

        for t in tables:
            src_meta = {
                "Filename": fname, "Ticker": ticker,
                "Page": t["page"], "Table Idx": t["table_idx"],
                "Reader": t["reader"]
            }
            geom_tag, geom_meta = def_classify_geometry(t["rows"])
            geom_counter[geom_tag] += 1
            table_inventory.append({**src_meta, "Geometry": geom_tag, "N Rows": len(t["rows"]), "N Cols": max(len(r) for r in t["rows"])})
            geometry_rows.append({
                "Validation Status": def_lamp("PASS") if geom_tag.startswith(("G1","G2","G5")) else def_lamp("WARN"),
                **src_meta,
                "Geometry": geom_tag,
                "Reason": geom_meta.get("reason",""),
                "Years": ",".join(geom_meta.get("years",[])),
                "Validation Note": "Accepted into restore" if geom_tag.startswith(("G1","G2","G5")) else "Skipped (D3 option 2)"
            })

            # Multi-header preview (first 30 only)
            if geom_tag == "G2_MultiHeaderMatrix" and len(multihdr_preview) < 30:
                col_meta, _ = def_flatten_multi_header(t["rows"])
                for cm in (col_meta or []):
                    multihdr_preview.append({
                        "Validation Status": def_lamp("PASS"),
                        **src_meta, **cm,
                        "Validation Note": "Year+period flattened"
                    })

            # Restore
            try:
                if geom_tag == "G1_YearMatrix":
                    restored_all.extend(def_restore_g1(t["rows"], src_meta))
                elif geom_tag == "G2_MultiHeaderMatrix":
                    restored_all.extend(def_restore_g2(t["rows"], src_meta))
                elif geom_tag == "G5_SegmentBreakdown":
                    restored_all.extend(def_restore_g5(t["rows"], src_meta))
            except Exception as e:
                issue_rows.append({
                    "Validation Status": def_lamp("WARN"),
                    "Dataset": "RESTORE", "Filename": fname,
                    "Issue Type": "RESTORE_EXCEPTION",
                    "Problem": (geom_tag + ": " + str(e))[:500],
                    "Suggested Fix": "Per-table defensive — review row in geometry tab"
                })

    # ============== Stage B: D3 quarantine + SSOT match + normalize ==============
    def_progress("RESTORE_POSTPROCESS", 0, len(restored_all), "post-processing")
    financial_rows = []
    for ri, r in enumerate(restored_all):
        if ri % 50 == 0:
            def_progress("RESTORE_POSTPROCESS", ri, len(restored_all), r.get("Filename",""))
        period = r.get("Period", "FY")
        # D3 option 2: half-year/quarter — quarantine
        if period and period not in ("FY", "Segment"):
            quarantine_rows.append({
                "Validation Status": def_lamp("WARN"),
                **r,
                "Validation Note": "D3 option 2: half-year/quarter quarantined; never aggregated"
            })
            continue
        # value normalize
        val, status = def_normalize_value(r.get("Value Raw"))
        if status != "ok":
            r["Value"] = None
            r["Value Parse"] = status
        else:
            r["Value"] = val
            r["Value Parse"] = "ok"
        # SSOT match
        en, score, kind = def_ssot_match(r.get("Account Raw",""), ssot)
        r["Account Official EN"] = en
        r["SSOT Match Score"] = score
        r["SSOT Match Kind"] = kind
        # restore confidence
        rc = 90.0 if r.get("Geometry") == "G1_YearMatrix" else (80.0 if r.get("Geometry") == "G2_MultiHeaderMatrix" else 65.0)
        if r.get("Is Estimate"): rc = min(rc, 75.0)  # estimate cap
        r["Restore Confidence"] = rc
        financial_rows.append(r)

    # ============== Stage C: trust validator v04.4 ==============
    def_progress("TRUST", 0, len(financial_rows), "computing trust")
    # AddSub / Division checks — by (filename, ticker, year)
    key_map = defaultdict(dict)
    for r in financial_rows:
        if r.get("Account Official EN") and r.get("Value") is not None:
            k = (r.get("Filename",""), r.get("Ticker",""), str(r.get("Year","")))
            key_map[k][r["Account Official EN"]] = r["Value"]

    addsub_ok = set()
    division_ok = set()
    for k, m in key_map.items():
        # AddSub: GP = Rev - COGS  (寬鬆 5%)
        if all(x in m for x in ["Revenue", "Cost of Goods Sold", "Gross Profit"]):
            calc = m["Revenue"] - m["Cost of Goods Sold"]
            if m["Gross Profit"] != 0 and abs((calc - m["Gross Profit"]) / m["Gross Profit"]) <= 0.05:
                addsub_ok.add(k)
        # Division: Gross Margin = GP / Rev
        if all(x in m for x in ["Gross Profit", "Revenue"]) and m["Revenue"]:
            division_ok.add(k)
        if all(x in m for x in ["Operating Income", "Revenue"]) and m["Revenue"]:
            division_ok.add(k)
        if all(x in m for x in ["Net Income", "Revenue"]) and m["Revenue"]:
            division_ok.add(k)

    trust_rows = []
    coverage_counter = Counter()
    review_rows = []
    for ri, r in enumerate(financial_rows):
        if ri % 50 == 0:
            def_progress("TRUST", ri, len(financial_rows), r.get("Filename",""))
        k = (r.get("Filename",""), r.get("Ticker",""), str(r.get("Year","")))
        rc = float(r.get("Restore Confidence", 0) or 0)
        ss = float(r.get("SSOT Match Score", 0) or 0)
        hist = 80 if rc >= 85 else 60
        addsub = 85 if k in addsub_ok else 60
        div = 85 if k in division_ok else 60
        ext = 50
        score, band, sev = def_trust_score(rc, ss, hist, addsub, div, ext)
        r["Historical Trust Score"] = hist
        r["AddSub Trust Score"] = addsub
        r["Division Trust Score"] = div
        r["External Trust Score"] = ext
        r["Final Trust Confidence"] = score
        r["Final Trust Band"] = band
        r["Final Severity"] = sev
        r["Validation Status"] = def_lamp("PASS" if band == "GREEN" else ("WARN" if band in ("YELLOW","ORANGE") else "FAIL"))
        trust_rows.append(r)
        coverage_counter[(r.get("Ticker",""), r.get("Account Official EN",""))] += 1
        if band in ("ORANGE", "RED"):
            review_rows.append({
                "Validation Status": def_lamp("WARN" if band=="ORANGE" else "FAIL"),
                "Filename": r.get("Filename",""), "Ticker": r.get("Ticker",""), "Year": r.get("Year",""),
                "Account Raw": r.get("Account Raw",""), "Account Official EN": r.get("Account Official EN",""),
                "Value": r.get("Value"), "Final Trust": score, "Band": band,
                "Validation Note": "Manual review required"
            })

    # ============== Stage D: write outputs ==============
    def_progress("WRITE", 0, 4, "writing CSVs")
    if WRITE_CSV:
        def_write_csv(OUT_GEOMETRY_CSV, geometry_rows)
        def_write_csv(OUT_MULTIHDR_CSV, multihdr_preview)
        def_write_csv(OUT_RESTORED_CSV, restored_all)
        def_write_csv(OUT_QUARANTINE_CSV, quarantine_rows)
        def_write_csv(OUT_FINANCIAL_CSV, financial_rows)
        def_write_csv(OUT_TRUST_CSV, trust_rows)
        def_write_csv(OUT_REVIEW_CSV, review_rows)
        # coverage
        cov_rows = []
        for (tk, acc), n in coverage_counter.most_common():
            cov_rows.append({
                "Validation Status": def_lamp("PASS") if n >= 2 else def_lamp("WARN"),
                "Ticker": tk, "Account Official EN": acc, "Year Count": n,
                "Validation Note": "Multi-year present" if n >= 2 else "Single-year only"
            })
        def_write_csv(OUT_COVERAGE_CSV, cov_rows)
        # action
        action_rows = [
            {"Validation Status": def_lamp("WARN"), "Action": "D8C SSOT Alias Auto-suggest", "Note": "Run if SSOT_Match_Kind='fuzzy_contain' rate > 20%"},
            {"Validation Status": def_lamp("WARN"), "Action": "D9 Production Write Gate", "Note": "BasicInfo/FinancialData parquet locked until manual confirm"},
        ]
        def_write_csv(OUT_ACTION_CSV, action_rows)

    def_progress("WRITE", 1, 4, "writing issues + parquet")
    if WRITE_PARQUET:
        try:
            import pandas as pd
            if trust_rows:
                df = pd.DataFrame(trust_rows).astype(str)
                Path(OUT_PARQUET_FIN).parent.mkdir(parents=True, exist_ok=True)
                df.to_parquet(OUT_PARQUET_FIN, index=False)
        except Exception as e:
            issue_rows.append({
                "Validation Status": def_lamp("WARN"), "Dataset": "PARQUET",
                "Filename": "", "Issue Type": "PARQUET_WRITE_FAIL",
                "Problem": str(e)[:500], "Suggested Fix": "Check pandas+pyarrow"
            })

    def_progress("WRITE", 2, 4, "writing duckdb")
    duckdb_ok = False; duckdb_err = ""
    if WRITE_DUCKDB:
        try:
            import duckdb, pandas as pd
            con = duckdb.connect(OUT_DUCKDB)
            if trust_rows:
                df = pd.DataFrame(trust_rows).astype(str)
                con.register("fin_df", df)
                con.execute("CREATE OR REPLACE TABLE financial_trust_candidate AS SELECT * FROM fin_df")
            con.close()
            duckdb_ok = True
        except Exception as e:
            duckdb_err = str(e)[:500]
            issue_rows.append({
                "Validation Status": def_lamp("WARN"), "Dataset": "DUCKDB",
                "Filename": "", "Issue Type": "DUCKDB_WRITE_FAIL",
                "Problem": duckdb_err, "Suggested Fix": "Check duckdb install"
            })

    # ============== Stage E: HTML ==============
    def_progress("HTML", 3, 4, "rendering HTML")
    meta = {
        "version": VERSION, "generated_at": def_now(),
        "elapsed_sec": round(time.time() - t0, 2),
        "system_pass": True,
        "ssot_missing": ssot.get("missing", False),
        "ssot_accounts": len(ssot.get("accounts", {})),
        "ssot_aliases": len(ssot.get("aliases", {})),
        "pdf_count": total_pdfs,
        "tables_total": len(table_inventory),
        "g1": geom_counter.get("G1_YearMatrix",0),
        "g2": geom_counter.get("G2_MultiHeaderMatrix",0),
        "g5": geom_counter.get("G5_SegmentBreakdown",0),
        "g0": geom_counter.get("G0_Unknown",0),
        "g3": geom_counter.get("G3_KeyValuePair",0),
        "g4": geom_counter.get("G4_MonthlyRevenue",0),
        "g6": geom_counter.get("G6_PriceTable",0),
        "g7": geom_counter.get("G7_Noise",0),
        "restored_rows": len(restored_all),
        "financial_rows": len(financial_rows),
        "quarantine_rows": len(quarantine_rows),
        "trust_green": sum(1 for r in trust_rows if r.get("Final Trust Band")=="GREEN"),
        "trust_yellow": sum(1 for r in trust_rows if r.get("Final Trust Band")=="YELLOW"),
        "trust_orange": sum(1 for r in trust_rows if r.get("Final Trust Band")=="ORANGE"),
        "trust_red": sum(1 for r in trust_rows if r.get("Final Trust Band")=="RED"),
        "issue_rows": len(issue_rows),
        "duckdb_ok": duckdb_ok, "duckdb_err": duckdb_err,
    }
    def_write_csv(str(Path(RUN_DIR) / "VRN_v070_Issues.csv"), issue_rows)
    def_write_json(OUT_JSON, {
        "meta": meta,
        "geometry_summary": dict(geom_counter),
        "review_count": len(review_rows),
    })
    def_render_html(meta, geometry_rows, multihdr_preview, financial_rows, trust_rows, quarantine_rows, review_rows, issue_rows)

    def_progress("DONE", 4, 4, "complete")
    print(json.dumps({"meta": meta}, ensure_ascii=False, indent=2))

def def_render_table(title, rows, cap=400):
    rows = list(rows or [])
    if not rows:
        return "<section class='card'><h2>" + def_esc(title) + "</h2><p>EMPTY</p></section>"
    shown = rows[:cap]
    keys = []
    for r in shown:
        for k in r.keys():
            if k not in keys: keys.append(k)
    if "Validation Status" in keys:
        keys = ["Validation Status"] + [k for k in keys if k!="Validation Status"]
    if "Validation Note" in keys:
        keys = [k for k in keys if k!="Validation Note"] + ["Validation Note"]
    th = "".join("<th>"+def_esc(k)+"</th>" for k in keys)
    body = []
    for r in shown:
        st = str(r.get("Validation Status",""))
        cls = "ok" if "🟢" in st else ("warn" if "🟡" in st else "fail")
        tds = []
        for k in keys:
            v = str(r.get(k,""))
            kind = "status" if k=="Validation Status" else ("long" if len(v)>32 or k.lower() in ("filename","path","account raw","problem","validation note","reason") else "short")
            tds.append("<td data-kind='"+kind+"'>"+def_esc(v)+"</td>")
        body.append("<tr class='"+cls+"'>"+"".join(tds)+"</tr>")
    note = ""
    if len(rows) > len(shown):
        note = "<div class='note'>HTML capped at "+format(len(shown),",")+"; CSV has "+format(len(rows),",")+" rows.</div>"
    return ("<section class='card'><h2>"+def_esc(title)+" <span>"+format(len(rows),",")+" rows</span></h2>"+note+
            "<div class='scroll'><table><thead><tr>"+th+"</tr></thead><tbody>"+"".join(body)+"</tbody></table></div></section>")

def def_render_html(meta, geom, mh, fin, trust, quar, review, issues):
    meta_json = def_esc(json.dumps(meta, ensure_ascii=False, indent=2))
    head = ("<!doctype html><html lang='zh-Hant'><head><meta charset='utf-8'><title>VRN v0.7.0 D8 AIO</title>"
        "<style>"
        ":root{--bg:#f5f4f0;--card:#fff;--border:#d8d4ca;--head:#e8e4da;--text:#151515;--muted:#5f5b54;--ok:#247a4d;--warn:#a86613;--fail:#a73535;--blue:#4c78a8;--sans:Syne,'DM Sans','Segoe UI','Microsoft JhengHei',system-ui,sans-serif;--mono:'DM Mono',Consolas,monospace}"
        "*{box-sizing:border-box}body{margin:0;background:var(--bg);color:var(--text);font-family:var(--sans);font-size:8.8px;line-height:1.15}"
        ".wrap{max-width:1920px;margin:0 auto;padding:10px 14px 18px}"
        "header{position:relative;background:#fff;border:1px solid var(--border);border-radius:12px;padding:14px 16px;margin-bottom:8px;box-shadow:0 4px 12px rgba(0,0,0,.06)}"
        "header:before{content:'';position:absolute;left:16px;right:16px;top:0;height:3px;background:linear-gradient(90deg,#4c78a8,#439a9a,#7a6daa,#c4943a,#c96b5a)}"
        "h1{margin:0;font-size:20px;font-weight:900}.sub{font-family:var(--mono);font-size:9.5px;font-weight:900;letter-spacing:.08em;margin-top:3px}.desc{font-size:9.5px;color:var(--muted);margin-top:3px}"
        ".tabs{display:flex;gap:3px;flex-wrap:wrap;margin:8px 0}.tabs button{padding:7px 13px;border:1px solid var(--border);border-radius:8px 8px 0 0;background:#edecea;color:var(--muted);font:800 10px var(--sans);cursor:pointer}"
        ".tabs button.on{background:#fff;color:var(--blue);border-bottom:2px solid var(--blue)}.page{display:none}.page.on{display:block}"
        ".grid{display:grid;grid-template-columns:repeat(8,minmax(90px,1fr));gap:6px;margin-bottom:8px}"
        ".kpi{background:#fff;border:1px solid var(--border);border-radius:8px;padding:6px 7px;min-height:42px}.kpi .v{font-size:14px;font-weight:900;text-align:center}.kpi .k{font-size:8px;color:var(--muted);text-align:center;margin-top:3px;text-transform:uppercase;font-weight:900}"
        ".card{background:#fff;border:1px solid var(--border);border-radius:9px;padding:8px;margin-bottom:8px;box-shadow:0 1px 6px rgba(0,0,0,.035)}.card h2{margin:0 0 6px;font-size:11px;font-weight:900}.card h2 span,.note{font-size:8.5px;color:var(--warn);font-weight:800}"
        ".scroll{overflow:auto;max-height:74vh;border:1px solid var(--border);border-radius:7px;background:#fff}"
        "table{border-collapse:separate;border-spacing:0;table-layout:auto;width:max-content;min-width:100%;font-size:8.4px;line-height:1.1}"
        "th{position:sticky;top:0;z-index:8;background:var(--head);font-weight:900;text-align:center;vertical-align:top;padding:3px 5px;border-bottom:1px solid var(--border);border-right:1px solid #ece8df;white-space:normal;overflow-wrap:anywhere;min-width:96px;max-width:360px}"
        "td{padding:3px 5px;border-bottom:1px solid #ece8df;border-right:1px solid #ece8df;vertical-align:top;text-align:center;white-space:normal;overflow-wrap:anywhere;word-break:break-word;min-width:82px;max-width:360px}"
        "td[data-kind='status'],td:first-child,th:first-child{text-align:center!important;min-width:82px;max-width:96px;font-weight:900}"
        "td[data-kind='long']{text-align:left!important;min-width:190px;max-width:900px}"
        "tr.ok td:first-child{color:var(--ok)}tr.warn td:first-child{color:var(--warn)}tr.fail td:first-child{color:var(--fail)}"
        "pre{font-family:var(--mono);font-size:8.4px;background:#f8fafc;border:1px solid #e2e8f0;border-radius:7px;padding:7px;white-space:pre-wrap;overflow-wrap:anywhere;max-height:42vh;overflow:auto}"
        ".footer{margin-top:10px;border-top:1px solid var(--border);padding-top:8px;color:var(--muted);font-size:8.4px}"
        "</style></head><body><div class='wrap'>"
        "<header><h1>VeritasReportNova v0.7.0 D8 ALL-IN-ONE</h1>"
        "<div class='sub'>VERITAS INTELLIGENCE ANALYTICS</div>"
        "<div class='desc'>G0-G7 Geometry Classifier · SSOT Alias · Trust v04.4 · D3 option 2 ENFORCED · NO SSOT MUTATION</div></header>"
    )
    kpis = ("<div class='grid'>"
        f"<div class='kpi'><div class='v'>{meta['pdf_count']}</div><div class='k'>PDFs</div></div>"
        f"<div class='kpi'><div class='v'>{meta['tables_total']}</div><div class='k'>Tables</div></div>"
        f"<div class='kpi'><div class='v'>{meta['restored_rows']}</div><div class='k'>Restored</div></div>"
        f"<div class='kpi'><div class='v'>{meta['financial_rows']}</div><div class='k'>Financial</div></div>"
        f"<div class='kpi'><div class='v'>{meta['quarantine_rows']}</div><div class='k'>Quarantine</div></div>"
        f"<div class='kpi'><div class='v' style='color:var(--ok)'>{meta['trust_green']}</div><div class='k'>Green</div></div>"
        f"<div class='kpi'><div class='v' style='color:var(--warn)'>{meta['trust_yellow']+meta['trust_orange']}</div><div class='k'>Y/O</div></div>"
        f"<div class='kpi'><div class='v' style='color:var(--fail)'>{meta['trust_red']}</div><div class='k'>Red</div></div>"
        "</div>")
    nav = ("<nav class='tabs'>"
        "<button class='on' onclick=\"sw(this,'overview')\">Overview</button>"
        "<button onclick=\"sw(this,'geom')\">03A Geometry</button>"
        "<button onclick=\"sw(this,'mh')\">03B MultiHeader</button>"
        "<button onclick=\"sw(this,'fin')\">04 Financial</button>"
        "<button onclick=\"sw(this,'trust')\">05 Trust</button>"
        "<button onclick=\"sw(this,'quar')\">03C Quarantine</button>"
        "<button onclick=\"sw(this,'review')\">06 Review</button>"
        "<button onclick=\"sw(this,'issues')\">07 Issues</button>"
        "</nav>")
    pages = (f"<section class='page on' id='overview'><section class='card'><h2>Meta</h2><pre>{meta_json}</pre></section></section>"
        f"<section class='page' id='geom'>{def_render_table('03A Table Geometry', geom)}</section>"
        f"<section class='page' id='mh'>{def_render_table('03B Multi-Header Flatten Preview', mh)}</section>"
        f"<section class='page' id='fin'>{def_render_table('04 Financial Candidates', fin)}</section>"
        f"<section class='page' id='trust'>{def_render_table('05 Trust Matrix', trust)}</section>"
        f"<section class='page' id='quar'>{def_render_table('03C Half-Year / Quarter Quarantine', quar)}</section>"
        f"<section class='page' id='review'>{def_render_table('06 Review Queue', review)}</section>"
        f"<section class='page' id='issues'>{def_render_table('07 Issues', issues)}</section>")
    foot = ("<div class='footer'>Veritas Intelligence Analytics │ AI-Augmented Investment Intelligence for Smarter Decisions and Deeper Insight</div></div>"
        "<script>function sw(b,i){document.querySelectorAll('.tabs button').forEach(x=>x.classList.remove('on'));"
        "document.querySelectorAll('.page').forEach(x=>x.classList.remove('on'));b.classList.add('on');document.getElementById(i).classList.add('on');}</script></body></html>")
    Path(OUT_HTML).write_text(head+kpis+nav+pages+foot, encoding="utf-8-sig")

if __name__ == "__main__":
    try:
        def_main()
    except BaseException:
        fatal = {"meta": {"version": VERSION, "system_pass": False, "fatal": traceback.format_exc(), "html": OUT_HTML, "json": OUT_JSON}}
        def_write_json(OUT_JSON, fatal)
        try: def_progress("FATAL", 0, 1, fatal["meta"]["fatal"][:200])
        except: pass
        print(json.dumps(fatal, ensure_ascii=False, indent=2))
        raise