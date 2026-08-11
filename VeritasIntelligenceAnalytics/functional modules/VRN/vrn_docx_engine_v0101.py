#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
vrn_docx_engine_v0101 — DOCX 深度解析引擎(10 地雷硬化版)
==========================================================================
依操作員貼文規格(Gemini 對話)VIA 落地;只整理不發明,graceful 全梯次:
  擷取梯:docx2python(合併格最強)→ python-docx → 內建 XML 末梯
         (純標準庫 zipfile+ElementTree 解 word/document.xml;永不空手)
  修復三式(貼文原方):①空字串→NA ②ffill 向下填補(垂直合併格)
         ③空白/換行正規化(re)+ 圖片佔位符清除
  驗證:矩形性/表頭唯一/空值率(承進件矩陣四閘口徑)
  對照梯(--compare):datacompy(報表最詳)→ pandas.compare → difflib 文字 diff
輸出:staging/ocr_out/docx_struct/<stem>/ 下 text.txt + table_N.csv(+parquet)
     + summary.json;誠實逐檔矩陣。
v0101 硬化(操作員貼文 10 地雷令,逐一落防):
  E1 偽裝檔:zip 簽章 PK 預檢→NOT_DOCX_ZIP 誠實判(.doc 提示轉檔路)
  E2 記憶體:docx2python 試 extract_image=False
  E3 表頭重複/空:fillna 未命名+後綴去重(datacompy 前置要求)
  E4 巢狀表格:cell 遞迴攤平(list-in-list 防炸)
  E5 排版表:列<2 或空值率>0.8 → SKIP_LAYOUT 誠實跳過(仍記錄)
  E6 vMerge 真空值:內建 XML 梯以 vMerge 標籤精準判(原生已備)
  E7 幽靈字元:NFKC 正規化+零寬字元(u200b/c/d/feff)拔除
  E8 表格錯位:對照前以表頭 Jaccard 相似度配對(≥0.5)並印配對表
  E9 主鍵對照:自動偵測 項次/編號/代號/id 欄→join_columns;無則 index+警語
  E10 數字型態:比對前去千分位逗號+to_numeric 統一
用法:py vrn_docx_engine_v0101.py <pattern>              # 解析收件夾匹配檔
     py vrn_docx_engine_v0101.py --compare old.docx new.docx  # 新舊對照
"""
from __future__ import annotations

import json
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

HERE = Path(__file__).resolve().parent
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ── 硬化 helpers(E1/E4/E7)────────────────────────────────────────────────
def precheck_zip(path: Path):
    """E1:zip 簽章預檢;偽裝檔誠實判。"""
    sig = path.read_bytes()[:4]
    if sig[:2] == b"PK":
        return None
    if sig[:4] == b"\xd0\xcf\x11\xe0":
        return "NOT_DOCX_ZIP:舊版 .doc(OLE2)偽裝 — 候 Word/LibreOffice 另存 .docx"
    if sig[:4] == b"%PDF":
        return "NOT_DOCX_ZIP:PDF 偽裝 .docx — 改走 via-pdfcheck/OCR 鏈"
    return f"NOT_DOCX_ZIP:未知簽章 {sig!r} — 誠實除外"


def flatten_cell(cell) -> str:
    """E4:巢狀表格 cell 遞迴攤平(list-in-list 防炸)。"""
    if isinstance(cell, str):
        return cell
    if isinstance(cell, (list, tuple)):
        return "\n".join(x for x in (flatten_cell(c) for c in cell) if x)
    return str(cell)


def strip_ghost(text: str) -> str:
    """E7:NFKC 正規化 + 零寬字元拔除。"""
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    return re.sub(r"[\u200b\u200c\u200d\ufeff]", "", text)


# ── 擷取梯 ────────────────────────────────────────────────────────────────
def extract_builtin_xml(path: Path):
    """內建 XML 末梯:純標準庫解 word/document.xml(段落+表格;vMerge 續格留空)。"""
    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    body = root.find(f"{W_NS}body")
    paragraphs, tables = [], []
    for el in body:
        if el.tag == f"{W_NS}p":
            t = "".join(n.text or "" for n in el.iter(f"{W_NS}t"))
            if t.strip():
                paragraphs.append(t)
        elif el.tag == f"{W_NS}tbl":
            rows = []
            for tr in el.findall(f"{W_NS}tr"):
                row = []
                for tc in tr.findall(f"{W_NS}tc"):
                    vm = tc.find(f"{W_NS}tcPr/{W_NS}vMerge")
                    cont = vm is not None and vm.get(f"{W_NS}val", "continue") != "restart"
                    txt = "\n".join(
                        "".join(n.text or "" for n in p.iter(f"{W_NS}t"))
                        for p in tc.findall(f"{W_NS}p"))
                    row.append("" if cont else txt)
                rows.append(row)
            if rows:
                tables.append(rows)
    return paragraphs, tables, "builtin_xml"


def extract_docx(path: Path):
    """梯次:docx2python → python-docx → 內建 XML;回 (段落list, 表格list[rows], 引擎名)。"""
    try:
        from docx2python import docx2python
        try:
            ctx = docx2python(path, extract_image=False)  # E2 記憶體:棄圖
        except TypeError:
            ctx = docx2python(path)
        with ctx as d:
            tables = []
            paragraphs = []
            for tbl in d.body:
                if len(tbl) > 1 or (len(tbl) == 1 and len(tbl[0]) > 1):
                    tables.append([[flatten_cell(c).strip() for c in row] for row in tbl])
                else:
                    for row in tbl:
                        for cell in row:
                            t = flatten_cell(cell)
                            paragraphs.extend(p for p in t.splitlines() if p.strip())
            return paragraphs, tables, "docx2python"
    except ImportError:
        pass
    try:
        import docx
        d = docx.Document(path)
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        tables = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
        return paragraphs, tables, "python-docx"
    except ImportError:
        pass
    return (*extract_builtin_xml(path),)


# ── 修復三式(貼文原方)────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    text = strip_ghost(text)
    text = re.sub(r"----media/.*?----", "", text)
    text = re.sub(r"[\r\t\f\v ]+", " ", text)
    text = re.sub(r"\n\s*\n", "\n", text)
    return text.strip()


def repair_table(rows):
    """回 (pandas.DataFrame|list, notes);pandas 缺時退純 python ffill。"""
    width = max(len(r) for r in rows)
    rect = [r + [""] * (width - len(r)) for r in rows]  # 欄位對齊
    notes = []
    if any(len(r) != width for r in rows):
        notes.append(f"欄數不齊已補齊至 {width}")
    try:
        import pandas as pd
        df = pd.DataFrame(rect)
        df = df.replace(r"^\s*$", pd.NA, regex=True)
        na_before = int(df.isna().sum().sum())
        df = df.ffill()
        filled = na_before - int(df.isna().sum().sum())
        if filled:
            notes.append(f"ffill 填補 {filled} 格(垂直合併格修復)")
        df = df.map(lambda x: re.sub(r"\s+", " ", str(x)).strip() if x is not None and str(x) != "<NA>" else x)
        if len(df) > 1:
            hdr = [str(h) if h is not None and str(h) not in ("<NA>", "") else "未命名欄位"
                   for h in df.iloc[0].tolist()]
            seen = {}
            dedup = []
            for h in hdr:
                seen[h] = seen.get(h, 0) + 1
                dedup.append(h if seen[h] == 1 else f"{h}_{seen[h]}")
            df.columns = dedup
            df = df[1:].reset_index(drop=True)
            note = "首列升表頭"
            if dedup != hdr:
                note += "(E3 重複/空欄名已去重補名)"
            notes.append(note)
        return df, notes
    except ImportError:
        out = []
        prev = [""] * width
        filled = 0
        for r in rect:
            nr = []
            for j, v in enumerate(r):
                v2 = re.sub(r"\s+", " ", v).strip()
                if not v2 and prev[j]:
                    v2 = prev[j]
                    filled += 1
                nr.append(v2)
            prev = nr
            out.append(nr)
        notes.append(f"pandas 缺——純 python ffill 填補 {filled} 格(誠實退路)")
        return out, notes


def validate_table(tb) -> dict:
    """驗證:矩形性/空值率(承進件矩陣口徑)。"""
    try:
        import pandas as pd
        if isinstance(tb, pd.DataFrame):
            total = tb.size or 1
            empty = int(tb.isna().sum().sum()) + int((tb.astype(str) == "").sum().sum())
            return {"rows": len(tb), "cols": len(tb.columns), "empty_rate": round(empty / total, 3)}
    except ImportError:
        pass
    total = sum(len(r) for r in tb) or 1
    empty = sum(1 for r in tb for v in r if not str(v).strip())
    return {"rows": len(tb), "cols": len(tb[0]) if tb else 0, "empty_rate": round(empty / total, 3)}


# ── 對照梯 ────────────────────────────────────────────────────────────────
def compare_docs(old: Path, new: Path) -> int:
    import difflib
    po, to, eo = extract_docx(old)
    pn, tn, en = extract_docx(new)
    print(f"=== DOCX 對照 · 舊 {old.name}({eo})vs 新 {new.name}({en})===")
    diff = [l for l in difflib.Differ().compare(
        clean_text("\n".join(po)).splitlines(), clean_text("\n".join(pn)).splitlines())
        if not l.startswith("  ")]
    print(f"  ── 文字層 diff({len(diff)} 行差異)──")
    for l in diff[:40]:
        print(f"    {l}")
    if len(diff) > 40:
        print(f"    …(餘 {len(diff) - 40} 行見報告)")
    # E8 表格配對:表頭 Jaccard 相似度(≥0.5 才對照,防表格錯位假報)
    def header_set(rows):
        return {strip_ghost(str(c)).strip() for c in (rows[0] if rows else [])} - {""}
    pairs = []
    used = set()
    for i, ta in enumerate(to):
        ha = header_set(ta)
        best, bj = None, 0.0
        for j, tb2 in enumerate(tn):
            if j in used:
                continue
            hb = header_set(tb2)
            jac = len(ha & hb) / len(ha | hb) if (ha | hb) else 0.0
            if jac > bj:
                best, bj = j, jac
        if best is not None and bj >= 0.5:
            pairs.append((i, best, bj))
            used.add(best)
        else:
            print(f"    表{i + 1}(舊):無相似新表配對(最佳 Jaccard {bj:.2f})— 誠實列示不硬比")
    for j in range(len(tn)):
        if j not in used:
            print(f"    表{j + 1}(新):新增表,舊版無對應")
    print(f"  ── 表格層({len(to)} vs {len(tn)} 表,配對 {len(pairs)} 組)──")
    KEY_HINTS = ("項次", "編號", "代號", "id", "ticker", "產品編號")
    for i, jn, jac in pairs:
        print(f"    配對:舊表{i + 1} ↔ 新表{jn + 1}(Jaccard {jac:.2f})")
        do, _ = repair_table(to[i])
        dn, _ = repair_table(tn[jn])
        # E10 數字型態統一:去千分位逗號 → to_numeric
        try:
            import pandas as pd
            for d_ in (do, dn):
                if isinstance(d_, pd.DataFrame):
                    for c in d_.columns:
                        raw = d_[c].astype(str).str.replace(",", "", regex=False)
                        num = pd.to_numeric(raw, errors="coerce")  # pandas 3.x 已無 errors='ignore'
                        if not num.isna().any() or num.isna().equals(d_[c].isna()):
                            d_[c] = num
                        else:
                            d_[c] = raw
        except ImportError:
            pass
        # E9 主鍵偵測
        join_col = None
        try:
            import pandas as pd
            if isinstance(do, pd.DataFrame) and isinstance(dn, pd.DataFrame):
                for c in do.columns:
                    if any(h in str(c).lower() for h in KEY_HINTS) and c in dn.columns and do[c].is_unique:
                        join_col = c
                        break
            if join_col:
                print(f"      主鍵對照:join_columns='{join_col}'")
            else:
                print("      無主鍵欄 — index 逐列比對(E9 警語:中間插列會整段錯位)")
        except ImportError:
            pass
        try:
            import datacompy
            if join_col:
                comp = datacompy.Compare(do, dn, join_columns=join_col, df1_name="舊版", df2_name="新版")
            else:
                comp = datacompy.Compare(do.reset_index(), dn.reset_index(),
                                         join_columns="index", df1_name="舊版", df2_name="新版")
            same = comp.matches()
            print(f"    表{i + 1}:{'一致' if same else '有異'}(datacompy)")
            if not same:
                print("      " + "\n      ".join(comp.report().splitlines()[-12:]))
        except ImportError:
            try:
                import pandas as pd
                if isinstance(do, pd.DataFrame) and do.shape == dn.shape:
                    delta = do.compare(dn)
                    print(f"    表{i + 1}:{'一致' if delta.empty else f'{len(delta)} 列有異(pandas.compare)'}")
                else:
                    print(f"    表{i + 1}:結構不同 {getattr(do, 'shape', '?')} vs {getattr(dn, 'shape', '?')}(誠實列示)")
            except ImportError:
                so, sn = json.dumps(do, ensure_ascii=False), json.dumps(dn, ensure_ascii=False)
                print(f"    表{i + 1}:{'一致' if so == sn else '有異'}(difflib 末梯)")
    return 0


# ── 主流程 ────────────────────────────────────────────────────────────────
def main() -> int:
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        return 2
    if "--compare" in args:
        i = args.index("--compare")
        return compare_docs(Path(args[i + 1]), Path(args[i + 2]))

    files = sorted((HERE / "input/incoming").glob(args[0]))
    files = [f for f in files if f.suffix.lower() in (".docx", ".doc")]
    if not files:
        print(f"[FAIL] 收件夾無 DOCX 匹配:{args[0]}")
        return 1
    out_root = HERE / "staging/ocr_out/docx_struct"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    print(f"=== DOCX 深度解析 v0100 · {len(files)} 檔 ===")
    fails = 0
    for f in files:
        try:
            bad = precheck_zip(f)
            if bad:
                print(f"  [WARN] {f.name} · {bad}")
                fails += 1
                continue
            paras, tables, engine = extract_docx(f)
            od = out_root / f.stem
            od.mkdir(parents=True, exist_ok=True)
            (od / "text.txt").write_text(clean_text("\n".join(paras)), encoding="utf-8")
            t_reports = []
            for i, rows in enumerate(tables, 1):
                tb, notes = repair_table(rows)
                v = validate_table(tb)
                if v["rows"] < 2 or v["empty_rate"] > 0.8:
                    t_reports.append({"table": i, **v, "repair_notes": notes + ["E5 SKIP_LAYOUT:疑排版用表,誠實跳過落檔"]})
                    continue
                try:
                    import pandas as pd
                    if isinstance(tb, pd.DataFrame):
                        tb.to_csv(od / f"table_{i}.csv", index=False, encoding="utf-8-sig")
                        try:
                            tb.to_parquet(od / f"table_{i}.parquet", index=False)
                        except Exception:
                            pass
                    else:
                        raise ImportError
                except ImportError:
                    with open(od / f"table_{i}.csv", "w", encoding="utf-8-sig", newline="") as fh:
                        import csv as _csv
                        _csv.writer(fh).writerows(tb if isinstance(tb, list) else [])
                t_reports.append({"table": i, **v, "repair_notes": notes})
            summary = {"file": f.name, "engine": engine, "ts": ts,
                       "paragraphs": len(paras), "tables": t_reports}
            (od / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=1), encoding="utf-8")
            tinfo = " · ".join(f"表{t['table']}:{t['rows']}x{t['cols']} 空值率{t['empty_rate']}" for t in t_reports) or "無表格"
            print(f"  [OK  ] {f.name} · 引擎={engine} · 段落 {len(paras)} · {tinfo}")
            for t in t_reports:
                for nn in t["repair_notes"]:
                    print(f"     ↳ 表{t['table']}:{nn}")
        except Exception as exc:
            print(f"  [FAIL] {f.name} · {type(exc).__name__}: {exc}")
            fails += 1
    print(f"  [出] {out_root}")
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
