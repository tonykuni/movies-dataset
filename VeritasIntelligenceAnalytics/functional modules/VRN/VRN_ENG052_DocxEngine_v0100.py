#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
vrn_docx_engine_v0100 — DOCX 深度解析引擎(文字+表格 擷取→修復→驗證對照)
==========================================================================
依操作員貼文規格(Gemini 對話)VIA 落地;只整理不發明,graceful 全梯次:
  擷取梯:docx2python(合併格最強)→ python-docx → 內建 XML 末梯
         (純標準庫 zipfile+ElementTree 解 word/document.xml;永不空手)
  修復三式(貼文原方):①空字串→NA ②ffill 向下填補(垂直合併格)
         ③空白/換行正規化(re)+ 圖片佔位符清除
  驗證:矩形性/表頭唯一/空值率(承進件矩陣四閘口徑)
  對照梯(--compare):datacompy(報表最詳)→ pandas.compare → difflib 文字 diff
輸出:staging/ocr_out/docx_struct/<stem>/ 下 text.txt + table_N.csv(+parquet)
     + summary.json;誠實逐檔矩陣。
用法:py vrn_docx_engine_v0100.py <pattern>              # 解析收件夾匹配檔
     py vrn_docx_engine_v0100.py --compare old.docx new.docx  # 新舊對照
"""
from __future__ import annotations
# ===== [VIA:ACCEL-BRIDGE:v0100] SuperAccel 加速器橋(全引擎導入令 2026-08-18;graceful 零行為變更) =====
try:
    import sys as _sa_sys
    from pathlib import Path as _sa_Path
    _sa_p = _sa_Path(__file__).resolve()
    while _sa_p.parent != _sa_p:
        if (_sa_p / "supportive modules" / "VIA_SuperAccel_Module.py").exists():
            _sa_sys.path.insert(0, str(_sa_p / "supportive modules"))
            break
        _sa_p = _sa_p.parent
    import VIA_SuperAccel_Module as VIA_ACCEL  # accel_map/fetch/pip_install/run_fast
except Exception:
    VIA_ACCEL = None  # graceful:加速器缺席零影響
# ===== [VIA:ACCEL-BRIDGE:END] =====

import json
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

HERE = Path(__file__).resolve().parent
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ── 擷取梯 ────────────────────────────────────────────────────────────────
def extract_builtin_xml(path: Path):
    """內建 XML 末梯:純標準庫解 word/document.xml(段落+表格;vMerge 續格留空)。"""
    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    body = root.find(f"{W_NS}body")
    paragraphs, tables = [], []
    for el in body:
        if el.tag == f"{W_NS}p":
            t = "".join(n.text or "" for n in el.iter(f"{W_NS}t"))
            if t.strip():
                paragraphs.append(t)
        elif el.tag == f"{W_NS}tbl":
            rows = []
            for tr in el.findall(f"{W_NS}tr"):
                row = []
                for tc in tr.findall(f"{W_NS}tc"):
                    vm = tc.find(f"{W_NS}tcPr/{W_NS}vMerge")
                    cont = vm is not None and vm.get(f"{W_NS}val", "continue") != "restart"
                    txt = "\n".join(
                        "".join(n.text or "" for n in p.iter(f"{W_NS}t"))
                        for p in tc.findall(f"{W_NS}p"))
                    row.append("" if cont else txt)
                rows.append(row)
            if rows:
                tables.append(rows)
    return paragraphs, tables, "builtin_xml"


def extract_docx(path: Path):
    """梯次:docx2python → python-docx → 內建 XML;回 (段落list, 表格list[rows], 引擎名)。"""
    try:
        from docx2python import docx2python
        with docx2python(path) as d:
            tables = []
            paragraphs = []
            for tbl in d.body:
                if len(tbl) > 1 or (len(tbl) == 1 and len(tbl[0]) > 1):
                    tables.append([["\n".join(c).strip() for c in row] for row in tbl])
                else:
                    for row in tbl:
                        for cell in row:
                            paragraphs.extend(p for p in cell if p.strip())
            return paragraphs, tables, "docx2python"
    except ImportError:
        pass
    try:
        import docx
        d = docx.Document(path)
        paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
        tables = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
        return paragraphs, tables, "python-docx"
    except ImportError:
        pass
    return (*extract_builtin_xml(path),)


# ── 修復三式(貼文原方)────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    text = re.sub(r"----media/.*?----", "", text)
    text = re.sub(r"[\r\t\f\v ]+", " ", text)
    text = re.sub(r"\n\s*\n", "\n", text)
    return text.strip()


def repair_table(rows):
    """回 (pandas.DataFrame|list, notes);pandas 缺時退純 python ffill。"""
    width = max(len(r) for r in rows)
    rect = [r + [""] * (width - len(r)) for r in rows]  # 欄位對齊
    notes = []
    if any(len(r) != width for r in rows):
        notes.append(f"欄數不齊已補齊至 {width}")
    try:
        import pandas as pd
        df = pd.DataFrame(rect)
        df = df.replace(r"^\s*$", pd.NA, regex=True)
        na_before = int(df.isna().sum().sum())
        df = df.ffill()
        filled = na_before - int(df.isna().sum().sum())
        if filled:
            notes.append(f"ffill 填補 {filled} 格(垂直合併格修復)")
        df = df.map(lambda x: re.sub(r"\s+", " ", str(x)).strip() if x is not None and str(x) != "<NA>" else x)
        if len(df) > 1:
            hdr = df.iloc[0].tolist()
            if len(set(map(str, hdr))) == len(hdr):
                df.columns = [str(h) for h in hdr]
                df = df[1:].reset_index(drop=True)
                notes.append("首列升表頭(唯一性過)")
            else:
                notes.append("首列表頭重複——保留數字欄名(誠實)")
        return df, notes
    except ImportError:
        out = []
        prev = [""] * width
        filled = 0
        for r in rect:
            nr = []
            for j, v in enumerate(r):
                v2 = re.sub(r"\s+", " ", v).strip()
                if not v2 and prev[j]:
                    v2 = prev[j]
                    filled += 1
                nr.append(v2)
            prev = nr
            out.append(nr)
        notes.append(f"pandas 缺——純 python ffill 填補 {filled} 格(誠實退路)")
        return out, notes


def validate_table(tb) -> dict:
    """驗證:矩形性/空值率(承進件矩陣口徑)。"""
    try:
        import pandas as pd
        if isinstance(tb, pd.DataFrame):
            total = tb.size or 1
            empty = int(tb.isna().sum().sum()) + int((tb.astype(str) == "").sum().sum())
            return {"rows": len(tb), "cols": len(tb.columns), "empty_rate": round(empty / total, 3)}
    except ImportError:
        pass
    total = sum(len(r) for r in tb) or 1
    empty = sum(1 for r in tb for v in r if not str(v).strip())
    return {"rows": len(tb), "cols": len(tb[0]) if tb else 0, "empty_rate": round(empty / total, 3)}


# ── 對照梯 ────────────────────────────────────────────────────────────────
def compare_docs(old: Path, new: Path) -> int:
    import difflib
    po, to, eo = extract_docx(old)
    pn, tn, en = extract_docx(new)
    print(f"=== DOCX 對照 · 舊 {old.name}({eo})vs 新 {new.name}({en})===")
    diff = [l for l in difflib.Differ().compare(
        clean_text("\n".join(po)).splitlines(), clean_text("\n".join(pn)).splitlines())
        if not l.startswith("  ")]
    print(f"  ── 文字層 diff({len(diff)} 行差異)──")
    for l in diff[:40]:
        print(f"    {l}")
    if len(diff) > 40:
        print(f"    …(餘 {len(diff) - 40} 行見報告)")
    n = min(len(to), len(tn))
    print(f"  ── 表格層({len(to)} vs {len(tn)} 表,對照前 {n} 表)──")
    for i in range(n):
        do, _ = repair_table(to[i])
        dn, _ = repair_table(tn[i])
        try:
            import datacompy
            comp = datacompy.Compare(do.reset_index(), dn.reset_index(),
                                     join_columns="index", df1_name="舊版", df2_name="新版")
            same = comp.matches()
            print(f"    表{i + 1}:{'一致' if same else '有異'}(datacompy)")
            if not same:
                print("      " + "\n      ".join(comp.report().splitlines()[-12:]))
        except ImportError:
            try:
                import pandas as pd
                if isinstance(do, pd.DataFrame) and do.shape == dn.shape:
                    delta = do.compare(dn)
                    print(f"    表{i + 1}:{'一致' if delta.empty else f'{len(delta)} 列有異(pandas.compare)'}")
                else:
                    print(f"    表{i + 1}:結構不同 {getattr(do, 'shape', '?')} vs {getattr(dn, 'shape', '?')}(誠實列示)")
            except ImportError:
                so, sn = json.dumps(do, ensure_ascii=False), json.dumps(dn, ensure_ascii=False)
                print(f"    表{i + 1}:{'一致' if so == sn else '有異'}(difflib 末梯)")
    return 0


# ── 主流程 ────────────────────────────────────────────────────────────────
def main() -> int:
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        return 2
    if "--compare" in args:
        i = args.index("--compare")
        return compare_docs(Path(args[i + 1]), Path(args[i + 2]))

    files = sorted((HERE / "input/incoming").glob(args[0]))
    files = [f for f in files if f.suffix.lower() in (".docx", ".doc")]
    if not files:
        print(f"[FAIL] 收件夾無 DOCX 匹配:{args[0]}")
        return 1
    out_root = HERE / "staging/ocr_out/docx_struct"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    print(f"=== DOCX 深度解析 v0100 · {len(files)} 檔 ===")
    fails = 0
    for f in files:
        try:
            paras, tables, engine = extract_docx(f)
            od = out_root / f.stem
            od.mkdir(parents=True, exist_ok=True)
            (od / "text.txt").write_text(clean_text("\n".join(paras)), encoding="utf-8")
            t_reports = []
            for i, rows in enumerate(tables, 1):
                tb, notes = repair_table(rows)
                v = validate_table(tb)
                try:
                    import pandas as pd
                    if isinstance(tb, pd.DataFrame):
                        tb.to_csv(od / f"table_{i}.csv", index=False, encoding="utf-8-sig")
                        try:
                            tb.to_parquet(od / f"table_{i}.parquet", index=False)
                        except Exception:
                            pass
                    else:
                        raise ImportError
                except ImportError:
                    with open(od / f"table_{i}.csv", "w", encoding="utf-8-sig", newline="") as fh:
                        import csv as _csv
                        _csv.writer(fh).writerows(tb if isinstance(tb, list) else [])
                t_reports.append({"table": i, **v, "repair_notes": notes})
            summary = {"file": f.name, "engine": engine, "ts": ts,
                       "paragraphs": len(paras), "tables": t_reports}
            (od / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=1), encoding="utf-8")
            tinfo = " · ".join(f"表{t['table']}:{t['rows']}x{t['cols']} 空值率{t['empty_rate']}" for t in t_reports) or "無表格"
            print(f"  [OK  ] {f.name} · 引擎={engine} · 段落 {len(paras)} · {tinfo}")
            for t in t_reports:
                for nn in t["repair_notes"]:
                    print(f"     ↳ 表{t['table']}:{nn}")
        except Exception as exc:
            print(f"  [FAIL] {f.name} · {type(exc).__name__}: {exc}")
            fails += 1
    print(f"  [出] {out_root}")
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
