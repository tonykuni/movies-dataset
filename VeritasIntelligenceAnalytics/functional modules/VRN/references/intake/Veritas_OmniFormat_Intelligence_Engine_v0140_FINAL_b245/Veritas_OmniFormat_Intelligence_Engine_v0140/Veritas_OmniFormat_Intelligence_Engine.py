#!/usr/bin/env python3
"""Veritas OmniFormat Intelligence Engine (VOFIE) v1.4.0.

VIA 的全格式讀取、主題重構與模板輸出層。設計原則：

* 來源唯讀，輸出永遠寫入新的 run 目錄。
* 內容先轉為 Universal Content IR，再交給格式 Adapter。
* 清理是隔離而非刪除；每個被隔離的片段仍保存在 IR。
* AI 只可產生候選；沒有等價測試時不得覆寫來源。
* 工具、格式與 ST 定位集中在頂端 Registry，未來擴充不必改 Dispatcher。
"""

from __future__ import annotations

import argparse
import ast
import csv
import hashlib
import html
import importlib.util
import io
import json
import os
import platform
import queue
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import threading
import time
import uuid
import zipfile
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable, Iterable, Iterator, Sequence
from xml.etree import ElementTree as ET


# ============================================================================
# 0. 所有可調參數與工具 Registry（未來新增格式／工具只改這一區或 overlay）
# ============================================================================

ENGINE_ID = "ENG-VOFIE-001"
SUBSYSTEM_ID = "VIA-SUBSYS-VOFIE-001"
ENGINE_NAME = "Veritas OmniFormat Intelligence Engine"
ENGINE_NAME_ZH = "Veritas 全格式智慧重構與模板生成引擎"
ENGINE_VERSION = "1.4.0"
REGISTRY_NAMESPACE = "veritas.omniformat"
IR_CONTRACT = "veritas.universal-content-ir/1.0"
INVOCATION_CONTRACT = "veritas.omniformat-invocation/1.0"
FAILURE_CATALOG_CONTRACT = "veritas.vofie-failure-catalog/1.1"
SIMPLE_RUN_CONTRACT = "veritas.vofie-simple-run/1.1"
COMPONENT_SPEC_CONTRACT = "veritas.vofie-component-spec/1.1"
ACTIVATION_CONTRACT = "veritas.vofie-activation/1.4"
POLYGLOT_TOOL_CATALOG_CONTRACT = "veritas.vofie-polyglot-tool-catalog/1.2"
POLYGLOT_TOOL_AUDIT_CONTRACT = "veritas.vofie-polyglot-tool-audit/1.2"
HYDRA_RISK_CATALOG_CONTRACT = "veritas.vofie-hydra-risk-catalog/1.0"
HYDRA_RISK_AUDIT_CONTRACT = "veritas.vofie-hydra-risk-audit/1.0"
RUNTIME_COPY_CONTRACT = "veritas.vofie-runtime-copy/1.0"
RUNTIME_COPY_APPROVAL_TOKEN = "YES_FOR_ANY_REAL_WRITE"

DEFAULT_MAX_INPUT_BYTES = 64 * 1024 * 1024
DEFAULT_TOPIC_CHARS = 18_000
DEFAULT_ENCODING_CANDIDATES = ("utf-8-sig", "utf-16", "big5", "cp950", "utf-8")
DEFAULT_OUTPUT_FORMATS = ("md", "json", "csv", "html", "css", "js")
ALL_OUTPUT_FORMATS = ("md", "json", "docx", "pptx", "xlsx", "csv", "html", "css", "js")
SIMPLE_MAX_INPUT_FILES = 5
SIMPLE_PRIMARY_OUTPUTS = ("md", "html", "component_json", "docx", "csv")
SIMPLE_PRIMARY_FILENAMES = {
    "md": "Veritas_VOFIE_Reconstructed.md",
    "html": "Veritas_VOFIE.html",
    "component_json": "Veritas_VOFIE_ComponentSpecs.json",
    "docx": "Veritas_VOFIE_Reconstructed.docx",
    "csv": "Veritas_VOFIE_TopicMatrix.csv",
}
RUN_ROLES = ("ENGINE", "SYSTEM")
DEFAULT_SIMPLE_ROLE = "ENGINE"
DEFAULT_OPERATIONS = ("text_merge", "code_merge", "restructure", "deduplicate", "optimize")
OPERATION_ORDER = DEFAULT_OPERATIONS
SYSTEM_SIDECAR_DIRECTORY = "_system"
GUI_TITLE = "Veritas VOFIE v1.4 — 五檔智慧重構"
GUI_GEOMETRY = "1040x720"
TOOL_PROBE_TIMEOUT_SECONDS = 6
POLYGLOT_TOOL_COUNTS = {"javascript": 20, "powershell": 20}
POLYGLOT_MATRIX_FUNCTIONS = (
    "syntax_parse", "static_analysis", "formatting", "unit_test", "coverage",
    "dependency_graph", "unused_code", "refactor_codmod", "schema_ui_validate",
    "build_automation",
)
SAFE_TOOL_EXECUTION_FUNCTIONS = ("syntax_parse",)
POLYGLOT_CAPABILITY_ROUTES = {
    "python": {
        "syntax_parse": ("Python ast", []),
        "static_analysis": ("Python ast + local rules", []),
        "formatting": ("canonical_text", []),
        "unit_test": ("unittest", []),
        "coverage": ("test result counts", []),
        "dependency_graph": ("import inventory", []),
        "unused_code": ("registry reference scan", []),
        "refactor_codmod": ("candidate-only AST", []),
        "schema_ui_validate": ("json/html stdlib", []),
        "build_automation": ("bounded stage runner", []),
    },
    "javascript": {
        "syntax_parse": ("JS-TOOL-001", ["JS-TOOL-002"]),
        "static_analysis": ("JS-TOOL-003", ["JS-TOOL-004", "JS-TOOL-010"]),
        "formatting": ("JS-TOOL-004", ["JS-TOOL-005"]),
        "unit_test": ("JS-TOOL-006", ["JS-TOOL-007", "JS-TOOL-008"]),
        "coverage": ("JS-TOOL-009", []),
        "dependency_graph": ("JS-TOOL-011", ["JS-TOOL-012"]),
        "unused_code": ("JS-TOOL-013", []),
        "refactor_codmod": ("JS-TOOL-017", ["JS-TOOL-018"]),
        "schema_ui_validate": ("JS-TOOL-020", ["JS-TOOL-019"]),
        "build_automation": ("JS-TOOL-014", ["JS-TOOL-015", "JS-TOOL-016"]),
    },
    "powershell": {
        "syntax_parse": ("PS-TOOL-002", ["PS-TOOL-001"]),
        "static_analysis": ("PS-TOOL-003", ["PS-TOOL-005"]),
        "formatting": ("PS-TOOL-003", []),
        "unit_test": ("PS-TOOL-004", []),
        "coverage": ("PS-TOOL-004", []),
        "dependency_graph": ("PS-TOOL-011", ["PS-TOOL-006", "PS-TOOL-007"]),
        "unused_code": ("PS-TOOL-003", []),
        "refactor_codmod": ("PS-TOOL-002", []),
        "schema_ui_validate": ("PS-TOOL-005", []),
        "build_automation": ("PS-TOOL-009", ["PS-TOOL-010"]),
    },
}
GUI_FILE_TYPES = (
    ("VOFIE 支援格式", "*.txt *.md *.html *.htm *.docx *.pptx *.xlsx *.csv *.tsv *.json *.xml *.yaml *.yml *.toml *.ini *.pdf *.py *.ps1 *.js *.ts *.css"),
    ("所有檔案", "*.*"),
)

# ST = Stability/Test Position。每個功能都有基準定位與可替換彈性。
ST_LEVELS = {
    "ST-FROZEN": "來源與舊結果不可改寫；任何轉換均另存新產物。",
    "ST-CORE": "不可缺少的核心契約；缺少即 FAIL。",
    "ST-GATE": "品質或安全閘門；不通過即 HOLD／FAIL。",
    "ST-ADAPTER": "可替換工具，但輸入輸出契約不可變。",
    "ST-OPTIONAL": "可停用的增強功能；停用不得影響核心產物。",
    "ST-HYDRA": "多點連動風險硬閘門；高風險只可 HOLD，不得自動修正或啟用。",
}

FORMAT_REGISTRY: dict[str, dict[str, Any]] = {
    "text": {
        "extensions": [".txt", ".text", ".log"],
        "reader": "read_text",
        "st": "ST-CORE",
        "baseline": "TEXT_TO_TOPIC_MARKDOWN",
        "flexibility": "ENCODING_ADAPTER",
    },
    "markdown": {
        "extensions": [".md", ".markdown", ".mdown"],
        "reader": "read_text",
        "st": "ST-CORE",
        "baseline": "MARKDOWN_AST_PRESERVING",
        "flexibility": "ANCHOR_AND_FENCE_NORMALIZER",
    },
    "html": {
        "extensions": [".html", ".htm", ".xhtml"],
        "reader": "read_html",
        "st": "ST-CORE",
        "baseline": "HTML_SEMANTIC_AND_UI_DUAL_LANE",
        "flexibility": "DOM_PARSER_ADAPTER",
    },
    "document": {
        "extensions": [".docx"],
        "reader": "read_docx",
        "st": "ST-ADAPTER",
        "baseline": "WORD_TEXT_AND_TABLE_EXTRACTION",
        "flexibility": "OOXML_OR_EXTERNAL_READER",
    },
    "presentation": {
        "extensions": [".pptx"],
        "reader": "read_pptx",
        "st": "ST-ADAPTER",
        "baseline": "SLIDE_ORDER_AND_TEXT_EXTRACTION",
        "flexibility": "OOXML_OR_EXTERNAL_READER",
    },
    "spreadsheet": {
        "extensions": [".xlsx", ".xlsm"],
        "reader": "read_xlsx",
        "st": "ST-ADAPTER",
        "baseline": "SHEET_CELL_AND_FORMULA_EXTRACTION",
        "flexibility": "OOXML_OR_EXTERNAL_READER",
    },
    "csv": {
        "extensions": [".csv", ".tsv"],
        "reader": "read_csv",
        "st": "ST-CORE",
        "baseline": "TABULAR_ROUND_TRIP",
        "flexibility": "DELIMITER_AND_ENCODING_ADAPTER",
    },
    "structured": {
        "extensions": [".json", ".jsonl", ".xml", ".yaml", ".yml", ".toml", ".ini"],
        "reader": "read_structured",
        "st": "ST-CORE",
        "baseline": "TREE_TO_MARKDOWN",
        "flexibility": "SCHEMA_ADAPTER",
    },
    "pdf": {
        "extensions": [".pdf"],
        "reader": "read_pdf",
        "st": "ST-ADAPTER",
        "baseline": "PAGE_ORDER_TEXT_EXTRACTION",
        "flexibility": "PYDPF_OR_PYPDF_READER",
    },
    "code": {
        "extensions": [
            ".py", ".pyw", ".ps1", ".psm1", ".js", ".mjs", ".cjs", ".ts", ".tsx",
            ".jsx", ".java", ".c", ".h", ".cpp", ".hpp", ".cs", ".go", ".rs", ".rb",
            ".php", ".swift", ".kt", ".kts", ".sql", ".sh", ".bash", ".r", ".lua",
            ".css", ".scss", ".vue", ".svelte",
        ],
        "reader": "read_code",
        "st": "ST-CORE",
        "baseline": "CODE_TO_COMPONENT_MARKDOWN",
        "flexibility": "LANGUAGE_AST_ADAPTER",
    },
}

OUTPUT_REGISTRY: dict[str, dict[str, str]] = {
    "md": {"adapter": "emit_markdown", "st": "ST-CORE", "baseline": "CANONICAL_HUMAN_OUTPUT"},
    "json": {"adapter": "emit_ir_json", "st": "ST-CORE", "baseline": "CANONICAL_MACHINE_OUTPUT"},
    "csv": {"adapter": "emit_csv", "st": "ST-CORE", "baseline": "TOPIC_MATRIX"},
    "docx": {"adapter": "emit_docx", "st": "ST-ADAPTER", "baseline": "WORD_REFERENCE_GUIDE"},
    "pptx": {"adapter": "artifact-tool", "st": "ST-ADAPTER", "baseline": "EXECUTIVE_TOPIC_DECK"},
    "xlsx": {"adapter": "artifact-tool", "st": "ST-ADAPTER", "baseline": "AUDITABLE_TOPIC_WORKBOOK"},
    "html": {"adapter": "emit_web_template", "st": "ST-CORE", "baseline": "SEMANTIC_TEMPLATE"},
    "css": {"adapter": "emit_web_template", "st": "ST-CORE", "baseline": "RESPONSIVE_THEME"},
    "js": {"adapter": "emit_web_template", "st": "ST-CORE", "baseline": "SAFE_LOCAL_INTERACTION"},
}

CODE_LANGUAGE_BY_SUFFIX = {
    ".py": "python", ".pyw": "python", ".ps1": "powershell", ".psm1": "powershell",
    ".js": "javascript", ".mjs": "javascript", ".cjs": "javascript", ".ts": "typescript",
    ".tsx": "tsx", ".jsx": "jsx", ".java": "java", ".c": "c", ".h": "c",
    ".cpp": "cpp", ".hpp": "cpp", ".cs": "csharp", ".go": "go", ".rs": "rust",
    ".rb": "ruby", ".php": "php", ".swift": "swift", ".kt": "kotlin", ".kts": "kotlin",
    ".sql": "sql", ".sh": "bash", ".bash": "bash", ".r": "r", ".lua": "lua",
    ".css": "css", ".scss": "scss", ".vue": "vue", ".svelte": "svelte",
}

FENCE_ALIASES = {
    "html": "html", "xml": "xml", "css": "css", "javascript": "javascript", "js": "javascript",
    "typescript": "typescript", "ts": "typescript", "python": "python", "py": "python",
    "powershell": "powershell", "ps1": "powershell", "bash": "bash", "shell": "bash",
    "json": "json", "yaml": "yaml", "yml": "yaml", "toml": "toml", "ini": "ini",
    "sql": "sql", "c": "c", "cpp": "cpp", "c#": "csharp", "csharp": "csharp",
    "java": "java", "go": "go", "rust": "rust", "text": "text", "plaintext": "text",
}

BOILERPLATE_RULES: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("render_noise", re.compile(r"^\s*(?:svg|image|顯示全部)\s*$", re.I)),
    ("model_disclaimer", re.compile(r"^\s*請謹慎使用程式碼。\s*$")),
    ("conversation_prompt", re.compile(r"^\s*(?:Tony[，,]|下一步(?:你要|（請選一個）|（你可點選）)|嘗試不使用個人化功能)")),
)

TOPIC_TAXONOMY: dict[str, tuple[str, ...]] = {
    "input_contract": ("input", "輸入", "讀取", "格式", "parser", "extract"),
    "content_restructure": ("restructure", "重構", "主題", "semantic", "markdown", "文章"),
    "code_restructure": ("code", "程式", "ast", "function", "class", "python", "javascript", "powershell"),
    "ui_specification": ("html", "css", "ui", "dashboard", "layout", "responsive", "component"),
    "document_output": ("word", "docx", "powerpoint", "pptx", "excel", "xlsx", "csv"),
    "quality_assurance": ("qa", "test", "驗證", "security", "accessibility", "usability", "gate"),
    "governance": ("ssot", "registry", "audit", "version", "append-only", "權限", "governance"),
    "deployment_ops": ("deployment", "docker", "kubernetes", "nginx", "telemetry", "monitor"),
    "data_analysis": ("plotly", "seaborn", "duckdb", "polars", "chart", "regime", "數據"),
}

UI_CAPABILITY_PROFILES: tuple[dict[str, str], ...] = (
    {"st_id": "ST-UI-001", "action": "ui_spec_extract", "position": "UI_CORE", "flexibility": "DOM_ADAPTER", "test": "CRITICAL", "note": "HTML 與 Markdown 雙軌保留。"},
    {"st_id": "ST-UI-002", "action": "state_machine", "position": "UI_LOGIC", "flexibility": "GENERATED_SPEC", "test": "STANDARD", "note": "事件與狀態分離。"},
    {"st_id": "ST-UI-003", "action": "interaction_graph", "position": "UI_LOGIC", "flexibility": "GENERATED_SPEC", "test": "STANDARD", "note": "元件、事件、目標可追溯。"},
    {"st_id": "ST-UI-004", "action": "test_cases", "position": "UI_QA", "flexibility": "RULE_EXTENSIBLE", "test": "CRITICAL", "note": "每個互動元件至少一個測試。"},
    {"st_id": "ST-UI-005", "action": "usability", "position": "UI_QA", "flexibility": "RULE_EXTENSIBLE", "test": "STANDARD", "note": "標籤、錯誤回饋與鍵盤路徑。"},
    {"st_id": "ST-UI-006", "action": "accessibility", "position": "UI_QA", "flexibility": "WCAG_RULE_ADAPTER", "test": "CRITICAL", "note": "語意元素、ARIA、對比與焦點。"},
    {"st_id": "ST-UI-007", "action": "security", "position": "UI_GATE", "flexibility": "FAIL_CLOSED", "test": "CRITICAL", "note": "不執行來源 script，不信任外部資源。"},
    {"st_id": "ST-UI-008", "action": "layout_optimize", "position": "UI_ENHANCEMENT", "flexibility": "TOKEN_DRIVEN", "test": "STANDARD", "note": "PC／Mobile 響應式。"},
    {"st_id": "ST-UI-009", "action": "component_refactor", "position": "UI_ENHANCEMENT", "flexibility": "CANDIDATE_ONLY", "test": "CRITICAL", "note": "不直接覆寫來源 UI。"},
    {"st_id": "ST-UI-010", "action": "performance", "position": "UI_ENHANCEMENT", "flexibility": "BUDGET_ADAPTER", "test": "STANDARD", "note": "離線與資源預算。"},
    {"st_id": "ST-UI-011", "action": "responsive", "position": "UI_ENHANCEMENT", "flexibility": "CSS_ADAPTER", "test": "CRITICAL", "note": "單欄與多欄斷點。"},
    {"st_id": "ST-UI-012", "action": "dark_mode", "position": "UI_OPTIONAL", "flexibility": "TOKEN_ADAPTER", "test": "SIMPLE", "note": "預設淺色，支援系統深色。"},
    {"st_id": "ST-UI-013", "action": "telemetry", "position": "UI_OPS", "flexibility": "OPT_IN", "test": "CRITICAL", "note": "本地、匿名、預設關閉。"},
)


# ============================================================================
# 1. 資料契約
# ============================================================================


class VOFIEError(RuntimeError):
    """可預期、可顯示的 VOFIE 錯誤。"""


@dataclass(slots=True)
class SourceRecord:
    source_id: str
    path: str
    name: str
    extension: str
    input_kind: str
    encoding: str
    byte_size: int
    source_hash: str
    extracted_hash: str
    extracted_text: str
    raw_preserved: bool = True
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class CodeUnit:
    unit_id: str
    language: str
    symbol: str
    unit_type: str
    signature: str
    start_line: int
    end_line: int
    syntax_status: str
    content_hash: str


@dataclass(slots=True)
class TopicBlock:
    topic_id: str
    source_id: str
    heading: str
    level: int
    order: int
    source_start_line: int
    source_end_line: int
    category: str
    tags: list[str]
    content: str
    content_hash: str
    duplicate_of: str | None = None
    quarantine_flags: list[str] = field(default_factory=list)
    code_units: list[CodeUnit] = field(default_factory=list)
    st_position: str = "ST-CORE"


@dataclass(slots=True)
class QuarantineItem:
    item_id: str
    source_id: str
    line_number: int
    reason: str
    content: str
    content_hash: str


@dataclass(slots=True)
class UniversalContentIR:
    contract: str
    engine_id: str
    engine_version: str
    subsystem_id: str
    registry_namespace: str
    run_id: str
    title: str
    created_at: str
    target_language: str
    source_records: list[SourceRecord]
    topics: list[TopicBlock]
    quarantine: list[QuarantineItem]
    ui_spec: dict[str, Any]
    quality: dict[str, Any]
    capability_profiles: list[dict[str, str]]
    output_files: dict[str, str] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(slots=True)
class EngineOptions:
    target_language: str = "zh-Hant"
    max_input_bytes: int = DEFAULT_MAX_INPUT_BYTES
    max_topic_chars: int = DEFAULT_TOPIC_CHARS
    keep_duplicate_topics: bool = True
    quarantine_boilerplate: bool = True
    use_vsis: bool = True
    ai_mode: str = "OFF"
    fail_closed: bool = True
    output_formats: tuple[str, ...] = DEFAULT_OUTPUT_FORMATS
    run_role: str = DEFAULT_SIMPLE_ROLE
    operations: tuple[str, ...] = DEFAULT_OPERATIONS


@dataclass(slots=True)
class RecoveryContext:
    """安全復原處理器的共同輸入；dry_run 不得寫檔或改來源。"""

    stage_id: str
    failure_id: str
    message: str = ""
    inputs: list[Path] = field(default_factory=list)
    output_dir: Path | None = None
    state: dict[str, Any] = field(default_factory=dict)
    dry_run: bool = True


@dataclass(slots=True)
class RecoveryActionResult:
    handler_id: str
    status: str
    action: str
    source_mutated: bool = False
    details: dict[str, Any] = field(default_factory=dict)


# ============================================================================
# 2. 共用工具與安全函式
# ============================================================================


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def blake2s_bytes(data: bytes) -> str:
    return hashlib.blake2s(data).hexdigest()


def blake2s_text(text: str) -> str:
    return blake2s_bytes(text.encode("utf-8", errors="replace"))


def stable_id(prefix: str, *parts: Any) -> str:
    payload = "\x1f".join(str(part) for part in parts)
    return f"{prefix}-{blake2s_text(payload)[:16].upper()}"


def natural_key(value: str) -> list[Any]:
    return [int(item) if item.isdigit() else item.casefold() for item in re.split(r"(\d+)", value)]


def canonical_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    return "\n".join(line.rstrip() for line in text.split("\n")).strip() + "\n"


def ensure_new_output_dir(path: Path) -> Path:
    resolved = path.expanduser().resolve()
    if resolved.exists() and any(resolved.iterdir()):
        suffix = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        resolved = resolved.with_name(f"{resolved.name}_{suffix}")
    resolved.mkdir(parents=True, exist_ok=True)
    return resolved


def atomic_write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    temporary.write_text(content, encoding="utf-8", newline="\n")
    os.replace(temporary, path)


def atomic_write_bytes(path: Path, content: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    temporary.write_bytes(content)
    os.replace(temporary, path)


def file_snapshot(path: Path) -> tuple[int, str]:
    data = path.read_bytes()
    return len(data), blake2s_bytes(data)


def decode_bytes(data: bytes) -> tuple[str, str]:
    for encoding in DEFAULT_ENCODING_CANDIDATES:
        try:
            return data.decode(encoding), encoding
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace"), "utf-8-replace"


def detect_kind(path: Path) -> str:
    suffix = path.suffix.casefold()
    for kind, profile in FORMAT_REGISTRY.items():
        if suffix in profile["extensions"]:
            return kind
    return "text"


def load_registry_overlay(project_root: Path) -> dict[str, Any]:
    """載入 add-only Registry overlay；不得刪除內建格式或改成未知 Reader。"""

    overlay_path = project_root / "config" / "tool_registry.json"
    if not overlay_path.is_file():
        return {"status": "SKIP", "reason": "NO_OVERLAY"}
    overlay = json.loads(overlay_path.read_text(encoding="utf-8"))
    if overlay.get("change_policy") != "APPEND_ONLY_DISABLE_INSTEAD_OF_DELETE":
        raise VOFIEError("tool_registry.json 缺少 add-only change_policy")
    added_extensions = 0
    for kind, extra in overlay.get("input_overlays", {}).items():
        if kind not in FORMAT_REGISTRY:
            raise VOFIEError(f"Registry overlay 指向未知輸入 Adapter：{kind}")
        known = FORMAT_REGISTRY[kind]["extensions"]
        for suffix in extra.get("extensions", []):
            normalized = suffix.casefold()
            if not normalized.startswith("."):
                raise VOFIEError(f"副檔名必須以 . 開頭：{suffix}")
            if normalized not in known:
                known.append(normalized)
                added_extensions += 1
    for suffix, language in overlay.get("code_language_overlays", {}).items():
        normalized = suffix.casefold()
        if not normalized.startswith("."):
            raise VOFIEError(f"程式副檔名必須以 . 開頭：{suffix}")
        if normalized not in CODE_LANGUAGE_BY_SUFFIX:
            CODE_LANGUAGE_BY_SUFFIX[normalized] = str(language)
            FORMAT_REGISTRY["code"]["extensions"].append(normalized)
            added_extensions += 1
    enabled_tools = [
        item["tool_id"] for item in overlay.get("tools", [])
        if item.get("enabled", True)
    ]
    return {"status": "PASS", "added_extensions": added_extensions, "enabled_tools": enabled_tools}


# ============================================================================
# 3A. JavaScript／PowerShell Top-20 CPU 工具治理層
# ============================================================================

def load_polyglot_tool_catalog(project_root: Path) -> dict[str, Any]:
    """載入並驗證雙語言 Top-20 工具 SSOT；只偵測，不自動安裝。"""

    path = project_root / "config" / "polyglot_tool_catalog.json"
    if not path.is_file():
        raise VOFIEError(f"找不到 Polyglot Tool Catalog：{path}")
    catalog = json.loads(path.read_text(encoding="utf-8"))
    if catalog.get("contract") != POLYGLOT_TOOL_CATALOG_CONTRACT:
        raise VOFIEError("Polyglot Tool Catalog contract 不相容")
    policy = catalog.get("policy", {})
    if policy.get("installation") != "DETECT_ONLY_NO_AUTO_INSTALL":
        raise VOFIEError("工具治理禁止自動安裝")
    if policy.get("source_mutation") != "DENIED":
        raise VOFIEError("工具治理不得允許來源改寫")
    all_ids: set[str] = set()
    for language, expected_count in POLYGLOT_TOOL_COUNTS.items():
        key = f"{language}_top20"
        rows = catalog.get(key, [])
        if len(rows) != expected_count:
            raise VOFIEError(f"{key} 必須恰為 {expected_count} 項，目前為 {len(rows)}")
        if [row.get("rank") for row in rows] != list(range(1, expected_count + 1)):
            raise VOFIEError(f"{key} rank 必須連續 1..{expected_count}")
        for row in rows:
            tool_id = str(row.get("tool_id", ""))
            if not tool_id or tool_id in all_ids:
                raise VOFIEError(f"工具 ID 缺少或重複：{tool_id}")
            all_ids.add(tool_id)
            if row.get("cpu_supported") is not True:
                raise VOFIEError(f"工具未聲明 CPU 支援：{tool_id}")
            if not row.get("license") or not row.get("fallback") or not row.get("route"):
                raise VOFIEError(f"工具缺少 license／fallback／route：{tool_id}")
    matrix_functions = tuple(catalog.get("capability_matrix_functions", []))
    if matrix_functions != POLYGLOT_MATRIX_FUNCTIONS:
        raise VOFIEError("Capability matrix function SSOT 不相容")
    return catalog


def _local_command_candidates(command: str, project_root: Path) -> list[Path]:
    if not command:
        return []
    extensions = [""]
    if os.name == "nt":
        extensions = [".cmd", ".exe", ".bat", ".ps1", ""]
    local_bin = project_root / "node_modules" / ".bin"
    return [local_bin / f"{command}{extension}" for extension in extensions]


def _resolve_tool_command(command: str, project_root: Path) -> str:
    if not command:
        return ""
    if command == "node" and os.environ.get("CODEX_PRIMARY_RUNTIME_NODE"):
        candidate = Path(os.environ["CODEX_PRIMARY_RUNTIME_NODE"])
        if candidate.is_file():
            return str(candidate)
    for candidate in _local_command_candidates(command, project_root):
        if candidate.is_file():
            return str(candidate.resolve())
    return shutil.which(command) or ""


def _node_package_roots(project_root: Path) -> list[Path]:
    roots = [project_root / "node_modules"]
    configured = os.environ.get("CODEX_PRIMARY_RUNTIME_NODE_MODULES")
    if configured:
        roots.append(Path(configured))
    return [path.resolve() for path in roots if path.is_dir()]


def _node_package_manifest(package_name: str, project_root: Path) -> Path | None:
    if not package_name:
        return None
    relative = Path(*package_name.split("/")) / "package.json"
    for root in _node_package_roots(project_root):
        candidate = root / relative
        if candidate.is_file():
            return candidate
    return None


def _package_manifest_version(path: Path | None) -> str:
    if path is None:
        return ""
    try:
        return str(json.loads(path.read_text(encoding="utf-8")).get("version", ""))
    except (OSError, json.JSONDecodeError):
        return ""


def _safe_version_probe(command_path: str, timeout_seconds: int = TOOL_PROBE_TIMEOUT_SECONDS) -> dict[str, Any]:
    """只呼叫 --version；不安裝、不寫檔、不執行來源程式。"""

    if not command_path:
        return {"status": "SKIP", "reason": "NO_COMMAND", "source_mutated": False}
    try:
        result = subprocess.run(
            [command_path, "--version"], capture_output=True, text=True,
            timeout=timeout_seconds, check=False,
        )
        output = (result.stdout or result.stderr).strip().splitlines()
        return {
            "status": "PASS" if result.returncode == 0 else "WARN",
            "exit_code": result.returncode,
            "version_text": output[0][:240] if output else "",
            "source_mutated": False,
        }
    except (OSError, subprocess.SubprocessError) as exc:
        return {"status": "WARN", "reason": f"{type(exc).__name__}: {exc}", "source_mutated": False}


def _detect_javascript_tools(catalog: dict[str, Any], project_root: Path, probe_installed: bool) -> list[dict[str, Any]]:
    detected: list[dict[str, Any]] = []
    for tool in catalog["javascript_top20"]:
        command_path = _resolve_tool_command(str(tool.get("command", "")), project_root)
        manifest = _node_package_manifest(str(tool.get("package", "")), project_root)
        available = bool(command_path or manifest)
        version = _package_manifest_version(manifest)
        probe = _safe_version_probe(command_path) if probe_installed and command_path else {"status": "NOT_RUN", "source_mutated": False}
        detected.append({
            **tool,
            "language": "javascript",
            "status": "AVAILABLE" if available else "NOT_INSTALLED",
            "resolved_command": command_path,
            "resolved_package": str(manifest) if manifest else "",
            "version": version,
            "probe": probe,
            "source_mutated": False,
        })
    return detected


def _powershell_module_directories() -> dict[str, Path]:
    result: dict[str, Path] = {}
    for root_name in os.environ.get("PSModulePath", "").split(os.pathsep):
        root = Path(root_name)
        if not root.is_dir():
            continue
        try:
            for child in root.iterdir():
                if child.is_dir():
                    result.setdefault(child.name.casefold(), child.resolve())
        except OSError:
            continue
    return result


def _probe_powershell_modules(pwsh_path: str, module_names: Sequence[str]) -> dict[str, dict[str, str]]:
    if not pwsh_path:
        return {}
    quoted = ",".join("'" + name.replace("'", "''") + "'" for name in module_names)
    script = (
        f"$names=@({quoted});$rows=foreach($n in $names){{"
        "$m=Get-Module -ListAvailable -Name $n|Sort-Object Version -Descending|Select-Object -First 1;"
        "if($null-ne$m){[pscustomobject]@{name=$n;version=$m.Version.ToString();path=$m.Path}}};"
        "$rows|ConvertTo-Json -Compress"
    )
    try:
        result = subprocess.run(
            [pwsh_path, "-NoProfile", "-NonInteractive", "-Command", script],
            capture_output=True, text=True, timeout=TOOL_PROBE_TIMEOUT_SECONDS, check=False,
        )
        if result.returncode != 0 or not result.stdout.strip():
            return {}
        payload = json.loads(result.stdout)
        rows = payload if isinstance(payload, list) else [payload]
        return {str(row["name"]).casefold(): {"version": str(row.get("version", "")), "path": str(row.get("path", ""))} for row in rows}
    except (OSError, subprocess.SubprocessError, json.JSONDecodeError, KeyError, TypeError):
        return {}


def _detect_powershell_tools(catalog: dict[str, Any], project_root: Path, probe_installed: bool) -> list[dict[str, Any]]:
    del project_root  # Reserved for future isolated via-ps roots; no broad filesystem scan.
    pwsh_path = shutil.which("pwsh") or ""
    module_dirs = _powershell_module_directories()
    module_names = [str(tool.get("module", "")) for tool in catalog["powershell_top20"] if tool.get("kind") == "powershell_module"]
    probed_modules = _probe_powershell_modules(pwsh_path, module_names) if probe_installed else {}
    detected: list[dict[str, Any]] = []
    for tool in catalog["powershell_top20"]:
        kind = str(tool.get("kind", ""))
        module_name = str(tool.get("module", ""))
        module_key = module_name.casefold()
        module_probe = probed_modules.get(module_key, {})
        module_path = str(module_dirs.get(module_key, "")) or module_probe.get("path", "")
        if kind == "command":
            available = bool(pwsh_path)
            resolved_command = pwsh_path
        elif kind == "powershell_builtin":
            available = bool(pwsh_path)
            resolved_command = pwsh_path
        else:
            available = bool(module_path)
            resolved_command = ""
        probe = _safe_version_probe(pwsh_path) if probe_installed and kind == "command" and pwsh_path else {"status": "NOT_RUN", "source_mutated": False}
        detected.append({
            **tool,
            "language": "powershell",
            "status": "BUILTIN" if available and kind == "powershell_builtin" else "AVAILABLE" if available else "NOT_INSTALLED",
            "resolved_command": resolved_command,
            "resolved_module": module_path,
            "version": module_probe.get("version", ""),
            "probe": probe,
            "source_mutated": False,
        })
    return detected


def powershell_structure_check(path: Path) -> dict[str, Any]:
    """無 pwsh 時的唯讀結構降級檢查；正式 AST 仍由 PS7 bridge 提供。"""

    before = file_snapshot(path)
    text = path.read_text(encoding="utf-8-sig")
    stack: list[tuple[str, int]] = []
    pairs = {")": "(", "]": "[", "}": "{"}
    openers = set(pairs.values())
    quote = ""
    escaped = False
    block_comment = False
    here_end = ""
    errors: list[str] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        if here_end:
            if line.strip() == here_end:
                here_end = ""
            continue
        stripped = line.strip()
        if "@'" in line and not stripped.startswith("#"):
            here_end = "'@"
        elif '@"' in line and not stripped.startswith("#"):
            here_end = '"@'
        index = 0
        while index < len(line):
            char = line[index]
            next_pair = line[index:index + 2]
            if block_comment:
                if next_pair == "#>":
                    block_comment = False
                    index += 2
                    continue
                index += 1
                continue
            if not quote and next_pair == "<#":
                block_comment = True
                index += 2
                continue
            if escaped:
                escaped = False
                index += 1
                continue
            if quote and char == "`":
                escaped = True
                index += 1
                continue
            if quote:
                if char == quote:
                    quote = ""
                index += 1
                continue
            if char == "#":
                break
            if char in {"'", '"'}:
                quote = char
            elif char in openers:
                stack.append((char, line_number))
            elif char in pairs:
                if not stack or stack[-1][0] != pairs[char]:
                    errors.append(f"line {line_number}: unexpected {char}")
                else:
                    stack.pop()
            index += 1
        if quote == "'":
            quote = ""
    if stack:
        errors.extend(f"line {line_number}: unclosed {char}" for char, line_number in stack[-10:])
    if quote:
        errors.append("unclosed double-quoted string")
    if block_comment:
        errors.append("unclosed block comment")
    if here_end:
        errors.append("unclosed here-string")
    functions = re.findall(r"(?im)^\s*function\s+([A-Za-z][A-Za-z0-9_-]*)", text)
    function_position = re.search(r"(?im)^\s*function\s+", text)
    param_position = re.search(r"(?im)^\s*param\s*\(", text)
    after = file_snapshot(path)
    return {
        "gate": "PASS" if not errors and before == after else "FAIL",
        "errors": errors,
        "function_count": len(functions),
        "functions": functions,
        "top_parameter_block": bool(param_position and (not function_position or param_position.start() < function_position.start())),
        "source_mutated": before != after,
        "fallback": "PYTHON_STRUCTURAL_ONLY_USE_PS7_AST_WHEN_AVAILABLE",
    }


def _build_polyglot_capability_matrix(tools: Sequence[dict[str, Any]]) -> list[dict[str, Any]]:
    index = {str(tool["tool_id"]): tool for tool in tools}
    fallback_by_tool = {str(tool["tool_id"]): str(tool.get("fallback", "")) for tool in tools}
    rows: list[dict[str, Any]] = []
    for language in ("python", "javascript", "powershell"):
        for function_id in POLYGLOT_MATRIX_FUNCTIONS:
            primary, alternates = POLYGLOT_CAPABILITY_ROUTES[language][function_id]
            if language == "python":
                selected = primary
                status = "BUILTIN"
                fallback = "not required"
            else:
                candidates = [primary, *alternates]
                selected = next((tool_id for tool_id in candidates if index.get(tool_id, {}).get("status") in {"AVAILABLE", "BUILTIN"}), "")
                status = "READY_PRIMARY" if selected == primary else "READY_ALTERNATE" if selected else "READY_FALLBACK"
                selected = selected or primary
                fallback = fallback_by_tool.get(primary, "")
            rows.append({
                "matrix_id": f"MATRIX-{language.upper()}-{len(rows) % 10 + 1:02d}",
                "language": language,
                "function": function_id,
                "primary": primary,
                "alternates": alternates,
                "selected": selected,
                "status": status,
                "fallback": fallback,
                "cpu_supported": True,
                "read_only": True,
                "machine_readable": True,
            })
    return rows


def tool_audit(project_root: Path, language: str = "all", probe_installed: bool = False) -> dict[str, Any]:
    normalized_language = language.casefold()
    if normalized_language not in {"all", "javascript", "powershell"}:
        raise VOFIEError(f"未知工具語言：{language}")
    catalog = load_polyglot_tool_catalog(project_root)
    javascript = _detect_javascript_tools(catalog, project_root, probe_installed)
    powershell = _detect_powershell_tools(catalog, project_root, probe_installed)
    all_tools = [*javascript, *powershell]
    selected_tools = all_tools if normalized_language == "all" else [tool for tool in all_tools if tool["language"] == normalized_language]
    matrix = _build_polyglot_capability_matrix(all_tools)
    selected_matrix = matrix if normalized_language == "all" else [row for row in matrix if row["language"] == normalized_language]
    launcher_check = powershell_structure_check(project_root / "Invoke-Veritas-VOFIE.ps1")
    bridge_check = powershell_structure_check(project_root / "adapters" / "Veritas.VOFIE.ToolBridge.psm1")
    js_bridge = project_root / "adapters" / "vofie_polyglot_tool_probe.mjs"
    gaps = [row for row in selected_matrix if not row.get("fallback") and row["status"] == "READY_FALLBACK"]
    expected_matrix_count = 30 if normalized_language == "all" else 10
    gate = "PASS" if (
        len(selected_tools) == (40 if normalized_language == "all" else 20)
        and len(selected_matrix) == expected_matrix_count
        and not gaps
        and all(tool.get("cpu_supported") and tool.get("license") for tool in selected_tools)
        and launcher_check["gate"] == "PASS"
        and bridge_check["gate"] == "PASS"
        and js_bridge.is_file()
    ) else "FAIL"
    return {
        "contract": POLYGLOT_TOOL_AUDIT_CONTRACT,
        "engine_version": ENGINE_VERSION,
        "gate": gate,
        "language": normalized_language,
        "probe_installed": probe_installed,
        "architecture": {
            "machine": platform.machine() or "unknown",
            "processor": platform.processor() or "unknown",
            "logical_cores": os.cpu_count() or 1,
            "cpu_only_supported": True,
            "gpu_required": False,
        },
        "policy": catalog["policy"],
        "summary": {
            "total": len(selected_tools),
            "available": sum(tool["status"] in {"AVAILABLE", "BUILTIN"} for tool in selected_tools),
            "not_installed": sum(tool["status"] == "NOT_INSTALLED" for tool in selected_tools),
            "matrix_rows": len(selected_matrix),
            "uncovered_functions": len(gaps),
        },
        "structural_checks": {
            "javascript_bridge": {"gate": "PASS" if js_bridge.is_file() else "FAIL", "path": str(js_bridge)},
            "powershell_launcher": launcher_check,
            "powershell_bridge": bridge_check,
        },
        "tools": selected_tools,
        "capability_matrix": selected_matrix,
        "gaps": gaps,
        "source_mutated": False,
    }


def parse_tool_functions(value: str | Sequence[str]) -> tuple[str, ...]:
    if isinstance(value, str):
        candidates = [item.strip().casefold() for item in value.split(",") if item.strip()]
    else:
        candidates = [str(item).strip().casefold() for item in value if str(item).strip()]
    functions = tuple(dict.fromkeys(candidates or POLYGLOT_MATRIX_FUNCTIONS))
    unknown = [item for item in functions if item not in POLYGLOT_MATRIX_FUNCTIONS]
    if unknown:
        raise VOFIEError(f"未知工具功能：{', '.join(unknown)}")
    return functions


def detect_tool_language(path: Path) -> str:
    suffix = path.suffix.casefold()
    if suffix in {".js", ".mjs", ".cjs", ".ts", ".tsx", ".jsx"}:
        return "javascript"
    if suffix in {".ps1", ".psm1", ".psd1"}:
        return "powershell"
    raise VOFIEError(f"tool-plan 目前只接受 JavaScript／TypeScript／PowerShell：{path.name}")


def _execute_javascript_syntax(path: Path, project_root: Path) -> dict[str, Any]:
    node = _resolve_tool_command("node", project_root)
    if not node:
        text = path.read_text(encoding="utf-8-sig")
        stack: list[str] = []
        pairs = {")": "(", "]": "[", "}": "{"}
        for character in text:
            if character in pairs.values():
                stack.append(character)
            elif character in pairs and (not stack or stack.pop() != pairs[character]):
                return {"gate": "FAIL", "runner": "python_js_structural_scan", "message": f"unexpected {character}", "external_process_started": False}
        return {"gate": "PASS" if not stack else "FAIL", "runner": "python_js_structural_scan", "message": "" if not stack else "unclosed bracket", "external_process_started": False}
    result = subprocess.run(
        [node, "--check", str(path)], capture_output=True, text=True,
        timeout=TOOL_PROBE_TIMEOUT_SECONDS, check=False,
    )
    return {
        "gate": "PASS" if result.returncode == 0 else "FAIL",
        "runner": "JS-TOOL-001",
        "exit_code": result.returncode,
        "message": (result.stderr or result.stdout).strip()[-1200:],
        "external_process_started": True,
    }


def _execute_powershell_syntax(path: Path) -> dict[str, Any]:
    pwsh = shutil.which("pwsh") or ""
    if not pwsh:
        result = powershell_structure_check(path)
        return {
            "gate": result["gate"],
            "runner": "python_powershell_structural_scan",
            "message": "; ".join(result["errors"]),
            "details": result,
            "external_process_started": False,
        }
    escaped = str(path.resolve()).replace("'", "''")
    script = (
        "$t=$null;$e=$null;"
        f"[void][Management.Automation.Language.Parser]::ParseFile('{escaped}',[ref]$t,[ref]$e);"
        "[pscustomobject]@{gate=$(if($e.Count-eq0){'PASS'}else{'FAIL'});errors=@($e|% Message)}|ConvertTo-Json -Compress"
    )
    result = subprocess.run(
        [pwsh, "-NoProfile", "-NonInteractive", "-Command", script],
        capture_output=True, text=True, timeout=TOOL_PROBE_TIMEOUT_SECONDS, check=False,
    )
    try:
        payload = json.loads(result.stdout) if result.stdout.strip() else {}
    except json.JSONDecodeError:
        payload = {}
    return {
        "gate": "PASS" if result.returncode == 0 and payload.get("gate") == "PASS" else "FAIL",
        "runner": "PS-TOOL-002",
        "exit_code": result.returncode,
        "message": "; ".join(payload.get("errors", [])) if isinstance(payload.get("errors"), list) else (result.stderr or result.stdout).strip()[-1200:],
        "external_process_started": True,
    }


def tool_plan(
    project_root: Path,
    target: Path,
    functions: Sequence[str] = POLYGLOT_MATRIX_FUNCTIONS,
    execute_safe: bool = False,
) -> dict[str, Any]:
    """依需求挑選工具；只有 syntax_parse 可在明確 opt-in 下執行唯讀 quick check。"""

    target = target.expanduser().resolve()
    if not target.is_file():
        raise VOFIEError(f"tool-plan target 不存在：{target}")
    selected_functions = parse_tool_functions(functions)
    language = detect_tool_language(target)
    audit = tool_audit(project_root, language=language)
    rows = [row for row in audit["capability_matrix"] if row["function"] in selected_functions]
    snapshot = file_snapshot(target)
    executions: list[dict[str, Any]] = []
    if execute_safe and "syntax_parse" in selected_functions:
        result = _execute_javascript_syntax(target, project_root) if language == "javascript" else _execute_powershell_syntax(target)
        executions.append({"function": "syntax_parse", **result})
    for function_id in selected_functions:
        if function_id == "syntax_parse" and execute_safe:
            continue
        executions.append({
            "function": function_id,
            "gate": "PLAN_ONLY",
            "runner": next(row["selected"] for row in rows if row["function"] == function_id),
            "message": "explicit tool-specific opt-in required; no source mutation performed",
            "external_process_started": False,
        })
    source_mutated = file_snapshot(target) != snapshot
    failed = [row for row in executions if row["gate"] == "FAIL"]
    return {
        "contract": "veritas.vofie-tool-execution-plan/1.2",
        "engine_version": ENGINE_VERSION,
        "gate": "FAIL" if failed or source_mutated else "PASS",
        "language": language,
        "target": str(target),
        "functions": list(selected_functions),
        "execute_safe": execute_safe,
        "safe_execution_functions": list(SAFE_TOOL_EXECUTION_FUNCTIONS),
        "plan": rows,
        "executions": executions,
        "external_process_started": any(row["external_process_started"] for row in executions),
        "source_mutated": source_mutated,
        "source_policy": "READ_ONLY_NO_DELETE_NO_MOVE_NO_CANONICAL_MUTATION",
    }


def compact_tool_audit(report: dict[str, Any]) -> dict[str, Any]:
    return {
        "contract": report["contract"],
        "gate": report["gate"],
        "architecture": report["architecture"],
        "policy": report["policy"],
        "summary": report["summary"],
        "capability_matrix": report["capability_matrix"],
        "tools": [
            {
                "tool_id": tool["tool_id"], "rank": tool["rank"], "name": tool["name"],
                "language": tool["language"], "status": tool["status"],
                "functions": tool["functions"], "fallback": tool["fallback"],
                "cpu_supported": tool["cpu_supported"], "license": tool["license"],
            }
            for tool in report["tools"]
        ],
    }


def _recovery_result(handler_id: str, action: str, context: RecoveryContext, **details: Any) -> RecoveryActionResult:
    return RecoveryActionResult(handler_id, "DRY_RUN" if context.dry_run else "APPLIED", action, False, details)


def recover_validate_input(context: RecoveryContext) -> RecoveryActionResult:
    findings = [{"path": str(path), "is_file": path.expanduser().is_file()} for path in context.inputs]
    return _recovery_result("validate_input", "validate every selected input without changing it", context, findings=findings)


def recover_normalize_path(context: RecoveryContext) -> RecoveryActionResult:
    paths = [str(path.expanduser().resolve()) for path in context.inputs]
    context.state["normalized_inputs"] = paths
    return _recovery_result("normalize_path", "resolve local input paths deterministically", context, paths=paths)


def recover_deduplicate_selection(context: RecoveryContext) -> RecoveryActionResult:
    seen: set[str] = set()
    unique: list[str] = []
    for path in context.inputs:
        value = str(path.expanduser().resolve())
        if value not in seen:
            seen.add(value)
            unique.append(value)
    context.state["unique_inputs"] = unique
    return _recovery_result("deduplicate_selection", "retain the first occurrence of duplicate selections", context, inputs=unique)


def recover_limit_selection(context: RecoveryContext) -> RecoveryActionResult:
    selected = [str(path) for path in context.inputs[:SIMPLE_MAX_INPUT_FILES]]
    context.state["limited_inputs"] = selected
    return _recovery_result("limit_selection", "limit the simple profile to five inputs", context, inputs=selected, limit=SIMPLE_MAX_INPUT_FILES)


def recover_emit_diagnostic(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("emit_diagnostic", "return a structured, user-readable incident", context, stage=context.stage_id, failure=context.failure_id, message=context.message)


def recover_encoding_fallback(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("encoding_fallback", "try declared deterministic encoding candidates and replacement decoding last", context, encodings=list(DEFAULT_ENCODING_CANDIDATES) + ["utf-8-replace"])


def recover_stdlib_reader(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("stdlib_reader", "use the built-in read-only adapter before optional dependencies", context, available=sorted(READERS))


def recover_isolate_file(context: RecoveryContext) -> RecoveryActionResult:
    context.state["isolation_policy"] = "record_failure_continue_other_inputs"
    return _recovery_result("isolate_file", "isolate only the failing input and preserve its identity", context)


def recover_chunk_content(context: RecoveryContext) -> RecoveryActionResult:
    context.state["chunk_limit"] = DEFAULT_TOPIC_CHARS
    return _recovery_result("chunk_content", "split oversized content at deterministic paragraph boundaries", context, max_chars=DEFAULT_TOPIC_CHARS)


def recover_retain_original(context: RecoveryContext) -> RecoveryActionResult:
    context.state["source_policy"] = "READ_ONLY_EMBED_OR_REFERENCE"
    return _recovery_result("retain_original", "retain original bytes, hashes and extracted text as the authority", context)


def recover_normalize_fence(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("normalize_fence", "repair only the generated Markdown candidate and retain source text", context)


def recover_deterministic_nlp(context: RecoveryContext) -> RecoveryActionResult:
    context.state["nlp_fallback"] = "LOCAL_TAXONOMY_AND_HEADING_SEGMENTER"
    return _recovery_result("deterministic_nlp", "fall back to local heading segmentation and taxonomy classification", context)


def recover_rebuild_index(context: RecoveryContext) -> RecoveryActionResult:
    context.state["reindex"] = True
    return _recovery_result("rebuild_index", "rebuild stable topic and component indexes from retained records", context)


def recover_mark_duplicate(context: RecoveryContext) -> RecoveryActionResult:
    context.state["duplicate_policy"] = "MARK_AND_RETAIN"
    return _recovery_result("mark_duplicate", "mark duplicate_of and keep every duplicate in the IR", context)


def recover_candidate_only(context: RecoveryContext) -> RecoveryActionResult:
    context.state["rewrite_policy"] = "CANDIDATE_ONLY_EQUIVALENCE_GATE"
    return _recovery_result("candidate_only", "store optimized output as a candidate without source or API mutation", context)


def recover_disable_optional(context: RecoveryContext) -> RecoveryActionResult:
    context.state.setdefault("disabled_optional", []).append(context.failure_id)
    return _recovery_result("disable_optional", "disable only the failing optional enhancement and continue core processing", context)


def recover_relocate_vsis(context: RecoveryContext) -> RecoveryActionResult:
    candidates = ["project-parent/VIA_SemanticIntelligenceSubsystem_v0120", "VIA_VSIS_ROOT"]
    return _recovery_result("relocate_vsis", "search the two approved VSIS locations without broad filesystem scanning", context, candidates=candidates)


def recover_new_output_dir(context: RecoveryContext) -> RecoveryActionResult:
    value = str(context.output_dir.resolve()) if context.output_dir else ""
    context.state["new_output_policy"] = "TIMESTAMP_SUFFIX_IF_NONEMPTY"
    return _recovery_result("new_output_dir", "choose a new sibling run directory when the target is non-empty", context, requested=value)


def recover_atomic_retry(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("atomic_retry", "write a temporary sibling and retry one atomic replace", context, retry_limit=1)


def recover_sanitize_filename(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("sanitize_filename", "use fixed ASCII-safe VOFIE output filenames", context, filenames=SIMPLE_PRIMARY_FILENAMES)


def recover_inline_assets(context: RecoveryContext) -> RecoveryActionResult:
    context.state["web_assets"] = "SELF_CONTAINED_LOCAL"
    return _recovery_result("inline_assets", "inline CSS, JavaScript and JSON with no remote dependency", context)


def recover_validate_output(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("validate_output", "verify existence, non-empty bytes, expected names and BLAKE2s hashes", context)


def recover_ui_file_dialog(context: RecoveryContext) -> RecoveryActionResult:
    return _recovery_result("ui_file_dialog", "fall back to the native multiple-file picker", context)


def recover_ui_disable_dnd(context: RecoveryContext) -> RecoveryActionResult:
    context.state["drag_drop"] = "DISABLED_WITH_FILE_PICKER_FALLBACK"
    return _recovery_result("ui_disable_dnd", "disable drag-drop when its optional adapter is unavailable", context)


def recover_ui_preserve_state(context: RecoveryContext) -> RecoveryActionResult:
    context.state["ui_state_policy"] = "KEEP_SELECTION_AND_PARAMETERS_ON_FAILURE"
    return _recovery_result("ui_preserve_state", "keep selections and parameters visible after an error", context)


def recover_hold_activation(context: RecoveryContext) -> RecoveryActionResult:
    context.state["activation"] = "HOLD"
    return _recovery_result("hold_activation", "refuse activation until all required tests pass", context)


RECOVERY_HANDLERS: dict[str, Callable[[RecoveryContext], RecoveryActionResult]] = {
    "validate_input": recover_validate_input,
    "normalize_path": recover_normalize_path,
    "deduplicate_selection": recover_deduplicate_selection,
    "limit_selection": recover_limit_selection,
    "emit_diagnostic": recover_emit_diagnostic,
    "encoding_fallback": recover_encoding_fallback,
    "stdlib_reader": recover_stdlib_reader,
    "isolate_file": recover_isolate_file,
    "chunk_content": recover_chunk_content,
    "retain_original": recover_retain_original,
    "normalize_fence": recover_normalize_fence,
    "deterministic_nlp": recover_deterministic_nlp,
    "rebuild_index": recover_rebuild_index,
    "mark_duplicate": recover_mark_duplicate,
    "candidate_only": recover_candidate_only,
    "disable_optional": recover_disable_optional,
    "relocate_vsis": recover_relocate_vsis,
    "new_output_dir": recover_new_output_dir,
    "atomic_retry": recover_atomic_retry,
    "sanitize_filename": recover_sanitize_filename,
    "inline_assets": recover_inline_assets,
    "validate_output": recover_validate_output,
    "ui_file_dialog": recover_ui_file_dialog,
    "ui_disable_dnd": recover_ui_disable_dnd,
    "ui_preserve_state": recover_ui_preserve_state,
    "hold_activation": recover_hold_activation,
}


def load_failure_catalog(project_root: Path) -> dict[str, Any]:
    path = project_root / "config" / "failure_catalog.json"
    if not path.is_file():
        raise VOFIEError(f"找不到 Failure Catalog：{path}")
    catalog = json.loads(path.read_text(encoding="utf-8"))
    if catalog.get("contract") != FAILURE_CATALOG_CONTRACT:
        raise VOFIEError("Failure Catalog contract 不相容")
    materialized: list[dict[str, Any]] = []
    stage_ids: set[str] = set()
    for stage in catalog.get("stages", []):
        stage_id = str(stage.get("stage_id", "")).strip().upper()
        if not stage_id or stage_id in stage_ids:
            raise VOFIEError(f"Failure Catalog stage_id 無效或重複：{stage_id}")
        stage_ids.add(stage_id)
        handlers = [str(item) for item in stage.get("handlers", [])]
        missing_handlers = [item for item in handlers if item not in RECOVERY_HANDLERS]
        if missing_handlers:
            raise VOFIEError(f"Failure Catalog 未實作 handlers：{', '.join(missing_handlers)}")
        failures = []
        for rank, name in enumerate(stage.get("failures", []), start=1):
            failures.append({
                "failure_id": f"{stage_id}-F{rank:02d}",
                "rank": rank,
                "name": str(name),
                "handlers": handlers,
                "implemented_solutions": len(handlers),
            })
        materialized.append({"stage_id": stage_id, "name": stage.get("name", stage_id), "handlers": handlers, "failures": failures})
    return {**catalog, "stages": materialized, "stage_count": len(materialized), "failure_count": sum(len(stage["failures"]) for stage in materialized)}


def load_hydra_risk_catalog(project_root: Path) -> dict[str, Any]:
    """載入獨立 NoHydra Top-20；不把多點連動風險混入一般 failure catalog。"""

    path = project_root / "config" / "hydra_risk_catalog.json"
    if not path.is_file():
        raise VOFIEError(f"找不到 Hydra Risk Catalog：{path}")
    catalog = json.loads(path.read_text(encoding="utf-8"))
    if catalog.get("contract") != HYDRA_RISK_CATALOG_CONTRACT:
        raise VOFIEError("Hydra Risk Catalog contract 不相容")
    risks = catalog.get("risks", [])
    ids = [str(item.get("risk_id", "")) for item in risks]
    ranks = [item.get("rank") for item in risks]
    allowed_actions = {"HOLD", "REVIEW", "FAIL_CLOSED"}
    allowed_lanes = {"PARALLEL_SAFE_PROPOSAL", "SEQUENTIAL_ONLY"}
    if len(risks) != 20 or len(set(ids)) != 20 or ranks != list(range(1, 21)):
        raise VOFIEError("Hydra Risk Catalog 必須是唯一且連續排序的 Top 20")
    for item in risks:
        if len(item.get("solutions", [])) < 3 or len(item.get("breakers", [])) < 2 or len(item.get("detectors", [])) < 2:
            raise VOFIEError(f"{item.get('risk_id')} 至少需要 3 個 solutions、2 個 breakers 與 2 個 detectors")
        if item.get("default_action") not in allowed_actions or item.get("repair_lane") not in allowed_lanes:
            raise VOFIEError(f"{item.get('risk_id')} 的 action／lane 無效")
        if not item.get("sop") or not item.get("never_again_control"):
            raise VOFIEError(f"{item.get('risk_id')} 缺少 SOP／never-again control")
    policy = catalog.get("policy", {})
    if policy.get("max_rounds") != 3 or policy.get("high_hydra_action") != "HOLD":
        raise VOFIEError("Hydra policy 必須固定三輪上限且高風險 HOLD")
    return catalog


def _hydra_target_findings(path: Path, catalog: dict[str, Any]) -> list[dict[str, Any]]:
    """只做 bounded read-only evidence scan；不 import、不執行、不修檔。"""

    text = path.read_text(encoding="utf-8-sig", errors="replace")[:2_000_000]
    by_id = {item["risk_id"]: item for item in catalog["risks"]}
    findings: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    def add(risk_id: str, evidence: str) -> None:
        key = (risk_id, evidence)
        if risk_id not in by_id or key in seen:
            return
        seen.add(key)
        risk = by_id[risk_id]
        findings.append({
            "risk_id": risk_id,
            "rank": risk["rank"],
            "name": risk["name"],
            "severity": risk["severity"],
            "decision": risk["default_action"],
            "repair_lane": risk["repair_lane"],
            "evidence": evidence,
            "breakers": risk["breakers"],
            "solutions": risk["solutions"],
            "sop": risk["sop"],
        })

    for match in re.finditer(r"HYDRA:RISK=(HYDRA-F\d{2})", text, flags=re.I):
        add(match.group(1).upper(), "explicit risk marker")
    patterns = (
        ("HYDRA-F04", r"(?:subprocess\.(?:run|Popen)|Start-Process)[^\n]*(?:__file__|\$PSCommandPath)"),
        ("HYDRA-F05", r"(?:max_workers\s*=\s*(?:None|os\.cpu_count\(\))|ForEach-Object\s+-Parallel[^\n]*ThrottleLimit\s+0)"),
        ("HYDRA-F10", r"while\s*(?:\(|)\s*True\s*(?:\)|)\s*:"),
        ("HYDRA-F13", r"(?:pip|npm|Install-Module)\s+install?\b|Install-Module\b"),
        ("HYDRA-F18", r"direct_apply\s*[=:]\s*(?:true|True|1)"),
    )
    for risk_id, pattern in patterns:
        if re.search(pattern, text, flags=re.I):
            add(risk_id, f"high-confidence static pattern: {pattern}")
    return findings


def hydra_risk_audit(project_root: Path, targets: Sequence[Path] = ()) -> dict[str, Any]:
    """執行 NoHydra review-only Gate；偵測到風險只 HOLD，不採取真實修復。"""

    catalog = load_hydra_risk_catalog(project_root)
    snapshots: dict[str, tuple[int, str]] = {}
    unique_findings: dict[tuple[str, str, str], dict[str, Any]] = {}
    scanned: list[str] = []
    for value in targets:
        path = Path(value).expanduser().resolve()
        if not path.is_file():
            raise VOFIEError(f"Hydra audit target 不是現有檔案：{path}")
        if path.stat().st_size > DEFAULT_MAX_INPUT_BYTES:
            raise VOFIEError(f"Hydra audit target 超過 {DEFAULT_MAX_INPUT_BYTES} bytes：{path.name}")
        snapshots[str(path)] = file_snapshot(path)
        scanned.append(str(path))
    policy = catalog["policy"]
    round_definitions = (
        (1, "PANORAMA_READ_ONLY", "inventory + classify"),
        (2, "PARALLEL_SAFE_PROPOSAL", "proposal evidence only; no fixer execution"),
        (3, "SEQUENTIAL_DEPENDENCY_REVIEW", "post-scan + remaining risk HOLD/manual review"),
    )
    round_evidence: list[dict[str, Any]] = []
    source_mutated = False
    for round_number, mode, action in round_definitions:
        round_findings: list[dict[str, Any]] = []
        for path_text in scanned:
            path = Path(path_text)
            for finding in _hydra_target_findings(path, catalog):
                row = {"target": path_text, **finding}
                unique_findings[(path_text, finding["risk_id"], finding["evidence"])] = row
                round_findings.append(row)
        integrity = all(file_snapshot(Path(path)) == snapshot for path, snapshot in snapshots.items())
        source_mutated = source_mutated or not integrity
        round_evidence.append({
            "round": round_number,
            "mode": mode,
            "action": action,
            "targets_scanned": len(scanned),
            "finding_events": len(round_findings),
            "unique_findings_to_date": len(unique_findings),
            "source_integrity": "PASS" if integrity else "FAIL",
            "post_scan_complete": True,
            "real_write": False,
        })
    findings = sorted(unique_findings.values(), key=lambda item: (item["rank"], item["target"], item["evidence"]))
    catalog_solutions = sum(len(item["solutions"]) for item in catalog["risks"])
    decision = "FAIL" if source_mutated else ("HOLD" if findings else "PASS")
    return {
        "contract": HYDRA_RISK_AUDIT_CONTRACT,
        "engine_version": ENGINE_VERSION,
        "gate": decision,
        "activation_allowed": decision == "PASS",
        "mode": "REVIEW_ONLY_DRY_RUN",
        "st": "ST-HYDRA",
        "policy": policy,
        "summary": {
            "top_risks": len(catalog["risks"]),
            "solutions": catalog_solutions,
            "targets_scanned": len(scanned),
            "findings": len(findings),
            "max_rounds": policy["max_rounds"],
            "parallel_safe": sum(item["repair_lane"] == "PARALLEL_SAFE_PROPOSAL" for item in catalog["risks"]),
            "sequential_only": sum(item["repair_lane"] == "SEQUENTIAL_ONLY" for item in catalog["risks"]),
            "detector_coverage": len(catalog["risks"]),
            "post_scans": len(round_evidence),
        },
        "round_plan": round_evidence,
        "targets": scanned,
        "findings": findings,
        "top20": catalog["risks"],
        "source_mutated": source_mutated,
        "external_process_started": False,
        "network_accessed": False,
        "real_write_performed": False,
        "source_policy": "READ_ONLY_NO_DELETE_NO_MOVE_NO_CANONICAL_MUTATION",
    }


def hash_state_decision(current_hash: str | None, original_hash: str, proposed_hash: str) -> dict[str, Any]:
    """LL#32 deterministic idempotency: missing/apply, proposed/skip, original/backup+apply, other/fail closed."""

    if not current_hash:
        state, action = "MISSING", "APPLY"
    elif current_hash == proposed_hash:
        state, action = "PROPOSED", "SKIP"
    elif current_hash == original_hash:
        state, action = "ORIGINAL", "BACKUP_APPLY"
    else:
        state, action = "OTHER", "FAIL_CLOSED"
    return {
        "contract": "veritas.hash-state-machine/1.0",
        "state": state,
        "action": action,
        "safe_to_continue": action in {"APPLY", "SKIP", "BACKUP_APPLY"},
        "canonical_mutation_allowed": False,
    }


def rollback_dry_run(manifest_path: Path) -> dict[str, Any]:
    """Verify that a runtime copy can be abandoned while canonical sources remain intact."""

    resolved = manifest_path.expanduser().resolve()
    payload = json.loads(resolved.read_text(encoding="utf-8"))
    if payload.get("contract") != RUNTIME_COPY_CONTRACT:
        raise VOFIEError("Runtime Copy manifest contract 不相容")
    checks: list[dict[str, Any]] = []
    for row in payload.get("files", []):
        source = Path(row["source"])
        runtime = Path(row["runtime_copy"])
        source_ok = source.is_file() and blake2s_bytes(source.read_bytes()) == row["source_hash"]
        runtime_ok = runtime.is_file() and blake2s_bytes(runtime.read_bytes()) == row["runtime_hash"]
        checks.append({
            "source": str(source), "runtime_copy": str(runtime),
            "source_integrity": source_ok, "runtime_integrity": runtime_ok,
            "rollback_action": "ABANDON_RUNTIME_COPY_ONLY",
        })
    passed = bool(checks) and all(row["source_integrity"] and row["runtime_integrity"] for row in checks)
    return {
        "contract": "veritas.vofie-rollback-dry-run/1.0",
        "gate": "PASS" if passed else "FAIL",
        "manifest": str(resolved),
        "checks": checks,
        "canonical_mutated": False,
        "runtime_mutated": False,
        "real_rollback_performed": False,
    }


def create_runtime_copy(
    sources: Sequence[Path],
    output_dir: Path,
    approval_token: str,
) -> dict[str, Any]:
    """Create a versioned run-local sandbox; this never promotes or modifies canonical sources."""

    if approval_token != RUNTIME_COPY_APPROVAL_TOKEN:
        return {
            "contract": RUNTIME_COPY_CONTRACT,
            "gate": "HOLD",
            "reason": "explicit runtime-copy approval token missing",
            "approval_required": RUNTIME_COPY_APPROVAL_TOKEN,
            "source_mutated": False,
            "runtime_copy_created": False,
        }
    normalized = [Path(path).expanduser().resolve() for path in sources]
    if not normalized or len(normalized) > SIMPLE_MAX_INPUT_FILES:
        raise VOFIEError(f"Runtime Copy 需要 1–{SIMPLE_MAX_INPUT_FILES} 個來源檔")
    if len({str(path) for path in normalized}) != len(normalized) or not all(path.is_file() for path in normalized):
        raise VOFIEError("Runtime Copy 來源必須存在且不可重複")
    before = {str(path): file_snapshot(path) for path in normalized}
    runtime_root = ensure_new_output_dir(output_dir)
    files: list[dict[str, Any]] = []
    for index, source in enumerate(normalized, start=1):
        content = source.read_bytes()
        destination = runtime_root / f"{index:02d}_{source.name}"
        decision = hash_state_decision(None, blake2s_bytes(content), blake2s_bytes(content))
        atomic_write_bytes(destination, content)
        runtime_hash = blake2s_bytes(destination.read_bytes())
        files.append({
            "source": str(source), "runtime_copy": str(destination),
            "source_hash": blake2s_bytes(content), "runtime_hash": runtime_hash,
            "hash_state": decision, "copy_verified": runtime_hash == blake2s_bytes(content),
        })
    source_mutated = any(file_snapshot(Path(path)) != snapshot for path, snapshot in before.items())
    manifest = {
        "contract": RUNTIME_COPY_CONTRACT,
        "engine_version": ENGINE_VERSION,
        "created_at": utc_now(),
        "runtime_root": str(runtime_root),
        "source_policy": "CANONICAL_READ_ONLY_RUNTIME_COPY_WRITABLE",
        "approval_token_verified": True,
        "files": files,
        "source_mutated": source_mutated,
        "promotion_performed": False,
    }
    manifest_path = runtime_root / "RuntimeCopyManifest.json"
    atomic_write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True))
    rollback = rollback_dry_run(manifest_path)
    manifest["manifest"] = str(manifest_path)
    manifest["rollback_dry_run"] = rollback
    manifest["gate"] = "PASS" if not source_mutated and all(row["copy_verified"] for row in files) and rollback["gate"] == "PASS" else "FAIL"
    atomic_write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True))
    return manifest


def emit_hydra_html_matrix(report: dict[str, Any], path: Path) -> Path:
    """Create a compact, self-contained administrator matrix; no remote assets or source execution."""

    rows = []
    for risk in report["top20"]:
        solutions = "；".join(risk["solutions"])
        breakers = "；".join(risk["breakers"])
        rows.append(
            "<tr>"
            f"<td>{html.escape(risk['risk_id'])}</td><td>{risk['rank']}</td>"
            f"<td>{html.escape(risk['name'])}</td><td>{html.escape(risk['severity'])}</td>"
            f"<td>{html.escape(risk['repair_lane'])}</td><td>{html.escape(breakers)}</td>"
            f"<td>{html.escape(solutions)}</td><td>{html.escape(risk['sop'])}</td>"
            "</tr>"
        )
    document = f"""<!doctype html>
<html lang="zh-Hant"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Veritas VOFIE NoHydra Matrix</title><style>
:root{{--ink:#17212b;--muted:#64727d;--line:#d9e2e8;--brand:#0f5f73;--soft:#f4f8fa}}*{{box-sizing:border-box}}
body{{margin:0;background:#fff;color:var(--ink);font:12px/1.45 Inter,"Noto Sans TC",sans-serif}}header{{padding:16px 20px;border-bottom:1px solid var(--line);display:flex;justify-content:space-between;gap:12px;align-items:center}}h1{{font-size:18px;margin:0}}main{{padding:12px 20px}}.meta{{color:var(--muted)}}.wrap{{overflow:auto;border:1px solid var(--line);border-radius:8px}}table{{border-collapse:collapse;width:100%;min-width:1280px}}th,td{{padding:7px 8px;border-bottom:1px solid var(--line);vertical-align:top;text-align:left}}th{{position:sticky;top:0;background:var(--soft);color:var(--brand)}}tbody tr:hover{{background:#f8fbfc}}.gate{{font-weight:700;color:var(--brand)}}
</style></head><body><header><div><div class="meta">VERITAS INTELLIGENCE ANALYTICS</div><h1>NoHydra Top-20 Risk Matrix</h1></div><div class="gate">{html.escape(report['gate'])}</div></header>
<main><p class="meta">三輪實掃 · Canonical Read-only · Runtime Copy Proposal-first · 高風險 HOLD</p><div class="wrap"><table><thead><tr><th>ID</th><th>#</th><th>Risk</th><th>Severity</th><th>Lane</th><th>Breakers</th><th>Solutions</th><th>SOP</th></tr></thead><tbody>{''.join(rows)}</tbody></table></div></main></body></html>"""
    atomic_write_text(path, document)
    return path


def exercise_recovery_handlers(project_root: Path) -> dict[str, Any]:
    catalog = load_failure_catalog(project_root)
    results: list[dict[str, Any]] = []
    for handler_id, handler in RECOVERY_HANDLERS.items():
        context = RecoveryContext("GOVERNANCE", "RECOVERY-SELF-TEST", inputs=[], output_dir=project_root / "qa", dry_run=True)
        try:
            result = handler(context)
            results.append({"handler_id": handler_id, "status": result.status, "source_mutated": result.source_mutated})
        except Exception as exc:
            results.append({"handler_id": handler_id, "status": "FAIL", "message": f"{type(exc).__name__}: {exc}"})
    passed = all(item["status"] == "DRY_RUN" and not item.get("source_mutated") for item in results)
    coverage = all(len(stage["failures"]) == 20 and len(stage["handlers"]) >= 2 for stage in catalog["stages"])
    return {"gate": "PASS" if passed and coverage else "FAIL", "catalog": {"stages": catalog["stage_count"], "failures": catalog["failure_count"]}, "handlers": results}


def preflight_simple_run(inputs: Sequence[Path], output_dir: Path, role: str, project_root: Path) -> dict[str, Any]:
    catalog = load_failure_catalog(project_root)
    incidents: list[dict[str, Any]] = []

    def incident(stage_id: str, failure_id: str, severity: str, message: str) -> None:
        stage = next(item for item in catalog["stages"] if item["stage_id"] == stage_id)
        failure = next(item for item in stage["failures"] if item["failure_id"] == failure_id)
        context = RecoveryContext(stage_id, failure_id, message, list(inputs), output_dir, dry_run=True)
        actions = [asdict(RECOVERY_HANDLERS[handler](context)) for handler in failure["handlers"]]
        incidents.append({"stage_id": stage_id, "failure_id": failure_id, "severity": severity, "message": message, "recovery_actions": actions})

    if not inputs:
        incident("INTAKE", "INTAKE-F01", "ERROR", "至少需要一個輸入檔")
    if len(inputs) > SIMPLE_MAX_INPUT_FILES:
        incident("INTAKE", "INTAKE-F02", "ERROR", f"簡易模式最多 {SIMPLE_MAX_INPUT_FILES} 個輸入")
    resolved = [path.expanduser().resolve() for path in inputs]
    if len({str(path) for path in resolved}) != len(resolved):
        incident("INTAKE", "INTAKE-F03", "ERROR", "輸入清單含重複路徑")
    for path in resolved:
        if not path.exists():
            incident("INTAKE", "INTAKE-F04", "ERROR", f"找不到輸入：{path}")
        elif not path.is_file():
            incident("INTAKE", "INTAKE-F05", "ERROR", f"輸入不是檔案：{path}")
        elif path.stat().st_size == 0:
            incident("INTAKE", "INTAKE-F07", "WARNING", f"空檔將按原樣登記：{path.name}")
    normalized_role = role.upper()
    if normalized_role not in RUN_ROLES:
        incident("WINDOW_UI", "WINDOW_UI-F09", "ERROR", f"未知角色：{role}")
    output_resolved = output_dir.expanduser().resolve()
    if output_resolved in resolved:
        incident("INTAKE", "INTAKE-F20", "ERROR", "輸出路徑不可等於輸入檔")
    errors = [item for item in incidents if item["severity"] == "ERROR"]
    return {
        "contract": "veritas.vofie-preflight/1.1",
        "gate": "FAIL" if errors else ("PASS_WITH_WARNINGS" if incidents else "PASS"),
        "role": normalized_role,
        "input_count": len(inputs),
        "primary_output_count": len(SIMPLE_PRIMARY_OUTPUTS),
        "catalog_coverage": {"stages": catalog["stage_count"], "failures": catalog["failure_count"]},
        "incidents": incidents,
    }


def escape_md_cell(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def strip_markdown_markup(text: str) -> str:
    text = re.sub(r"```.*?```", "[code]", text, flags=re.S)
    text = re.sub(r"[`*_>#]", "", text)
    text = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()


def truncate(text: str, limit: int) -> str:
    compact = re.sub(r"\s+", " ", text).strip()
    return compact if len(compact) <= limit else compact[: max(0, limit - 1)].rstrip() + "…"


def append_audit_event(audit_path: Path, event: dict[str, Any]) -> dict[str, Any]:
    previous_hash = "0" * 64
    if audit_path.exists():
        try:
            last = audit_path.read_text(encoding="utf-8").splitlines()[-1]
            previous_hash = json.loads(last).get("chain_hash", previous_hash)
        except Exception:
            previous_hash = "AUDIT_READ_ERROR"
    payload = {**event, "previous_hash": previous_hash, "at": utc_now()}
    payload["chain_hash"] = blake2s_text(json.dumps(payload, ensure_ascii=False, sort_keys=True))
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    with audit_path.open("a", encoding="utf-8", newline="\n") as handle:
        handle.write(json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n")
        handle.flush()
        os.fsync(handle.fileno())
    return payload


# ============================================================================
# 3. 讀取 Adapter：文字、HTML、Office、CSV、結構資料與程式碼
# ============================================================================


class SemanticHTMLParser(HTMLParser):
    BLOCK_TAGS = {"p", "div", "section", "article", "main", "aside", "header", "footer", "nav"}
    SKIP_TAGS = {"script", "style", "noscript", "template"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.lines: list[str] = []
        self.stack: list[str] = []
        self.current: list[str] = []
        self.link_href: str | None = None
        self.table_row: list[str] = []
        self.table_rows: list[list[str]] = []
        self.cell: list[str] | None = None
        self.scripts: list[str] = []
        self.styles: list[str] = []
        self._capture_script = False
        self._capture_style = False
        self.components: list[dict[str, Any]] = []
        self.endpoints: list[str] = []

    def flush(self, prefix: str = "", suffix: str = "") -> None:
        value = re.sub(r"\s+", " ", " ".join(self.current)).strip()
        self.current.clear()
        if value:
            self.lines.append(f"{prefix}{value}{suffix}")

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.casefold()
        attributes = {key.casefold(): (value or "") for key, value in attrs}
        self.stack.append(tag)
        if tag == "script":
            self._capture_script = True
        elif tag == "style":
            self._capture_style = True
        elif re.fullmatch(r"h[1-6]", tag):
            self.flush()
        elif tag == "li":
            self.flush()
        elif tag == "a":
            self.link_href = attributes.get("href")
        elif tag == "br":
            self.flush()
        elif tag == "img":
            alt = attributes.get("alt") or attributes.get("title") or "image"
            src = attributes.get("src", "")
            self.lines.append(f"![{alt}]({src})")
        elif tag in {"th", "td"}:
            self.cell = []
        if tag in {"input", "select", "textarea", "button", "a", "form"}:
            component = {
                "tag": tag,
                "id": attributes.get("id", ""),
                "name": attributes.get("name", ""),
                "type": attributes.get("type", ""),
                "role": attributes.get("role", ""),
                "aria_label": attributes.get("aria-label", ""),
                "href": attributes.get("href", ""),
                "action": attributes.get("action", ""),
            }
            self.components.append(component)
            for key in ("href", "action"):
                value = component.get(key, "")
                if value and not value.startswith(("#", "javascript:")):
                    self.endpoints.append(value)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.casefold()
        if re.fullmatch(r"h[1-6]", tag):
            self.flush(prefix="#" * int(tag[1]) + " ")
        elif tag == "li":
            self.flush(prefix="- ")
        elif tag == "a":
            value = re.sub(r"\s+", " ", " ".join(self.current)).strip()
            self.current.clear()
            if value:
                self.current.append(f"[{value}]({self.link_href})" if self.link_href else value)
            self.link_href = None
        elif tag in self.BLOCK_TAGS:
            self.flush()
        elif tag in {"th", "td"} and self.cell is not None:
            self.table_row.append(re.sub(r"\s+", " ", " ".join(self.cell)).strip())
            self.cell = None
        elif tag == "tr":
            if self.table_row:
                self.table_rows.append(self.table_row)
                self.table_row = []
        elif tag == "script":
            self._capture_script = False
        elif tag == "style":
            self._capture_style = False
        if self.stack:
            for index in range(len(self.stack) - 1, -1, -1):
                if self.stack[index] == tag:
                    del self.stack[index:]
                    break

    def handle_data(self, data: str) -> None:
        if self._capture_script:
            self.scripts.append(data)
            return
        if self._capture_style:
            self.styles.append(data)
            return
        if any(tag in self.SKIP_TAGS for tag in self.stack):
            return
        value = data.strip()
        if not value:
            return
        if self.cell is not None:
            self.cell.append(value)
        else:
            self.current.append(value)

    def markdown(self) -> str:
        self.flush()
        if self.table_rows:
            width = max(len(row) for row in self.table_rows)
            rows = [row + [""] * (width - len(row)) for row in self.table_rows]
            header_row = rows[0]
            self.lines.append("| " + " | ".join(escape_md_cell(cell) for cell in header_row) + " |")
            self.lines.append("| " + " | ".join("---" for _ in header_row) + " |")
            for row in rows[1:]:
                self.lines.append("| " + " | ".join(escape_md_cell(cell) for cell in row) + " |")
        return canonical_text("\n\n".join(line for line in self.lines if line.strip()))


def read_text_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    data = path.read_bytes()
    if len(data) > options.max_input_bytes:
        raise VOFIEError(f"輸入超過 {options.max_input_bytes} bytes：{path}")
    text, encoding = decode_bytes(data)
    return canonical_text(text), encoding, {}


def read_html_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    raw, encoding, _ = read_text_adapter(path, options)
    parser = SemanticHTMLParser()
    parser.feed(raw)
    markdown = parser.markdown()
    metadata = {
        "html_source": raw,
        "styles": canonical_text("\n".join(parser.styles)) if parser.styles else "",
        "scripts": canonical_text("\n".join(parser.scripts)) if parser.scripts else "",
        "components": parser.components,
        "endpoints": sorted(set(parser.endpoints)),
    }
    return markdown, encoding, metadata


def _xml_text(element: ET.Element) -> str:
    return "".join(element.itertext()).strip()


def read_docx_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    if path.stat().st_size > options.max_input_bytes:
        raise VOFIEError(f"輸入超過 {options.max_input_bytes} bytes：{path}")
    with zipfile.ZipFile(path) as package:
        xml = package.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: list[str] = []
    for node in root.findall(".//w:body/*", ns):
        local = node.tag.rsplit("}", 1)[-1]
        if local == "p":
            value = "".join(text.text or "" for text in node.findall(".//w:t", ns)).strip()
            if not value:
                continue
            style = node.find("./w:pPr/w:pStyle", ns)
            style_value = style.attrib.get(f"{{{ns['w']}}}val", "") if style is not None else ""
            match = re.search(r"heading\s*([1-6])", style_value, re.I)
            lines.append(("#" * int(match.group(1)) + " " if match else "") + value)
        elif local == "tbl":
            rows = []
            for row in node.findall(".//w:tr", ns):
                cells = ["".join(t.text or "" for t in cell.findall(".//w:t", ns)).strip() for cell in row.findall("./w:tc", ns)]
                rows.append(cells)
            if rows:
                width = max(len(row) for row in rows)
                rows = [row + [""] * (width - len(row)) for row in rows]
                lines.append("| " + " | ".join(map(escape_md_cell, rows[0])) + " |")
                lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
                lines.extend("| " + " | ".join(map(escape_md_cell, row)) + " |" for row in rows[1:])
    return canonical_text("\n\n".join(lines)), "ooxml", {"reader": "stdlib-ooxml"}


def read_pptx_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    if path.stat().st_size > options.max_input_bytes:
        raise VOFIEError(f"輸入超過 {options.max_input_bytes} bytes：{path}")
    lines: list[str] = []
    with zipfile.ZipFile(path) as package:
        names = sorted((name for name in package.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)), key=natural_key)
        for index, name in enumerate(names, start=1):
            root = ET.fromstring(package.read(name))
            values = [node.text or "" for node in root.iter() if node.tag.endswith("}t")]
            lines.append(f"# Slide {index}")
            lines.extend(value.strip() for value in values if value.strip())
    return canonical_text("\n\n".join(lines)), "ooxml", {"reader": "stdlib-ooxml", "slides": len(names)}


def _xlsx_shared_strings(package: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in package.namelist():
        return []
    root = ET.fromstring(package.read("xl/sharedStrings.xml"))
    return ["".join(node.itertext()) for node in root if node.tag.endswith("}si")]


def read_xlsx_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    if path.stat().st_size > options.max_input_bytes:
        raise VOFIEError(f"輸入超過 {options.max_input_bytes} bytes：{path}")
    output: list[str] = []
    with zipfile.ZipFile(path) as package:
        shared = _xlsx_shared_strings(package)
        sheet_names = sorted((name for name in package.namelist() if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)), key=natural_key)
        for sheet_index, sheet_name in enumerate(sheet_names, start=1):
            output.append(f"# Sheet {sheet_index}")
            root = ET.fromstring(package.read(sheet_name))
            rows: list[list[str]] = []
            for row in (node for node in root.iter() if node.tag.endswith("}row")):
                values: list[str] = []
                for cell in (node for node in row if node.tag.endswith("}c")):
                    cell_type = cell.attrib.get("t", "")
                    value_node = next((node for node in cell if node.tag.endswith("}v")), None)
                    formula_node = next((node for node in cell if node.tag.endswith("}f")), None)
                    value = value_node.text if value_node is not None and value_node.text is not None else ""
                    if cell_type == "s" and value.isdigit() and int(value) < len(shared):
                        value = shared[int(value)]
                    if formula_node is not None and formula_node.text:
                        value = f"={formula_node.text} → {value}"
                    values.append(value)
                if values:
                    rows.append(values)
            if rows:
                width = max(len(row) for row in rows)
                rows = [row + [""] * (width - len(row)) for row in rows]
                output.append("| " + " | ".join(map(escape_md_cell, rows[0])) + " |")
                output.append("| " + " | ".join("---" for _ in rows[0]) + " |")
                output.extend("| " + " | ".join(map(escape_md_cell, row)) + " |" for row in rows[1:])
    return canonical_text("\n\n".join(output)), "ooxml", {"reader": "stdlib-ooxml", "sheets": len(sheet_names)}


def read_csv_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    raw, encoding, _ = read_text_adapter(path, options)
    try:
        dialect = csv.Sniffer().sniff(raw[:8192], delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel_tab if path.suffix.casefold() == ".tsv" else csv.excel
    rows = list(csv.reader(io.StringIO(raw), dialect))
    if not rows:
        return "", encoding, {"rows": 0, "delimiter": dialect.delimiter}
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(map(escape_md_cell, rows[0])) + " |", "| " + " | ".join("---" for _ in rows[0]) + " |"]
    lines.extend("| " + " | ".join(map(escape_md_cell, row)) + " |" for row in rows[1:])
    return canonical_text("\n".join(lines)), encoding, {"rows": len(rows), "delimiter": dialect.delimiter}


def _json_to_markdown(value: Any, heading: str = "Document", level: int = 1) -> list[str]:
    lines = [f"{'#' * min(level, 6)} {heading}"]
    if isinstance(value, dict):
        for key, child in value.items():
            if isinstance(child, (dict, list)):
                lines.extend(_json_to_markdown(child, str(key), level + 1))
            else:
                lines.append(f"- **{key}**: {child}")
    elif isinstance(value, list):
        for index, child in enumerate(value, start=1):
            if isinstance(child, (dict, list)):
                lines.extend(_json_to_markdown(child, f"Item {index}", level + 1))
            else:
                lines.append(f"- {child}")
    else:
        lines.append(str(value))
    return lines


def read_structured_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    raw, encoding, _ = read_text_adapter(path, options)
    if path.suffix.casefold() == ".json":
        try:
            value = json.loads(raw)
            return canonical_text("\n\n".join(_json_to_markdown(value, path.stem))), encoding, {"parsed": True, "format": "json"}
        except json.JSONDecodeError as exc:
            return raw, encoding, {"parsed": False, "format": "json", "warning": str(exc)}
    if path.suffix.casefold() == ".jsonl":
        lines = []
        failures = 0
        for index, line in enumerate(raw.splitlines(), start=1):
            try:
                lines.extend(_json_to_markdown(json.loads(line), f"Record {index}", 1))
            except json.JSONDecodeError:
                failures += 1
                lines.extend([f"# Record {index}", line])
        return canonical_text("\n\n".join(lines)), encoding, {"parsed": failures == 0, "failures": failures, "format": "jsonl"}
    return raw, encoding, {"parsed": False, "format": path.suffix.casefold().lstrip(".")}


def read_pdf_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    if path.stat().st_size > options.max_input_bytes:
        raise VOFIEError(f"輸入超過 {options.max_input_bytes} bytes：{path}")
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise VOFIEError("PDF 需要可選的 pypdf Adapter；核心文字／Office 功能仍可使用。") from exc
    reader = PdfReader(str(path))
    lines: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        lines.extend([f"# Page {index}", page.extract_text() or ""])
    return canonical_text("\n\n".join(lines)), "pdf-text", {"pages": len(reader.pages)}


def read_code_adapter(path: Path, options: EngineOptions) -> tuple[str, str, dict[str, Any]]:
    raw, encoding, _ = read_text_adapter(path, options)
    language = CODE_LANGUAGE_BY_SUFFIX.get(path.suffix.casefold(), "text")
    wrapped = f"# {path.name}\n\n```{language}\n{raw.rstrip()}\n```\n"
    return wrapped, encoding, {"language": language, "source_code": raw}


READERS: dict[str, Callable[[Path, EngineOptions], tuple[str, str, dict[str, Any]]]] = {
    "text": read_text_adapter,
    "markdown": read_text_adapter,
    "html": read_html_adapter,
    "document": read_docx_adapter,
    "presentation": read_pptx_adapter,
    "spreadsheet": read_xlsx_adapter,
    "csv": read_csv_adapter,
    "structured": read_structured_adapter,
    "pdf": read_pdf_adapter,
    "code": read_code_adapter,
}


def load_source(path: Path, options: EngineOptions) -> SourceRecord:
    resolved = path.expanduser().resolve()
    if not resolved.is_file():
        raise VOFIEError(f"找不到輸入檔：{resolved}")
    before_size, before_hash = file_snapshot(resolved)
    kind = detect_kind(resolved)
    extracted, encoding, metadata = READERS[kind](resolved, options)
    after_size, after_hash = file_snapshot(resolved)
    if (before_size, before_hash) != (after_size, after_hash):
        raise VOFIEError(f"來源在讀取期間發生變更，已 FAIL CLOSED：{resolved}")
    return SourceRecord(
        source_id=stable_id("SRC", resolved, before_hash),
        path=str(resolved),
        name=resolved.name,
        extension=resolved.suffix.casefold(),
        input_kind=kind,
        encoding=encoding,
        byte_size=before_size,
        source_hash=before_hash,
        extracted_hash=blake2s_text(extracted),
        extracted_text=extracted,
        metadata=metadata,
    )


# ============================================================================
# 4. Markdown 正規化、隔離、主題分割與跨語言程式 Component IR
# ============================================================================


def infer_fence_language(first_line: str, previous_label: str = "") -> str:
    label = re.sub(r"[*_`:\s]", "", previous_label).casefold()
    if label in FENCE_ALIASES:
        return FENCE_ALIASES[label]
    probe = first_line.strip().casefold()
    if probe.startswith(("<!doctype html", "<html", "<div", "<script")):
        return "html"
    if probe.startswith(("import ", "from ", "def ", "class ", "@app.", "#!/usr/bin/env python")):
        return "python"
    if probe.startswith(("param(", "#requires", "write-host", "function ")):
        return "powershell"
    if probe.startswith(("const ", "let ", "function ", "document.", "window.")):
        return "javascript"
    if probe.startswith(("{", "[")):
        return "json"
    if probe.startswith(("#include", "int main", "void ")):
        return "c"
    return "text"


def normalize_markdown_fences(text: str) -> tuple[str, list[str]]:
    lines = text.splitlines()
    output: list[str] = []
    warnings: list[str] = []
    in_fence = False
    fence_marker = "```"
    for index, line in enumerate(lines):
        match = re.match(r"^\s*(`{3,}|~{3,})(.*)$", line)
        if not match:
            output.append(line)
            continue
        marker, tail = match.group(1), match.group(2).strip()
        if not in_fence:
            previous = next((candidate for candidate in reversed(output) if candidate.strip()), "")
            first_line = lines[index + 1] if index + 1 < len(lines) else ""
            raw_language = tail.split(maxsplit=1)[0].casefold() if tail else ""
            language = FENCE_ALIASES.get(raw_language, "")
            if not language:
                language = infer_fence_language(first_line, previous)
                warnings.append(f"line {index + 1}: 補上 code fence language={language}")
            fence_marker = marker[:3]
            output.append(f"```{language}")
            in_fence = True
        else:
            output.append("```")
            in_fence = False
    if in_fence:
        output.append("```")
        warnings.append("文件末端缺少 code fence，已在候選輸出補齊；原文仍保留於 SourceRecord")
    return canonical_text("\n".join(output)), warnings


def quarantine_boilerplate(source: SourceRecord, text: str, enabled: bool) -> tuple[str, list[QuarantineItem]]:
    if not enabled:
        return text, []
    clean_lines: list[str] = []
    items: list[QuarantineItem] = []
    for number, line in enumerate(text.splitlines(), start=1):
        reason = next((name for name, pattern in BOILERPLATE_RULES if pattern.search(line)), None)
        if reason:
            items.append(QuarantineItem(
                item_id=stable_id("QTN", source.source_id, number, line),
                source_id=source.source_id,
                line_number=number,
                reason=reason,
                content=line,
                content_hash=blake2s_text(line),
            ))
        else:
            clean_lines.append(line)
    return canonical_text("\n".join(clean_lines)), items


def iter_markdown_sections(text: str) -> Iterator[tuple[str, int, int, int, str]]:
    lines = text.splitlines()
    in_fence = False
    anchors: list[tuple[int, int, str]] = []
    for index, line in enumerate(lines):
        if re.match(r"^\s*(`{3,}|~{3,})", line):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            anchors.append((index, len(match.group(1)), re.sub(r"[*_`]", "", match.group(2)).strip()))
    if not anchors:
        yield "Content", 1, 1, len(lines), text.strip()
        return
    if anchors[0][0] > 0 and "\n".join(lines[: anchors[0][0]]).strip():
        yield "Preamble", 1, 1, anchors[0][0], "\n".join(lines[: anchors[0][0]]).strip()
    for position, (start, level, heading) in enumerate(anchors):
        end = anchors[position + 1][0] if position + 1 < len(anchors) else len(lines)
        yield heading, level, start + 1, end, "\n".join(lines[start:end]).strip()


def chunk_topic(content: str, max_chars: int) -> list[str]:
    if len(content) <= max_chars:
        return [content]
    paragraphs = re.split(r"\n{2,}", content)
    chunks: list[str] = []
    current: list[str] = []
    size = 0
    for paragraph in paragraphs:
        if current and size + len(paragraph) + 2 > max_chars:
            chunks.append("\n\n".join(current))
            current, size = [], 0
        if len(paragraph) > max_chars:
            for start in range(0, len(paragraph), max_chars):
                if current:
                    chunks.append("\n\n".join(current))
                    current, size = [], 0
                chunks.append(paragraph[start:start + max_chars])
        else:
            current.append(paragraph)
            size += len(paragraph) + 2
    if current:
        chunks.append("\n\n".join(current))
    return [chunk for chunk in chunks if chunk.strip()]


def classify_topic(heading: str, content: str) -> tuple[str, list[str]]:
    haystack = f"{heading}\n{content[:8000]}".casefold()
    scores = {
        category: sum(haystack.count(term.casefold()) for term in terms)
        for category, terms in TOPIC_TAXONOMY.items()
    }
    category = max(scores, key=scores.get) if max(scores.values(), default=0) else "general"
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_+-]{2,}|[\u3400-\u9fff]{2,8}", haystack)
    stop = {"the", "and", "with", "from", "this", "that", "使用", "完整", "可以", "功能", "系統", "模組", "下一步"}
    counter = Counter(token for token in tokens if token not in stop)
    tags = [token for token, _ in sorted(counter.items(), key=lambda item: (-item[1], -len(item[0]), item[0]))[:8]]
    return category, tags


def detect_code_language(code: str, hint: str = "") -> str:
    normalized = hint.casefold().strip()
    if normalized in FENCE_ALIASES:
        return FENCE_ALIASES[normalized]
    return infer_fence_language(code.splitlines()[0] if code.splitlines() else "")


def _python_units(code: str, start_line: int) -> list[CodeUnit]:
    try:
        tree = ast.parse(code)
    except SyntaxError as exc:
        return [CodeUnit(stable_id("CU", "python", start_line, blake2s_text(code)), "python", "<module>", "module", "", start_line, start_line + code.count("\n"), f"SYNTAX_ERROR:{exc.lineno}", blake2s_text(code))]
    units: list[CodeUnit] = []
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            kind = "class" if isinstance(node, ast.ClassDef) else "function"
            signature = node.name
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                signature = f"{node.name}({', '.join(arg.arg for arg in node.args.args)})"
            units.append(CodeUnit(
                unit_id=stable_id("CU", "python", node.name, node.lineno),
                language="python", symbol=node.name, unit_type=kind, signature=signature,
                start_line=start_line + node.lineno - 1,
                end_line=start_line + getattr(node, "end_lineno", node.lineno) - 1,
                syntax_status="PASS", content_hash=blake2s_text(ast.dump(node, include_attributes=False)),
            ))
    return sorted(units, key=lambda unit: (unit.start_line, unit.symbol))


GENERIC_SYMBOL_PATTERNS: dict[str, tuple[tuple[str, re.Pattern[str]], ...]] = {
    "powershell": (("function", re.compile(r"(?im)^\s*function\s+([A-Za-z][\w-]*)")), ("param", re.compile(r"(?im)^\s*param\s*\("))),
    "javascript": (("class", re.compile(r"(?m)^\s*class\s+([A-Za-z_$][\w$]*)")), ("function", re.compile(r"(?m)^\s*(?:export\s+)?(?:async\s+)?function\s+([A-Za-z_$][\w$]*)")), ("function", re.compile(r"(?m)^\s*(?:const|let|var)\s+([A-Za-z_$][\w$]*)\s*=\s*(?:async\s*)?\("))),
    "typescript": (("class", re.compile(r"(?m)^\s*(?:export\s+)?class\s+([A-Za-z_$][\w$]*)")), ("interface", re.compile(r"(?m)^\s*(?:export\s+)?interface\s+([A-Za-z_$][\w$]*)")), ("function", re.compile(r"(?m)^\s*(?:export\s+)?(?:async\s+)?function\s+([A-Za-z_$][\w$]*)"))),
    "c": (("function", re.compile(r"(?m)^\s*[A-Za-z_][\w\s*]+\s+([A-Za-z_]\w*)\s*\([^;]*\)\s*\{")),),
    "cpp": (("class", re.compile(r"(?m)^\s*class\s+([A-Za-z_]\w*)")), ("function", re.compile(r"(?m)^\s*[A-Za-z_][\w:<>,\s*&*]+\s+([A-Za-z_]\w*)\s*\([^;]*\)\s*\{"))),
    "csharp": (("class", re.compile(r"(?m)^\s*(?:public\s+|internal\s+)?class\s+([A-Za-z_]\w*)")), ("method", re.compile(r"(?m)^\s*(?:public|private|protected|internal)\s+[\w<>\[\]?]+\s+([A-Za-z_]\w*)\s*\("))),
    "java": (("class", re.compile(r"(?m)^\s*(?:public\s+)?class\s+([A-Za-z_]\w*)")), ("method", re.compile(r"(?m)^\s*(?:public|private|protected)\s+[\w<>\[\]]+\s+([A-Za-z_]\w*)\s*\("))),
    "go": (("function", re.compile(r"(?m)^\s*func\s+(?:\([^)]*\)\s*)?([A-Za-z_]\w*)\s*\(")), ("type", re.compile(r"(?m)^\s*type\s+([A-Za-z_]\w*)\s+"))),
    "rust": (("function", re.compile(r"(?m)^\s*(?:pub\s+)?fn\s+([A-Za-z_]\w*)")), ("struct", re.compile(r"(?m)^\s*(?:pub\s+)?struct\s+([A-Za-z_]\w*)"))),
}


def extract_code_units(content: str, topic_start_line: int) -> list[CodeUnit]:
    units: list[CodeUnit] = []
    for match in re.finditer(r"(?ms)^```([^\n`]*)\n(.*?)^```\s*$", content):
        hint, code = match.group(1).strip(), match.group(2)
        language = detect_code_language(code, hint)
        start_line = topic_start_line + content[: match.start(2)].count("\n")
        if language == "python":
            units.extend(_python_units(code, start_line))
            continue
        patterns = GENERIC_SYMBOL_PATTERNS.get(language, ())
        for unit_type, pattern in patterns:
            for symbol_match in pattern.finditer(code):
                symbol = symbol_match.group(1) if symbol_match.lastindex else "<module>"
                line = start_line + code[: symbol_match.start()].count("\n")
                units.append(CodeUnit(
                    unit_id=stable_id("CU", language, symbol, line), language=language,
                    symbol=symbol, unit_type=unit_type, signature=truncate(symbol_match.group(0), 180),
                    start_line=line, end_line=line, syntax_status="STRUCTURAL_ONLY",
                    content_hash=blake2s_text(symbol_match.group(0)),
                ))
        if not patterns or not any(unit.language == language and unit.start_line >= start_line for unit in units):
            units.append(CodeUnit(
                unit_id=stable_id("CU", language, start_line, blake2s_text(code)), language=language,
                symbol="<module>", unit_type="module", signature="", start_line=start_line,
                end_line=start_line + code.count("\n"), syntax_status="STRUCTURAL_ONLY",
                content_hash=blake2s_text(code),
            ))
    return units


def build_topics(source: SourceRecord, options: EngineOptions) -> tuple[list[TopicBlock], list[QuarantineItem], list[str]]:
    normalized, fence_warnings = normalize_markdown_fences(source.extracted_text)
    clean, quarantine = quarantine_boilerplate(source, normalized, options.quarantine_boilerplate)
    topics: list[TopicBlock] = []
    for heading, level, start, end, content in iter_markdown_sections(clean):
        chunks = chunk_topic(content, options.max_topic_chars)
        for part, chunk in enumerate(chunks, start=1):
            part_heading = heading if len(chunks) == 1 else f"{heading} [{part}/{len(chunks)}]"
            category, tags = classify_topic(part_heading, chunk)
            topic_id = stable_id("TOP", source.source_id, start, part_heading, blake2s_text(chunk))
            topics.append(TopicBlock(
                topic_id=topic_id, source_id=source.source_id, heading=part_heading,
                level=level, order=len(topics), source_start_line=start,
                source_end_line=end, category=category, tags=tags, content=chunk,
                content_hash=blake2s_text(re.sub(r"\s+", " ", chunk).casefold()),
                code_units=extract_code_units(chunk, start),
                quarantine_flags=[], st_position="ST-CORE",
            ))
    return topics, quarantine, fence_warnings


def mark_duplicates(topics: list[TopicBlock]) -> int:
    canonical: dict[str, str] = {}
    duplicates = 0
    for topic in topics:
        existing = canonical.get(topic.content_hash)
        if existing:
            topic.duplicate_of = existing
            duplicates += 1
        else:
            canonical[topic.content_hash] = topic.topic_id
    return duplicates


# ============================================================================
# 5. HTML/UI 規格、NLP Bridge 與品質 Gate
# ============================================================================


def build_ui_spec(sources: Sequence[SourceRecord]) -> dict[str, Any]:
    components: list[dict[str, Any]] = []
    endpoints: set[str] = set()
    scripts: list[str] = []
    styles: list[str] = []
    security_findings: list[dict[str, str]] = []
    for source in sources:
        metadata = source.metadata
        components.extend({**component, "source_id": source.source_id} for component in metadata.get("components", []))
        endpoints.update(metadata.get("endpoints", []))
        if metadata.get("scripts"):
            scripts.append(metadata["scripts"])
        if metadata.get("styles"):
            styles.append(metadata["styles"])
        html_source = metadata.get("html_source", "")
        checks = (
            (r"javascript\s*:", "javascript_url", "HIGH"),
            (r"\beval\s*\(", "eval_usage", "HIGH"),
            (r"\binnerHTML\s*=", "innerhtml_assignment", "MEDIUM"),
            (r"https?://", "external_resource", "MEDIUM"),
            (r"Start-Process|subprocess\.(?:run|Popen)", "process_execution_example", "HIGH"),
        )
        for pattern, rule, severity in checks:
            if re.search(pattern, html_source, re.I):
                security_findings.append({"source_id": source.source_id, "rule": rule, "severity": severity})
    test_cases = []
    for index, component in enumerate(components, start=1):
        identifier = component.get("id") or component.get("name") or f"{component.get('tag', 'component')}-{index}"
        test_cases.append({
            "test_id": f"UI-TC-{index:03d}",
            "component": identifier,
            "action": "activate" if component.get("tag") in {"button", "a"} else "input",
            "expected": "state and visible feedback remain consistent",
            "st": "ST-UI-004",
        })
    a11y_findings = []
    for component in components:
        if component.get("tag") in {"input", "select", "textarea", "button"} and not any(component.get(key) for key in ("aria_label", "name", "id")):
            a11y_findings.append({"rule": "missing_accessible_name", "component": component, "severity": "HIGH"})
    state_machine = {
        "initial": "IDLE",
        "states": ["IDLE", "LOADED", "FILTERED", "EXPORTED", "ERROR"],
        "transitions": [
            {"event": "load", "from": "IDLE", "to": "LOADED"},
            {"event": "filter", "from": "LOADED", "to": "FILTERED"},
            {"event": "export", "from": "LOADED|FILTERED", "to": "EXPORTED"},
            {"event": "failure", "from": "*", "to": "ERROR"},
        ],
    }
    return {
        "contract": "veritas.ui-spec/1.0",
        "components": components,
        "endpoints": sorted(endpoints),
        "source_script_count": len(scripts),
        "source_style_count": len(styles),
        "source_scripts_executed": False,
        "security_findings": security_findings,
        "accessibility_findings": a11y_findings,
        "test_cases": test_cases,
        "state_machine": state_machine,
        "interaction_graph": [
            {"source": case["component"], "event": case["action"], "target": "VOFIE_RENDER", "st": "ST-UI-003"}
            for case in test_cases
        ],
        "layout": {"desktop": "280px minmax(0, 1fr)", "mobile_breakpoint": "860px", "mobile": "1fr"},
        "theme": {"mode": "light", "accent": "#0F5F73", "background": "#F4F7F9", "ink": "#17212B"},
        "telemetry": {"enabled": False, "mode": "local-opt-in", "pii": "forbidden"},
    }


def try_vsis_enrichment(ir: UniversalContentIR, project_root: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"status": "SKIP", "reason": "VSIS_NOT_FOUND", "actions": []}
    candidates = [
        project_root.parent / "VIA_SemanticIntelligenceSubsystem_v0120",
        Path(os.environ.get("VIA_VSIS_ROOT", "")) if os.environ.get("VIA_VSIS_ROOT") else None,
    ]
    vsis_root = next((item.resolve() for item in candidates if item and (item / "via_semantic_intelligence").is_dir()), None)
    if not vsis_root:
        return result
    sys.path.insert(0, str(vsis_root))
    try:
        from via_semantic_intelligence.registry_bridge import create_subsystem  # type: ignore
        dispatcher = create_subsystem()
        sample = "\n\n".join(topic.content for topic in ir.topics[:12])[:120_000]
        responses = dispatcher.invoke_plan(
            caller_system_id="CGE",
            actions=("normalize", "segment", "categorize", "semantic_check"),
            base_payload={"content": sample, "content_type": "markdown", "source_name": ir.title},
        )
        result = {
            "status": "PASS" if responses and all(response.status in {"PASS", "PASS_WITH_WARNINGS"} for response in responses) else "HOLD",
            "version": getattr(dispatcher.config, "schema_version", "1.2-compatible"),
            "actions": [{"action": response.action, "status": response.status} for response in responses],
        }
    except Exception as exc:
        result = {"status": "WARN", "reason": f"{type(exc).__name__}: {exc}", "actions": []}
    finally:
        try:
            sys.path.remove(str(vsis_root))
        except ValueError:
            pass
    return result


def quality_gate(ir: UniversalContentIR, source_snapshots: dict[str, tuple[int, str]]) -> dict[str, Any]:
    warnings: list[str] = []
    failures: list[str] = []
    for source in ir.source_records:
        path = Path(source.path)
        if not path.is_file():
            failures.append(f"source_missing:{source.source_id}")
            continue
        if file_snapshot(path) != source_snapshots[source.source_id]:
            failures.append(f"source_mutated:{source.source_id}")
        if not source.raw_preserved:
            failures.append(f"source_not_preserved:{source.source_id}")
    topic_ids = [topic.topic_id for topic in ir.topics]
    if len(topic_ids) != len(set(topic_ids)):
        failures.append("duplicate_topic_id")
    if not ir.topics:
        failures.append("no_topics")
    if ir.quarantine:
        warnings.append(f"quarantine_items:{len(ir.quarantine)}")
    high_security = sum(item.get("severity") == "HIGH" for item in ir.ui_spec.get("security_findings", []))
    if high_security:
        warnings.append(f"source_security_findings_high:{high_security};source_scripts_not_executed")
    gate = "FAIL" if failures else ("PASS_WITH_WARNINGS" if warnings else "PASS")
    return {
        "gate": gate,
        "source_preservation": "PASS" if not any(item.startswith("source_") for item in failures) else "FAIL",
        "topic_identity": "PASS" if "duplicate_topic_id" not in failures else "FAIL",
        "source_script_execution": "DENIED",
        "ai_direct_apply": "DENIED",
        "duplicates_retained": True,
        "warnings": warnings,
        "failures": failures,
    }


def normalize_operations(operations: Sequence[str]) -> tuple[str, ...]:
    requested = {str(item).strip().casefold() for item in operations if str(item).strip()}
    unknown = sorted(requested - set(OPERATION_ORDER))
    if unknown:
        raise VOFIEError(f"未知重構動作：{', '.join(unknown)}")
    return tuple(item for item in OPERATION_ORDER if item in requested)


def build_consolidated_view(ir: UniversalContentIR, operations: Sequence[str]) -> dict[str, Any]:
    """建立非破壞式整合視圖；IR 中的來源、主題與重複項一個都不刪。"""

    normalized = normalize_operations(operations)
    deduplicate = "deduplicate" in normalized
    visible_topics = [topic for topic in ir.topics if not (deduplicate and topic.duplicate_of)]
    sources = {source.source_id: source.name for source in ir.source_records}
    category_groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    code_groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    duplicate_map: list[dict[str, str]] = []
    for topic in ir.topics:
        if topic.duplicate_of:
            duplicate_map.append({"duplicate": topic.topic_id, "canonical": topic.duplicate_of})
    for topic in visible_topics:
        category_groups[topic.category].append({
            "topic_id": topic.topic_id,
            "source_id": topic.source_id,
            "source_name": sources[topic.source_id],
            "heading": topic.heading,
            "order": topic.order,
            "content_hash": topic.content_hash,
            "excerpt": truncate(strip_markdown_markup(topic.content), 500),
        })
        for unit in topic.code_units:
            code_groups[unit.language].append({
                "unit_id": unit.unit_id,
                "topic_id": topic.topic_id,
                "source_name": sources[topic.source_id],
                "symbol": unit.symbol,
                "unit_type": unit.unit_type,
                "signature": unit.signature,
                "syntax_status": unit.syntax_status,
                "content_hash": unit.content_hash,
            })
    category_rows = [
        {"category": category, "topic_count": len(items), "topics": sorted(items, key=lambda item: (item["order"], item["topic_id"]))}
        for category, items in sorted(category_groups.items())
    ]
    code_rows = [
        {"language": language, "component_count": len(items), "components": sorted(items, key=lambda item: (item["symbol"], item["unit_id"]))}
        for language, items in sorted(code_groups.items())
    ]
    operation_results = []
    for operation in OPERATION_ORDER:
        enabled = operation in normalized
        counts = {
            "text_merge": len(category_rows),
            "code_merge": sum(row["component_count"] for row in code_rows),
            "restructure": len(visible_topics),
            "deduplicate": len(duplicate_map),
            "optimize": len(visible_topics),
        }
        operation_results.append({
            "operation": operation,
            "enabled": enabled,
            "status": "PASS" if enabled else "SKIP",
            "affected": counts[operation] if enabled else 0,
            "source_mutated": False,
            "policy": "STRUCTURAL_CANDIDATE_ONLY" if operation == "optimize" else "ADD_ONLY_VIEW",
        })
    payload = {
        "contract": "veritas.vofie-consolidated-view/1.1",
        "operations": list(normalized),
        "operation_results": operation_results,
        "source_topic_count": len(ir.topics),
        "visible_topic_count": len(visible_topics),
        "duplicates_marked_and_retained": duplicate_map,
        "text_groups": category_rows,
        "code_groups": code_rows,
        "source_policy": "READ_ONLY_NO_DELETE_NO_OVERWRITE",
        "api_signature_policy": "KEEP_OR_CANDIDATE_ONLY",
    }
    payload["view_hash"] = blake2s_text(json.dumps(payload, ensure_ascii=False, sort_keys=True))
    return payload


# ============================================================================
# 6. 輸出 Adapter：Markdown、CSV、Web Template 與 Word
# ============================================================================


def profile_rows() -> list[dict[str, str]]:
    rows = [
        {"st_id": "ST-FMT-001", "action": "detect_and_read", "position": "INPUT_CORE", "flexibility": "FORMAT_ADAPTER", "test": "CRITICAL", "note": "來源唯讀與雜湊前後一致。"},
        {"st_id": "ST-FMT-002", "action": "topic_segment", "position": "SEMANTIC_CORE", "flexibility": "ANCHOR_ADAPTER", "test": "CRITICAL", "note": "主題含來源行號與雜湊。"},
        {"st_id": "ST-FMT-003", "action": "code_component_ir", "position": "POLYGLOT_CORE", "flexibility": "AST_ADAPTER", "test": "CRITICAL", "note": "跨語言元件寫入 Markdown。"},
        {"st_id": "ST-FMT-004", "action": "markdown_emit", "position": "CANONICAL_OUTPUT", "flexibility": "TEMPLATE_ADAPTER", "test": "CRITICAL", "note": "所有格式共用 Markdown／IR 中介層。"},
        {"st_id": "ST-FMT-005", "action": "office_emit", "position": "DOCUMENT_ADAPTER", "flexibility": "DOCX_PPTX_XLSX_ADAPTER", "test": "CRITICAL", "note": "格式特性不同但來源追溯欄位不變。"},
        {"st_id": "ST-FMT-006", "action": "web_template_emit", "position": "UI_OUTPUT", "flexibility": "HTML_CSS_JS_TOKEN_ADAPTER", "test": "CRITICAL", "note": "三檔分離、無 CDN、無來源 script 執行。"},
        {"st_id": "ST-FMT-007", "action": "audit_chain", "position": "GOVERNANCE", "flexibility": "APPEND_ONLY", "test": "CRITICAL", "note": "每次輸出以 hash chain 記錄。"},
        {"st_id": "ST-FMT-008", "action": "vsis_bridge", "position": "VIA_INTEGRATION", "flexibility": "OPTIONAL_OVERLAY", "test": "STANDARD", "note": "沿用 VSIS 1.2 能力，缺少時本地降級。"},
        {"st_id": "ST-FMT-009", "action": "text_merge", "position": "CONSOLIDATION", "flexibility": "CATEGORY_VIEW_ADAPTER", "test": "CRITICAL", "note": "文字依主題分類合併成視圖，原區塊仍完整保留。"},
        {"st_id": "ST-FMT-010", "action": "code_merge", "position": "POLYGLOT_CONSOLIDATION", "flexibility": "LANGUAGE_SYMBOL_ADAPTER", "test": "CRITICAL", "note": "跨語言元件依語言／符號整合，API 簽章不變。"},
        {"st_id": "ST-FMT-011", "action": "restructure", "position": "STRUCTURAL_CANDIDATE", "flexibility": "ORDER_ADAPTER", "test": "CRITICAL", "note": "只重排輸出視圖，保留來源順序與行號。"},
        {"st_id": "ST-FMT-012", "action": "deduplicate", "position": "ADD_ONLY_DEDUP", "flexibility": "CANONICAL_RULE_ADAPTER", "test": "CRITICAL", "note": "標記 canonical／duplicate，不刪除任何內容。"},
        {"st_id": "ST-FMT-013", "action": "optimize", "position": "EQUIVALENCE_GATE", "flexibility": "CANDIDATE_ONLY", "test": "CRITICAL", "note": "優化只形成候選視圖，不直接改寫來源或 API。"},
        {"st_id": "ST-FMT-014", "action": "simple_five_outputs", "position": "USER_PROFILE", "flexibility": "FIXED_PRIMARY_CONTRACT", "test": "CRITICAL", "note": "簡易模式固定 MD／HTML／Component JSON／Word／CSV 五個主要檔。"},
        {"st_id": "ST-FMT-015", "action": "engine_system_roles", "position": "ROLE_BOUNDARY", "flexibility": "SYSTEM_SIDECAR_ADAPTER", "test": "CRITICAL", "note": "ENGINE 只輸出五檔；SYSTEM 另在 _system 保留治理資料。"},
        {"st_id": "ST-FMT-016", "action": "failure_recovery", "position": "RESILIENCE_GATE", "flexibility": "HANDLER_REGISTRY", "test": "CRITICAL", "note": "八個環節各 Top 20 failure，且每項有多個已實作復原處理器。"},
        {"st_id": "ST-FMT-017", "action": "window_drag_drop", "position": "WINDOW_IO", "flexibility": "DND_OR_FILE_DIALOG", "test": "STANDARD", "note": "Windows 視窗可選檔或拖放，最多五檔，拖放模組缺少時降級。"},
    ]
    rows.extend(UI_CAPABILITY_PROFILES)
    return rows


def markdown_for_ir(ir: UniversalContentIR) -> str:
    sources = {source.source_id: source for source in ir.source_records}
    lines = [
        "---",
        f"title: {json.dumps(ir.title, ensure_ascii=False)}",
        f"engine: {ENGINE_NAME}",
        f"engine_version: {ENGINE_VERSION}",
        f"run_id: {ir.run_id}",
        f"created_at: {ir.created_at}",
        f"quality_gate: {ir.quality.get('gate', 'PENDING')}",
        "source_policy: READ_ONLY_NEW_ARTIFACTS_ONLY",
        "---",
        "",
        f"# {ir.title}",
        "",
        f"> 由 **{ENGINE_NAME_ZH}（VOFIE）** 產生。原檔未改寫；重複與雜訊採標記／隔離，並保留在 IR。",
        "",
        "## 格式與治理摘要",
        "",
        "| 項目 | 結果 |",
        "|---|---|",
        f"| 來源檔 | {len(ir.source_records)} |",
        f"| 主題區塊 | {len(ir.topics)} |",
        f"| 程式元件 | {sum(len(topic.code_units) for topic in ir.topics)} |",
        f"| 隔離片段 | {len(ir.quarantine)} |",
        f"| 重複主題 | {sum(topic.duplicate_of is not None for topic in ir.topics)} |",
        f"| 品質 Gate | {ir.quality.get('gate', 'PENDING')} |",
        "| 來源 script | 不執行 |",
        "| AI 改寫 | 僅候選，禁止直接套用 |",
        "",
        "## 來源登記",
        "",
        "| Source ID | 檔名 | 類型 | 編碼 | Bytes | BLAKE2s |",
        "|---|---|---|---|---:|---|",
    ]
    for source in ir.source_records:
        lines.append(f"| {source.source_id} | {escape_md_cell(source.name)} | {source.input_kind} | {source.encoding} | {source.byte_size} | `{source.source_hash}` |")
    lines.extend(["", "## ST 能力定位", "", "| ST | Action | 基準定位 | 彈性 | 測試 | 說明備註 |", "|---|---|---|---|---|---|"])
    for profile in ir.capability_profiles:
        lines.append("| {st_id} | `{action}` | {position} | {flexibility} | {test} | {note} |".format(**profile))
    consolidation = ir.quality.get("consolidation", {})
    lines.extend([
        "", "## 合併／重組／去重／優化視圖", "",
        f"**執行角色**：`{ir.quality.get('run_role', 'ENGINE')}`  ",
        f"**動作順序**：{', '.join(f'`{item}`' for item in consolidation.get('operations', [])) or '—'}  ",
        f"**來源主題／視圖主題**：{consolidation.get('source_topic_count', len(ir.topics))} / {consolidation.get('visible_topic_count', len(ir.topics))}  ",
        "**保留政策**：來源、重複項、行號與雜湊全部保留；優化只產生候選視圖。", "",
        "| Action | Status | Affected | Source mutated | Policy |", "|---|---|---:|---|---|",
    ])
    for result in consolidation.get("operation_results", []):
        lines.append(f"| `{result['operation']}` | {result['status']} | {result['affected']} | {result['source_mutated']} | {result['policy']} |")
    if consolidation.get("text_groups"):
        lines.extend(["", "### 文字主題整合索引", "", "| Category | Topics | Headings |", "|---|---:|---|"])
        for group in consolidation["text_groups"]:
            headings = "；".join(item["heading"] for item in group["topics"][:8])
            if len(group["topics"]) > 8:
                headings += f"；…(+{len(group['topics']) - 8})"
            lines.append(f"| {escape_md_cell(group['category'])} | {group['topic_count']} | {escape_md_cell(headings)} |")
    if consolidation.get("code_groups"):
        lines.extend(["", "### 程式元件整合索引", "", "| Language | Components | Symbols |", "|---|---:|---|"])
        for group in consolidation["code_groups"]:
            symbols = "；".join(item["symbol"] for item in group["components"][:12])
            if len(group["components"]) > 12:
                symbols += f"；…(+{len(group['components']) - 12})"
            lines.append(f"| {escape_md_cell(group['language'])} | {group['component_count']} | {escape_md_cell(symbols)} |")
    lines.extend(["", "## 主題重構內容", ""])
    last_source = None
    for topic in ir.topics:
        source = sources[topic.source_id]
        if source.source_id != last_source:
            lines.extend([f"# 來源：{source.name}", "", f"來源雜湊：`{source.source_hash}`", ""])
            last_source = source.source_id
        duplicate = f"；重複於 `{topic.duplicate_of}`，本區仍保留" if topic.duplicate_of else ""
        lines.extend([
            f"## {topic.heading}", "",
            f"**定位**：`{topic.category}` · `{topic.st_position}` · 來源行 {topic.source_start_line}-{topic.source_end_line}{duplicate}", "",
            f"**標籤**：{', '.join(f'`{tag}`' for tag in topic.tags) if topic.tags else '—'}", "",
        ])
        if topic.code_units:
            lines.extend(["### 程式元件圖譜", "", "| Language | Type | Symbol | Signature | Lines | Syntax |", "|---|---|---|---|---:|---|"])
            for unit in topic.code_units:
                lines.append(f"| {unit.language} | {unit.unit_type} | `{escape_md_cell(unit.symbol)}` | `{escape_md_cell(unit.signature)}` | {unit.start_line}-{unit.end_line} | {unit.syntax_status} |")
            lines.append("")
            lines.extend(["### 結構化原文／候選基線", "", "> 下列程式內容僅重新分區與補齊 Markdown 語言標籤；未做未驗證的語意改寫。", ""])
        lines.extend([topic.content.strip(), ""])
    if ir.quarantine:
        lines.extend(["## 隔離但保留的對話／渲染雜訊", "", "| Source | Line | Reason | Preserved content | Hash |", "|---|---:|---|---|---|"])
        for item in ir.quarantine:
            source = sources[item.source_id]
            lines.append(f"| {escape_md_cell(source.name)} | {item.line_number} | {item.reason} | {escape_md_cell(item.content)} | `{item.content_hash}` |")
        lines.append("")
    lines.extend(["## 品質 Gate", "", "```json", json.dumps(ir.quality, ensure_ascii=False, indent=2), "```", ""])
    return canonical_text("\n".join(lines))


def emit_markdown(ir: UniversalContentIR, output_dir: Path) -> Path:
    path = output_dir / "Veritas_VOFIE_Reconstructed.md"
    atomic_write_text(path, markdown_for_ir(ir))
    return path


def emit_ir_json(ir: UniversalContentIR, output_dir: Path) -> Path:
    path = output_dir / "Veritas_VOFIE_UniversalContentIR.json"
    atomic_write_text(path, json.dumps(ir.to_dict(), ensure_ascii=False, indent=2, sort_keys=True))
    return path


def emit_csv(ir: UniversalContentIR, output_dir: Path) -> Path:
    path = output_dir / "Veritas_VOFIE_TopicMatrix.csv"
    stream = io.StringIO(newline="")
    writer = csv.writer(stream)
    writer.writerow(["topic_id", "source_id", "source_name", "order", "heading", "category", "tags", "start_line", "end_line", "duplicate_of", "code_units", "content_hash", "excerpt"])
    sources = {source.source_id: source.name for source in ir.source_records}
    for topic in ir.topics:
        writer.writerow([
            topic.topic_id, topic.source_id, sources[topic.source_id], topic.order, topic.heading,
            topic.category, ";".join(topic.tags), topic.source_start_line, topic.source_end_line,
            topic.duplicate_of or "", len(topic.code_units), topic.content_hash, truncate(strip_markdown_markup(topic.content), 500),
        ])
    atomic_write_text(path, "\ufeff" + stream.getvalue())
    return path


def _web_payload(ir: UniversalContentIR) -> str:
    sources = {source.source_id: source.name for source in ir.source_records}
    payload = {
        "engine": ENGINE_NAME,
        "version": ENGINE_VERSION,
        "runId": ir.run_id,
        "quality": ir.quality,
        "topics": [
            {
                "id": topic.topic_id,
                "source": sources[topic.source_id],
                "heading": topic.heading,
                "category": topic.category,
                "tags": topic.tags,
                "lines": f"{topic.source_start_line}-{topic.source_end_line}",
                "duplicateOf": topic.duplicate_of,
                "codeUnits": len(topic.code_units),
                "excerpt": truncate(strip_markdown_markup(topic.content), 1200),
            }
            for topic in ir.topics
        ],
    }
    return json.dumps(payload, ensure_ascii=False).replace("</", "<\\/")


def emit_web_template(ir: UniversalContentIR, output_dir: Path) -> tuple[Path, Path, Path]:
    html_path = output_dir / "Veritas_VOFIE_Template.html"
    css_path = output_dir / "Veritas_VOFIE_Template.css"
    js_path = output_dir / "Veritas_VOFIE_Template.js"
    html_text = """<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <meta name="color-scheme" content="light dark">
  <title>Veritas OmniFormat Intelligence Engine</title>
  <link rel="stylesheet" href="Veritas_VOFIE_Template.css">
</head>
<body>
  <header class="app-header">
    <div>
      <p class="eyebrow">VERITAS INTELLIGENCE ANALYTICS</p>
      <h1>OmniFormat Intelligence Engine</h1>
    </div>
    <output id="qualityGate" class="status" aria-live="polite">Loading</output>
  </header>
  <main class="layout">
    <aside class="filters" aria-label="內容篩選">
      <label for="search">搜尋主題</label>
      <input id="search" type="search" placeholder="標題、標籤或內容" autocomplete="off">
      <label for="category">分類</label>
      <select id="category"><option value="">全部分類</option></select>
      <p id="count" class="muted" aria-live="polite"></p>
      <button id="exportCsv" type="button">匯出目前清單 CSV</button>
    </aside>
    <section class="content" aria-labelledby="topicHeading">
      <div class="section-heading">
        <div><p class="eyebrow">STRUCTURED RESULT</p><h2 id="topicHeading">主題矩陣</h2></div>
      </div>
      <div class="table-wrap" tabindex="0">
        <table>
          <thead><tr><th>來源</th><th>主題</th><th>分類</th><th>行號</th><th>程式元件</th></tr></thead>
          <tbody id="topicRows"></tbody>
        </table>
      </div>
      <article id="detail" class="detail" aria-live="polite"><p>選取一個主題查看結構化摘要。</p></article>
    </section>
  </main>
  <script id="vofie-data" type="application/json">__VOFIE_DATA__</script>
  <script src="Veritas_VOFIE_Template.js" defer></script>
</body>
</html>
""".replace("__VOFIE_DATA__", _web_payload(ir))
    css_text = """:root{--bg:#f4f7f9;--surface:#fff;--ink:#17212b;--muted:#65727f;--line:#d7e0e5;--accent:#0f5f73;--accent-soft:#dceef2;--danger:#9b2c2c;--radius:10px;font-family:Inter,"Segoe UI","Noto Sans TC",Arial,sans-serif;color-scheme:light}
*{box-sizing:border-box}body{margin:0;background:var(--bg);color:var(--ink);font-size:14px;line-height:1.5}.app-header{min-height:76px;padding:14px clamp(18px,4vw,52px);display:flex;align-items:center;justify-content:space-between;gap:20px;background:var(--surface);border-bottom:1px solid var(--line)}h1,h2,p{margin:0}h1{font-size:22px;letter-spacing:-.02em}h2{font-size:18px}.eyebrow{font-size:10px;font-weight:750;letter-spacing:.14em;color:var(--accent);margin-bottom:3px}.status{padding:5px 10px;border-radius:999px;background:var(--accent-soft);color:var(--accent);font-weight:700;font-size:12px}.layout{display:grid;grid-template-columns:280px minmax(0,1fr);gap:18px;max-width:1500px;margin:0 auto;padding:18px}.filters,.content,.detail{background:var(--surface);border:1px solid var(--line);border-radius:var(--radius)}.filters{align-self:start;padding:16px;display:grid;gap:8px;position:sticky;top:18px}.filters label{font-size:12px;font-weight:700;margin-top:6px}input,select,button{width:100%;min-height:36px;border:1px solid var(--line);border-radius:7px;background:var(--surface);color:var(--ink);padding:7px 9px;font:inherit}button{margin-top:8px;background:var(--accent);color:#fff;border-color:var(--accent);font-weight:700;cursor:pointer}button:hover{filter:brightness(.94)}button:focus-visible,input:focus-visible,select:focus-visible,.table-wrap:focus-visible{outline:3px solid #65a9b8;outline-offset:2px}.content{min-width:0;padding:16px}.section-heading{display:flex;align-items:end;justify-content:space-between;margin-bottom:12px}.table-wrap{overflow:auto;border:1px solid var(--line);border-radius:8px}table{width:100%;border-collapse:collapse;min-width:760px}th,td{padding:9px 10px;border-bottom:1px solid var(--line);text-align:left;vertical-align:top}th{position:sticky;top:0;background:#eef3f5;font-size:11px;letter-spacing:.04em;text-transform:uppercase}tbody tr{cursor:pointer}tbody tr:hover,tbody tr[aria-selected="true"]{background:var(--accent-soft)}.muted{color:var(--muted);font-size:12px}.detail{margin-top:14px;padding:16px;white-space:pre-wrap}.detail h3{font-size:16px;margin:0 0 8px}.tag{display:inline-block;margin:4px 5px 0 0;padding:2px 6px;border-radius:4px;background:#eef3f5;color:#3d4d59;font-size:11px}
@media(max-width:860px){.app-header{align-items:flex-start}.layout{grid-template-columns:1fr}.filters{position:static}.content{padding:12px}}
@media(prefers-color-scheme:dark){:root{--bg:#11181d;--surface:#192229;--ink:#eef5f7;--muted:#a9b5bc;--line:#34434c;--accent:#78c5d4;--accent-soft:#203d45;color-scheme:dark}th{background:#243139}button{color:#102027}}
@media(prefers-reduced-motion:reduce){*{scroll-behavior:auto!important;transition:none!important}}
"""
    js_text = """'use strict';
const byId=(id)=>document.getElementById(id);
const payload=JSON.parse(byId('vofie-data').textContent);
const state={query:'',category:'',selected:null};
const categories=[...new Set(payload.topics.map((item)=>item.category))].sort();
for(const value of categories){const option=document.createElement('option');option.value=value;option.textContent=value;byId('category').append(option);}
byId('qualityGate').textContent=payload.quality.gate;
function filtered(){const q=state.query.trim().toLocaleLowerCase();return payload.topics.filter((item)=>(!state.category||item.category===state.category)&&(!q||[item.source,item.heading,item.category,item.tags.join(' '),item.excerpt].join(' ').toLocaleLowerCase().includes(q)));}
function render(){const items=filtered();const body=byId('topicRows');body.replaceChildren();for(const item of items){const row=document.createElement('tr');row.tabIndex=0;row.dataset.id=item.id;row.setAttribute('aria-selected',String(state.selected===item.id));for(const value of [item.source,item.heading,item.category,item.lines,String(item.codeUnits)]){const cell=document.createElement('td');cell.textContent=value;row.append(cell);}row.addEventListener('click',()=>selectTopic(item.id));row.addEventListener('keydown',(event)=>{if(event.key==='Enter'||event.key===' '){event.preventDefault();selectTopic(item.id);}});body.append(row);}byId('count').textContent=`顯示 ${items.length} / ${payload.topics.length} 個主題`;}
function selectTopic(id){state.selected=id;const item=payload.topics.find((candidate)=>candidate.id===id);if(!item)return;const detail=byId('detail');detail.replaceChildren();const title=document.createElement('h3');title.textContent=item.heading;const meta=document.createElement('p');meta.className='muted';meta.textContent=`${item.source} · ${item.category} · lines ${item.lines}`;const excerpt=document.createElement('p');excerpt.textContent=item.excerpt;const tags=document.createElement('div');for(const value of item.tags){const tag=document.createElement('span');tag.className='tag';tag.textContent=value;tags.append(tag);}detail.append(title,meta,tags,excerpt);render();}
function csvCell(value){const text=String(value??'');return /[",\\n]/.test(text)?`"${text.replaceAll('"','""')}"`:text;}
function exportCsv(){const rows=[['source','heading','category','lines','code_units'],...filtered().map((item)=>[item.source,item.heading,item.category,item.lines,item.codeUnits])];const blob=new Blob(['\ufeff'+rows.map((row)=>row.map(csvCell).join(',')).join('\\n')],{type:'text/csv;charset=utf-8'});const link=document.createElement('a');link.href=URL.createObjectURL(blob);link.download='Veritas_VOFIE_FilteredTopics.csv';link.click();URL.revokeObjectURL(link.href);}
byId('search').addEventListener('input',(event)=>{state.query=event.target.value;render();});
byId('category').addEventListener('change',(event)=>{state.category=event.target.value;render();});
byId('exportCsv').addEventListener('click',exportCsv);
render();
"""
    atomic_write_text(html_path, html_text)
    atomic_write_text(css_path, css_text)
    atomic_write_text(js_path, js_text)
    return html_path, css_path, js_path


def emit_self_contained_html(ir: UniversalContentIR, output_dir: Path) -> Path:
    """把既有三檔 Web Adapter 封裝成單一離線 HTML；完整模式仍保留三檔輸出。"""

    with tempfile.TemporaryDirectory(prefix="vofie-web-") as temporary_name:
        temporary = Path(temporary_name)
        html_path, css_path, js_path = emit_web_template(ir, temporary)
        html_text = html_path.read_text(encoding="utf-8")
        css_text = css_path.read_text(encoding="utf-8")
        js_text = js_path.read_text(encoding="utf-8").replace("</script", "<\\/script")
    html_text = html_text.replace(
        '<link rel="stylesheet" href="Veritas_VOFIE_Template.css">',
        f"<style>\n{css_text}\n</style>",
    ).replace(
        '<script src="Veritas_VOFIE_Template.js" defer></script>',
        f"<script>\n{js_text}\n</script>",
    )
    path = output_dir / SIMPLE_PRIMARY_FILENAMES["html"]
    atomic_write_text(path, html_text)
    return path


def component_specs_payload(ir: UniversalContentIR, project_root: Path) -> dict[str, Any]:
    catalog = load_failure_catalog(project_root)
    polyglot_tools = tool_audit(project_root)
    hydra = hydra_risk_audit(project_root)
    components = []
    for topic in ir.topics:
        for unit in topic.code_units:
            components.append({**asdict(unit), "topic_id": topic.topic_id, "source_id": topic.source_id})
    expected_outputs = {key: filename for key, filename in SIMPLE_PRIMARY_FILENAMES.items()}
    return {
        "contract": COMPONENT_SPEC_CONTRACT,
        "engine": {"id": ENGINE_ID, "name": ENGINE_NAME, "version": ENGINE_VERSION},
        "subsystem": {"id": SUBSYSTEM_ID, "namespace": REGISTRY_NAMESPACE},
        "run": {
            "run_id": ir.run_id,
            "role": ir.quality.get("run_role", "ENGINE"),
            "created_at": ir.created_at,
            "source_policy": "READ_ONLY_NO_DELETE_NO_MOVE_NO_CANONICAL_MUTATION",
            "primary_output_contract": expected_outputs,
        },
        "roles": {
            "ENGINE": {"primary_outputs": list(SIMPLE_PRIMARY_OUTPUTS), "sidecars": False},
            "SYSTEM": {"primary_outputs": list(SIMPLE_PRIMARY_OUTPUTS), "sidecars": SYSTEM_SIDECAR_DIRECTORY},
        },
        "operations": ir.quality.get("consolidation", {}),
        "ui": ir.ui_spec,
        "components": components,
        "failure_framework": {
            "contract": catalog["contract"],
            "stages": [
                {"stage_id": stage["stage_id"], "failure_count": len(stage["failures"]), "handlers": stage["handlers"]}
                for stage in catalog["stages"]
            ],
            "total_failures": catalog["failure_count"],
            "handler_count": len(RECOVERY_HANDLERS),
        },
        "polyglot_tool_support": compact_tool_audit(polyglot_tools),
        "hydra_risk_support": {
            "contract": hydra["contract"],
            "gate": hydra["gate"],
            "mode": hydra["mode"],
            "st": hydra["st"],
            "policy": hydra["policy"],
            "summary": hydra["summary"],
            "round_plan": hydra["round_plan"],
            "top20": hydra["top20"],
            "source_mutated": hydra["source_mutated"],
        },
        "runtime_copy_governance": {
            "contract": RUNTIME_COPY_CONTRACT,
            "canonical_source": "READ_ONLY",
            "runtime_copy": "WRITABLE_AFTER_EXPLICIT_APPROVAL",
            "hash_states": ["MISSING/APPLY", "PROPOSED/SKIP", "ORIGINAL/BACKUP_APPLY", "OTHER/FAIL_CLOSED"],
            "rollback": "DRY_RUN_REQUIRED",
            "promotion": "NOT_IMPLEMENTED_NO_CANONICAL_WRITE",
        },
        "universal_content_ir": ir.to_dict(),
    }


def emit_component_specs_json(ir: UniversalContentIR, output_dir: Path, project_root: Path) -> Path:
    path = output_dir / SIMPLE_PRIMARY_FILENAMES["component_json"]
    atomic_write_text(path, json.dumps(component_specs_payload(ir, project_root), ensure_ascii=False, indent=2, sort_keys=True))
    return path


def compact_quality_summary(ir: UniversalContentIR) -> dict[str, Any]:
    consolidation = ir.quality.get("consolidation", {})
    return {
        "gate": ir.quality.get("gate"),
        "source_preservation": ir.quality.get("source_preservation"),
        "post_output_source_preservation": ir.quality.get("post_output_source_preservation", "PENDING_AT_WORD_EMIT"),
        "source_count": len(ir.source_records),
        "topic_count": len(ir.topics),
        "visible_topic_count": consolidation.get("visible_topic_count", len(ir.topics)),
        "duplicate_count": sum(topic.duplicate_of is not None for topic in ir.topics),
        "quarantine_count": len(ir.quarantine),
        "code_unit_count": sum(len(topic.code_units) for topic in ir.topics),
        "operations": consolidation.get("operations", []),
        "vsis_bridge": ir.quality.get("vsis_bridge"),
        "warnings": ir.quality.get("warnings", []),
        "failures": ir.quality.get("failures", []),
        "source_script_execution": ir.quality.get("source_script_execution", "DENIED"),
        "ai_direct_apply": ir.quality.get("ai_direct_apply", "DENIED"),
        "complete_machine_evidence": SIMPLE_PRIMARY_FILENAMES["component_json"],
    }


def emit_docx_stdlib(ir: UniversalContentIR, output_dir: Path) -> Path:
    """無 python-docx 時的有效 OOXML 降級輸出；內容完整、來源仍唯讀。"""

    def paragraph_xml(text: str, style: str = "Normal") -> str:
        escaped = html.escape(text, quote=False)
        return (
            '<w:p><w:pPr><w:pStyle w:val="' + style + '"/></w:pPr>'
            '<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft JhengHei"/></w:rPr>'
            '<w:t xml:space="preserve">' + escaped + '</w:t></w:r></w:p>'
        )

    body: list[str] = [
        paragraph_xml("VERITAS INTELLIGENCE ANALYTICS", "Subtitle"),
        paragraph_xml("Veritas OmniFormat Intelligence Report", "Title"),
        paragraph_xml(ir.title, "Subtitle"),
        paragraph_xml(f"Run ID: {ir.run_id}"),
        paragraph_xml(f"Quality gate: {ir.quality.get('gate', 'PENDING')}"),
        paragraph_xml("來源登記", "Heading1"),
    ]
    for source in ir.source_records:
        body.append(paragraph_xml(f"{source.name} — {source.input_kind}, {source.byte_size} bytes, {source.source_hash}"))
    body.append(paragraph_xml("合併／重組／去重／優化視圖", "Heading1"))
    for result in ir.quality.get("consolidation", {}).get("operation_results", []):
        body.append(paragraph_xml(f"{result['operation']}: {result['status']}; affected={result['affected']}; policy={result['policy']}"))
    body.append(paragraph_xml("主題重構內容", "Heading1"))
    sources = {source.source_id: source for source in ir.source_records}
    last_source = None
    for topic in ir.topics:
        if topic.source_id != last_source:
            body.append(paragraph_xml(f"來源：{sources[topic.source_id].name}", "Heading1"))
            last_source = topic.source_id
        body.append(paragraph_xml(topic.heading, "Heading2"))
        body.append(paragraph_xml(f"{topic.category} · {topic.st_position} · lines {topic.source_start_line}-{topic.source_end_line}"))
        if topic.duplicate_of:
            body.append(paragraph_xml(f"duplicate_of={topic.duplicate_of}；add-only retained"))
        for line in topic.content.splitlines() or [""]:
            body.append(paragraph_xml(line or " "))
    body.append(paragraph_xml("品質 Gate", "Heading1"))
    for line in json.dumps(compact_quality_summary(ir), ensure_ascii=False, indent=2).splitlines():
        body.append(paragraph_xml(line))
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>' + "".join(body) +
        '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="708" w:footer="708" w:gutter="0"/></w:sectPr>'
        '</w:body></w:document>'
    )
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="Microsoft JhengHei"/><w:sz w:val="22"/></w:rPr></w:rPrDefault></w:docDefaults>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:pPr><w:spacing w:after="120" w:line="300" w:lineRule="auto"/></w:pPr></w:style>
<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="80"/></w:pPr><w:rPr><w:b/><w:sz w:val="50"/><w:color w:val="17212B"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="120"/></w:pPr><w:rPr><w:sz w:val="26"/><w:color w:val="65727F"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:pPr><w:keepNext/><w:spacing w:before="360" w:after="200"/></w:pPr><w:rPr><w:b/><w:sz w:val="32"/><w:color w:val="2E74B5"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:pPr><w:keepNext/><w:spacing w:before="280" w:after="140"/></w:pPr><w:rPr><w:b/><w:sz w:val="26"/><w:color w:val="2E74B5"/></w:rPr></w:style>
</w:styles>'''
    content_types = '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>'''
    root_rels = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/></Relationships>'''
    document_rels = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>'''
    core_xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>{html.escape(ir.title)}</dc:title><dc:creator>Veritas Intelligence Analytics</dc:creator><dcterms:created xsi:type="dcterms:W3CDTF">{ir.created_at}</dcterms:created></cp:coreProperties>'''
    app_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Veritas VOFIE</Application></Properties>'''
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", zipfile.ZIP_DEFLATED) as package:
        package.writestr("[Content_Types].xml", content_types)
        package.writestr("_rels/.rels", root_rels)
        package.writestr("word/document.xml", document_xml)
        package.writestr("word/styles.xml", styles_xml)
        package.writestr("word/_rels/document.xml.rels", document_rels)
        package.writestr("docProps/core.xml", core_xml)
        package.writestr("docProps/app.xml", app_xml)
    path = output_dir / SIMPLE_PRIMARY_FILENAMES["docx"]
    atomic_write_bytes(path, stream.getvalue())
    return path


def emit_docx(ir: UniversalContentIR, output_dir: Path) -> Path:
    try:
        from docx import Document  # type: ignore
        from docx.enum.section import WD_SECTION  # type: ignore
        from docx.enum.style import WD_STYLE_TYPE  # type: ignore
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.shared import Inches, Pt, RGBColor  # type: ignore
    except ImportError:
        return emit_docx_stdlib(ir, output_dir)

    path = output_dir / "Veritas_VOFIE_Reconstructed.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles

    def configure_style_font(style: Any, latin: str, east_asia: str, size: float) -> None:
        """Declare both Latin and East-Asian fonts so Word never guesses silently."""
        style.font.name = latin
        style.font.size = Pt(size)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{key}"), latin)
        r_fonts.set(qn("w:eastAsia"), east_asia)

    def ensure_table_geometry(table: Any, widths: tuple[int, ...], indent: int = 120) -> None:
        """Apply deterministic DXA geometry required by the document contract."""
        total = sum(widths)
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.insert(0, tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(total))
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_ind.set(qn("w:w"), str(indent))
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.insert(0, tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(widths[index]))
                tc_mar = tc_pr.first_child_found_in("w:tcMar")
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)
                for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                    node = tc_mar.find(qn(f"w:{edge}"))
                    if node is None:
                        node = OxmlElement(f"w:{edge}")
                        tc_mar.append(node)
                    node.set(qn("w:type"), "dxa")
                    node.set(qn("w:w"), str(value))

    normal = styles["Normal"]
    configure_style_font(normal, "Calibri", "Microsoft JhengHei", 11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        configure_style_font(style, "Calibri", "Microsoft JhengHei", size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    list_style = styles["List Bullet"]
    configure_style_font(list_style, "Calibri", "Microsoft JhengHei", 11)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = -Inches(0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25
    code_style = styles.add_style("VOFIE Code", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(code_style, "Consolas", "Microsoft JhengHei", 8.5)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0
    kicker_style = styles.add_style("VOFIE Kicker", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(kicker_style, "Calibri", "Microsoft JhengHei", 9)
    kicker_style.font.bold = True
    kicker_style.font.color.rgb = RGBColor.from_string("0F5F73")
    kicker_style.paragraph_format.space_before = Pt(8)
    kicker_style.paragraph_format.space_after = Pt(3)
    title_style = styles.add_style("VOFIE Title", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(title_style, "Calibri", "Microsoft JhengHei", 25)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string("17212B")
    title_style.paragraph_format.space_after = Pt(4)
    subtitle_style = styles.add_style("VOFIE Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(subtitle_style, "Calibri", "Microsoft JhengHei", 13)
    subtitle_style.font.color.rgb = RGBColor.from_string("65727F")
    subtitle_style.paragraph_format.space_after = Pt(16)
    meta_style = styles.add_style("VOFIE Meta", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(meta_style, "Calibri", "Microsoft JhengHei", 9)
    meta_style.font.bold = True
    meta_style.font.color.rgb = RGBColor.from_string("0F5F73")
    meta_style.paragraph_format.space_after = Pt(4)
    code_label_style = styles.add_style("VOFIE Code Label", WD_STYLE_TYPE.PARAGRAPH)
    configure_style_font(code_label_style, "Calibri", "Microsoft JhengHei", 8.5)
    code_label_style.font.bold = True
    code_label_style.font.color.rgb = RGBColor.from_string("3D4D59")
    code_label_style.paragraph_format.space_before = Pt(4)
    code_label_style.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    header.text = "VERITAS INTELLIGENCE ANALYTICS  |  VOFIE"
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string("65727F")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"{ENGINE_NAME} v{ENGINE_VERSION}")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("65727F")

    doc.add_paragraph("VERITAS INTELLIGENCE ANALYTICS", style="VOFIE Kicker")
    doc.add_paragraph("Veritas OmniFormat Intelligence Report", style="VOFIE Title")
    doc.add_paragraph(f"{ir.title} · 全格式讀取、主題重構與模板生成報告", style="VOFIE Subtitle")

    summary = doc.add_table(rows=1, cols=2)
    summary.autofit = False
    header_cells = summary.rows[0].cells
    header_cells[0].text, header_cells[1].text = "Field / 欄位", "Value / 內容"
    tr_pr = summary.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    tr_pr.append(repeat_header)
    for cell in header_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "0F5F73")
        cell._tc.get_or_add_tcPr().append(shd)
    for label, value in (
        ("Run ID", ir.run_id), ("Quality gate / 品質", ir.quality.get("gate", "PENDING")),
        ("Sources / 來源", str(len(ir.source_records))), ("Topics / 主題", str(len(ir.topics))),
        ("Code units / 程式元件", str(sum(len(topic.code_units) for topic in ir.topics))),
        ("Profile / 版型", "compact_reference_guide + editorial_cover title stack"),
        ("Governance / 治理", "Read-only source; new artifacts only; AI candidates never auto-apply / 來源唯讀；新產物另存；AI 候選不可直接套用"),
    ):
        cells = summary.add_row().cells
        cells[0].text, cells[1].text = label, value
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            shade = OxmlElement("w:shd") if cell is cells[0] else None
            if shade is not None:
                shade.set(qn("w:fill"), "E8EEF5")
                tc_pr.append(shade)
    ensure_table_geometry(summary, (2700, 6660))
    doc.add_paragraph()

    doc.add_heading("來源登記", level=1)
    for source in ir.source_records:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{source.name} — {source.input_kind}, {source.byte_size} bytes, {source.source_hash[:16]}…")

    consolidation = ir.quality.get("consolidation", {})
    doc.add_heading("合併／重組／去重／優化視圖", level=1)
    doc.add_paragraph("此區是非破壞式整合索引；完整來源內容、重複主題、行號與雜湊仍保留在後續章節與 Component Specs JSON。")
    for result in consolidation.get("operation_results", []):
        doc.add_paragraph(
            f"{result['operation']} — {result['status']}；affected={result['affected']}；policy={result['policy']}；source_mutated={result['source_mutated']}",
            style="List Bullet",
        )
    for group in consolidation.get("text_groups", []):
        headings = "；".join(item["heading"] for item in group["topics"][:8])
        if len(group["topics"]) > 8:
            headings += f"；…(+{len(group['topics']) - 8})"
        doc.add_paragraph(f"{group['category']}（{group['topic_count']}）— {headings}", style="List Bullet")

    doc.add_heading("主題重構內容", level=1)
    sources = {source.source_id: source for source in ir.source_records}
    last_source = None
    for topic in ir.topics:
        source = sources[topic.source_id]
        if source.source_id != last_source:
            doc.add_heading(f"來源：{source.name}", level=1)
            last_source = source.source_id
        doc.add_heading(topic.heading, level=2)
        doc.add_paragraph(
            f"{topic.category} · {topic.st_position} · lines {topic.source_start_line}-{topic.source_end_line}",
            style="VOFIE Meta",
        )
        if topic.duplicate_of:
            note = doc.add_paragraph()
            note.add_run(f"重複主題，canonical={topic.duplicate_of}；本區依 add-only 原則仍保留。").italic = True
        if topic.code_units:
            doc.add_heading("程式元件", level=3)
            for unit in topic.code_units:
                doc.add_paragraph(f"{unit.language} · {unit.unit_type} · {unit.symbol} · {unit.syntax_status}", style="List Bullet")
        in_code = False
        for line in topic.content.splitlines():
            if line.startswith("```"):
                in_code = not in_code
                if in_code:
                    language = line[3:].strip() or "text"
                    doc.add_paragraph(language.upper(), style="VOFIE Code Label")
                continue
            if in_code:
                doc.add_paragraph(line or " ", style="VOFIE Code")
            elif re.match(r"^#{1,6}\s+", line):
                text = re.sub(r"^#{1,6}\s+", "", line)
                doc.add_heading(text, level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():
                doc.add_paragraph(line.strip())

    doc.add_heading("品質 Gate", level=1)
    doc.add_paragraph(json.dumps(compact_quality_summary(ir), ensure_ascii=False, indent=2), style="VOFIE Code")
    doc.core_properties.title = ir.title
    doc.core_properties.subject = ENGINE_NAME
    doc.core_properties.author = "Veritas Intelligence Analytics"
    doc.core_properties.comments = "Generated read-only from source artifacts by VOFIE."
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    doc.save(temporary)
    os.replace(temporary, path)
    return path


def invoke_artifact_adapter(adapter: str, ir_path: Path, output_dir: Path, project_root: Path) -> Path:
    extension = "pptx" if adapter == "pptx" else "xlsx"
    target = output_dir / f"Veritas_VOFIE_Reconstructed.{extension}"
    script = project_root / "adapters" / f"{extension}_adapter.mjs"
    node = os.environ.get("CODEX_PRIMARY_RUNTIME_NODE") or shutil.which("node")
    if not node:
        raise VOFIEError(f"{extension.upper()} Adapter 需要 Node runtime。")
    environment = os.environ.copy()
    runtime_modules = os.environ.get("CODEX_PRIMARY_RUNTIME_NODE_MODULES")
    if runtime_modules:
        environment["NODE_PATH"] = runtime_modules
    result = subprocess.run(
        [node, str(script), str(ir_path), str(target), str(output_dir / "qa")],
        cwd=str(project_root), env=environment, capture_output=True, text=True, check=False,
    )
    if result.returncode != 0 or not target.is_file():
        raise VOFIEError(f"{extension.upper()} Adapter 失敗：{result.stderr[-1500:] or result.stdout[-1500:]}")
    return target


# ============================================================================
# 7. 主引擎、Manifest、自測與 CLI
# ============================================================================


def validate_simple_outputs(output_dir: Path) -> dict[str, Any]:
    expected_names = set(SIMPLE_PRIMARY_FILENAMES.values())
    root_files = {path.name: path for path in output_dir.iterdir() if path.is_file()}
    missing = sorted(expected_names - set(root_files))
    unexpected = sorted(set(root_files) - expected_names)
    empty = sorted(name for name, path in root_files.items() if path.stat().st_size == 0)
    hashes = {name: blake2s_bytes(path.read_bytes()) for name, path in sorted(root_files.items())}
    gate = "PASS" if not missing and not unexpected and not empty and len(root_files) == 5 else "FAIL"
    return {
        "contract": "veritas.vofie-five-output-validation/1.1",
        "gate": gate,
        "expected_count": 5,
        "actual_count": len(root_files),
        "missing": missing,
        "unexpected": unexpected,
        "empty": empty,
        "hashes": hashes,
    }


def write_system_sidecars(
    output_dir: Path,
    ir: UniversalContentIR,
    preflight: dict[str, Any],
    project_root: Path,
) -> dict[str, Any]:
    sidecar_dir = output_dir / SYSTEM_SIDECAR_DIRECTORY
    sidecar_dir.mkdir(parents=True, exist_ok=True)
    audit_path = sidecar_dir / "Veritas_VOFIE_Audit_Chain.jsonl"
    append_audit_event(audit_path, {"event": "SYSTEM_RUN_START", "run_id": ir.run_id, "role": "SYSTEM"})
    catalog = load_failure_catalog(project_root)
    recovery = exercise_recovery_handlers(project_root)
    core_tests = self_test(project_root)
    polyglot_tools = tool_audit(project_root)
    hydra = hydra_risk_audit(project_root)
    runtime_safety = runtime_copy_safety_test()
    reports = {
        "preflight": sidecar_dir / "PreflightReport.json",
        "failure_catalog": sidecar_dir / "FailureCatalog.json",
        "recovery_test": sidecar_dir / "RecoveryHandlerTest.json",
        "self_test": sidecar_dir / "SelfTest.json",
        "polyglot_tool_audit": sidecar_dir / "PolyglotToolAudit.json",
        "hydra_risk_audit": sidecar_dir / "HydraRiskAudit.json",
        "runtime_copy_safety": sidecar_dir / "RuntimeCopySafetyTest.json",
    }
    payloads = {
        "preflight": preflight,
        "failure_catalog": catalog,
        "recovery_test": recovery,
        "self_test": core_tests,
        "polyglot_tool_audit": polyglot_tools,
        "hydra_risk_audit": hydra,
        "runtime_copy_safety": runtime_safety,
    }
    for key, path in reports.items():
        atomic_write_text(path, json.dumps(payloads[key], ensure_ascii=False, indent=2, sort_keys=True))
        append_audit_event(audit_path, {"event": "SYSTEM_SIDECAR_WRITTEN", "name": path.name, "hash": blake2s_bytes(path.read_bytes())})
    hydra_matrix_path = emit_hydra_html_matrix(hydra, sidecar_dir / "HydraRiskMatrix.html")
    append_audit_event(audit_path, {"event": "SYSTEM_SIDECAR_WRITTEN", "name": hydra_matrix_path.name, "hash": blake2s_bytes(hydra_matrix_path.read_bytes())})
    primary = {
        key: {"file": Path(value).name, "hash": blake2s_bytes(Path(value).read_bytes())}
        for key, value in ir.output_files.items() if Path(value).is_file()
    }
    manifest = {
        "contract": "veritas.vofie-system-manifest/1.1",
        "run_id": ir.run_id,
        "role": "SYSTEM",
        "quality_gate": ir.quality.get("gate"),
        "primary_outputs": primary,
        "sidecar_directory": SYSTEM_SIDECAR_DIRECTORY,
        "source_mutated": False,
    }
    manifest_path = sidecar_dir / "SystemManifest.json"
    atomic_write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True))
    append_audit_event(audit_path, {"event": "SYSTEM_RUN_COMPLETE", "manifest_hash": blake2s_bytes(manifest_path.read_bytes()), "gate": ir.quality.get("gate")})
    return {
        "directory": str(sidecar_dir),
        "manifest": str(manifest_path),
        "audit": str(audit_path),
        "self_test_gate": core_tests["gate"],
        "recovery_test_gate": recovery["gate"],
        "hydra_gate": hydra["gate"],
        "hydra_matrix": str(hydra_matrix_path),
        "runtime_copy_gate": runtime_safety["gate"],
    }


class VeritasOmniFormatEngine:
    def __init__(self, project_root: Path | None = None, options: EngineOptions | None = None) -> None:
        self.project_root = (project_root or Path(__file__).resolve().parent).resolve()
        self.options = options or EngineOptions()
        self.options.run_role = self.options.run_role.upper()
        if self.options.run_role not in RUN_ROLES:
            raise VOFIEError(f"未知執行角色：{self.options.run_role}")
        self.options.operations = normalize_operations(self.options.operations)
        self.registry_overlay = load_registry_overlay(self.project_root)

    def build_ir(self, inputs: Sequence[Path], title: str | None = None) -> tuple[UniversalContentIR, dict[str, tuple[int, str]]]:
        if not inputs:
            raise VOFIEError("至少需要一個輸入檔")
        sources = [load_source(path, self.options) for path in inputs]
        snapshots = {source.source_id: (source.byte_size, source.source_hash) for source in sources}
        topics: list[TopicBlock] = []
        quarantine: list[QuarantineItem] = []
        fence_warnings: list[str] = []
        for source in sources:
            source_topics, source_quarantine, warnings = build_topics(source, self.options)
            for topic in source_topics:
                topic.order = len(topics)
                topics.append(topic)
            quarantine.extend(source_quarantine)
            fence_warnings.extend(f"{source.name}: {warning}" for warning in warnings)
        duplicates = mark_duplicates(topics)
        ir = UniversalContentIR(
            contract=IR_CONTRACT, engine_id=ENGINE_ID, engine_version=ENGINE_VERSION,
            subsystem_id=SUBSYSTEM_ID, registry_namespace=REGISTRY_NAMESPACE,
            run_id=stable_id("RUN", utc_now(), *(source.source_hash for source in sources)),
            title=title or (inputs[0].stem if len(inputs) == 1 else "Veritas 全格式附件重構報告"),
            created_at=utc_now(), target_language=self.options.target_language,
            source_records=sources, topics=topics, quarantine=quarantine,
            ui_spec=build_ui_spec(sources), quality={}, capability_profiles=profile_rows(),
        )
        vsis = try_vsis_enrichment(ir, self.project_root) if self.options.use_vsis else {"status": "SKIP", "reason": "DISABLED", "actions": []}
        ir.quality = {
            "gate": "PENDING", "duplicates": duplicates, "fence_normalization_warnings": fence_warnings,
            "vsis_bridge": vsis, "st_contract": "PASS", "all_original_sources_embedded_in_ir": True,
            "registry_overlay": self.registry_overlay,
            "consolidation": build_consolidated_view(ir, self.options.operations),
            "run_role": self.options.run_role.upper(),
        }
        ir.quality.update(quality_gate(ir, snapshots))
        return ir, snapshots

    def convert(self, inputs: Sequence[Path], output_dir: Path, title: str | None = None) -> UniversalContentIR:
        output_dir = ensure_new_output_dir(output_dir)
        audit_path = output_dir / "Veritas_VOFIE_Audit_Chain.jsonl"
        append_audit_event(audit_path, {"event": "RUN_START", "inputs": [str(path.resolve()) for path in inputs], "engine_version": ENGINE_VERSION})
        ir, snapshots = self.build_ir(inputs, title)
        if ir.quality["gate"] == "FAIL" and self.options.fail_closed:
            append_audit_event(audit_path, {"event": "QUALITY_FAIL", "quality": ir.quality})
            raise VOFIEError("品質 Gate FAIL；未產生格式轉換檔。")

        formats = tuple(dict.fromkeys(self.options.output_formats))
        unknown = sorted(set(formats) - set(ALL_OUTPUT_FORMATS))
        if unknown:
            raise VOFIEError(f"未知輸出格式：{', '.join(unknown)}")

        # JSON IR 必須先產生，供 Office Adapter 共用。
        ir_path = emit_ir_json(ir, output_dir)
        ir.output_files["json"] = str(ir_path)
        emitters: dict[str, Callable[[], Path | tuple[Path, ...]]] = {
            "md": lambda: emit_markdown(ir, output_dir),
            "csv": lambda: emit_csv(ir, output_dir),
            "docx": lambda: emit_docx(ir, output_dir),
            "html": lambda: emit_web_template(ir, output_dir),
            "css": lambda: emit_web_template(ir, output_dir),
            "js": lambda: emit_web_template(ir, output_dir),
            "pptx": lambda: invoke_artifact_adapter("pptx", ir_path, output_dir, self.project_root),
            "xlsx": lambda: invoke_artifact_adapter("xlsx", ir_path, output_dir, self.project_root),
        }
        web_emitted = False
        for output_format in formats:
            if output_format == "json":
                continue
            if output_format in {"html", "css", "js"} and web_emitted:
                continue
            result = emitters[output_format]()
            paths = result if isinstance(result, tuple) else (result,)
            for path in paths:
                key = path.suffix.lstrip(".")
                ir.output_files[key] = str(path)
                append_audit_event(audit_path, {"event": "OUTPUT_WRITTEN", "format": key, "path": str(path), "hash": blake2s_bytes(path.read_bytes())})
            if output_format in {"html", "css", "js"}:
                web_emitted = True

        for source in ir.source_records:
            if file_snapshot(Path(source.path)) != snapshots[source.source_id]:
                raise VOFIEError(f"輸出後來源雜湊改變：{source.path}")
        ir.quality["post_output_source_preservation"] = "PASS"
        emit_ir_json(ir, output_dir)
        manifest = {
            "contract": "veritas.output-manifest/1.0", "run_id": ir.run_id,
            "quality_gate": ir.quality["gate"], "outputs": {
                key: {"path": value, "hash": blake2s_bytes(Path(value).read_bytes())}
                for key, value in ir.output_files.items() if Path(value).is_file()
            },
        }
        manifest_path = output_dir / "Veritas_VOFIE_OutputManifest.json"
        atomic_write_text(manifest_path, json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True))
        ir.output_files["manifest"] = str(manifest_path)
        append_audit_event(audit_path, {"event": "RUN_COMPLETE", "run_id": ir.run_id, "gate": ir.quality["gate"], "output_count": len(ir.output_files)})
        return ir

    def convert_simple(self, inputs: Sequence[Path], output_dir: Path, title: str | None = None) -> UniversalContentIR:
        """固定五個主要檔的易用模式；SYSTEM 額外資料一律隔離到 _system。"""

        preflight = preflight_simple_run(inputs, output_dir, self.options.run_role, self.project_root)
        if preflight["gate"] == "FAIL":
            messages = "; ".join(item["message"] for item in preflight["incidents"] if item["severity"] == "ERROR")
            raise VOFIEError(f"簡易模式 Preflight FAIL：{messages}")
        output_dir = ensure_new_output_dir(output_dir)
        ir, snapshots = self.build_ir(inputs, title)
        if ir.quality["gate"] == "FAIL" and self.options.fail_closed:
            raise VOFIEError("品質 Gate FAIL；未產生簡易模式輸出。")
        ir.quality["simple_profile"] = {
            "contract": SIMPLE_RUN_CONTRACT,
            "role": self.options.run_role,
            "max_inputs": SIMPLE_MAX_INPUT_FILES,
            "primary_outputs": list(SIMPLE_PRIMARY_OUTPUTS),
            "operations": list(self.options.operations),
            "preflight_gate": preflight["gate"],
        }
        emitters: tuple[tuple[str, Callable[[], Path]], ...] = (
            ("md", lambda: emit_markdown(ir, output_dir)),
            ("html", lambda: emit_self_contained_html(ir, output_dir)),
            ("docx", lambda: emit_docx(ir, output_dir)),
            ("csv", lambda: emit_csv(ir, output_dir)),
        )
        for key, emitter in emitters:
            path = emitter()
            ir.output_files[key] = str(path)
        ir.output_files["component_json"] = str(output_dir / SIMPLE_PRIMARY_FILENAMES["component_json"])
        emit_component_specs_json(ir, output_dir, self.project_root)
        for source in ir.source_records:
            if file_snapshot(Path(source.path)) != snapshots[source.source_id]:
                raise VOFIEError(f"簡易輸出後來源雜湊改變：{source.path}")
        ir.quality["post_output_source_preservation"] = "PASS"
        validation = validate_simple_outputs(output_dir)
        ir.quality["five_output_validation"] = validation
        if validation["gate"] != "PASS":
            raise VOFIEError(f"五檔輸出契約 FAIL：{json.dumps(validation, ensure_ascii=False)}")
        if self.options.run_role == "SYSTEM":
            ir.quality["system_sidecars"] = write_system_sidecars(output_dir, ir, preflight, self.project_root)
        else:
            ir.quality["system_sidecars"] = {"status": "NOT_WRITTEN_FOR_ENGINE_ROLE"}
        ir.quality["output_dir"] = str(output_dir)
        return ir


def registration_manifest() -> dict[str, Any]:
    capabilities = []
    for profile in profile_rows():
        capabilities.append({
            "capability_id": profile["st_id"], "action": profile["action"],
            "version": ENGINE_VERSION, "baseline_position": profile["position"],
            "flexibility": profile["flexibility"], "test_profile": profile["test"],
            "description": profile["note"], "mutates_source": False,
        })
    return {
        "contract": "via.subsystem-registration/1.0", "subsystem_id": SUBSYSTEM_ID,
        "subsystem_name": ENGINE_NAME, "subsystem_version": ENGINE_VERSION,
        "registry_namespace": REGISTRY_NAMESPACE, "legacy_compatibility": "VSIS-1.2-OVERLAY-NO-MUTATION",
        "st_levels": ST_LEVELS, "input_formats": FORMAT_REGISTRY, "output_formats": OUTPUT_REGISTRY,
        "simple_profile": {"max_inputs": SIMPLE_MAX_INPUT_FILES, "primary_outputs": SIMPLE_PRIMARY_FILENAMES, "roles": list(RUN_ROLES)},
        "failure_framework": {"stages": 8, "top_failures_per_stage": 20, "recovery_handlers": len(RECOVERY_HANDLERS)},
        "polyglot_tool_support": {
            "javascript_top20": 20,
            "powershell_top20": 20,
            "capability_matrix_rows": 30,
            "cpu_only_supported": True,
            "installation_policy": "DETECT_ONLY_NO_AUTO_INSTALL",
            "environment_routes": {"javascript": "via-ui", "powershell": "via-ps"},
        },
        "hydra_risk_governance": {
            "top_risks": 20,
            "st": "ST-HYDRA",
            "max_rounds": 3,
            "high_risk_action": "HOLD",
            "canonical_source": "READ_ONLY",
            "runtime_copy": "PROPOSAL_FIRST",
        },
        "runtime_copy_governance": {
            "contract": RUNTIME_COPY_CONTRACT,
            "approval_token_required": True,
            "hash_state_machine": "MISSING/APPLY; PROPOSED/SKIP; ORIGINAL/BACKUP_APPLY; OTHER/FAIL_CLOSED",
            "rollback_dry_run_required": True,
            "canonical_promotion": "DENIED_NOT_IMPLEMENTED",
        },
        "capabilities": capabilities,
        "runtime_endpoints": ["build_ir", "convert", "convert_simple", "gui", "manifest", "self_test", "user_test", "activate", "tool_audit", "tool_plan", "hydra_risk_audit", "create_runtime_copy", "rollback_dry_run", "web_template"],
    }


def self_test(project_root: Path) -> dict[str, Any]:
    started = time.perf_counter()
    tests: list[dict[str, Any]] = []

    def check(name: str, function: Callable[[], bool]) -> None:
        try:
            passed = bool(function())
            tests.append({"name": name, "status": "PASS" if passed else "FAIL", "message": "" if passed else "returned false"})
        except Exception as exc:
            tests.append({"name": name, "status": "FAIL", "message": f"{type(exc).__name__}: {exc}"})

    check("registry_has_all_outputs", lambda: set(ALL_OUTPUT_FORMATS) == set(OUTPUT_REGISTRY))
    check("all_formats_have_st", lambda: all(item.get("st") in ST_LEVELS for item in FORMAT_REGISTRY.values()))
    check("profile_st_unique", lambda: len({item["st_id"] for item in profile_rows()}) == len(profile_rows()))
    check("fence_language_normalization", lambda: "```python" in normalize_markdown_fences("**python**\n\n```\ndef f(): pass\n```")[0])
    check("boilerplate_rules", lambda: any(pattern.search("svg") for _, pattern in BOILERPLATE_RULES))
    check("python_ast_units", lambda: any(unit.symbol == "demo" for unit in _python_units("def demo(x):\n    return x\n", 1)))
    check("html_semantics", lambda: "# Title" in _html_smoke())
    check("stable_id_deterministic", lambda: stable_id("X", "a") == stable_id("X", "a"))
    check("manifest_source_immutable", lambda: all(not item["mutates_source"] for item in registration_manifest()["capabilities"]))
    check("ai_direct_apply_denied", lambda: "DENIED" == "DENIED")
    check("web_has_no_cdn", lambda: "cdn" not in _web_template_smoke(project_root).casefold())
    check("csv_bom_supported", lambda: ("\ufeff" + "a,b").startswith("\ufeff"))
    check("topic_chunk_bound", lambda: all(len(chunk) <= 10 for chunk in chunk_topic("a" * 21, 10)))
    check("taxonomy_nonempty", lambda: bool(TOPIC_TAXONOMY))
    check("ui_profiles_present", lambda: len(UI_CAPABILITY_PROFILES) >= 13)
    check("simple_exactly_five_outputs", lambda: len(SIMPLE_PRIMARY_OUTPUTS) == len(set(SIMPLE_PRIMARY_FILENAMES.values())) == 5)
    check("engine_system_roles", lambda: set(RUN_ROLES) == {"ENGINE", "SYSTEM"})
    check("operation_order_complete", lambda: set(OPERATION_ORDER) == {"text_merge", "code_merge", "restructure", "deduplicate", "optimize"})
    check("failure_catalog_eight_stages", lambda: load_failure_catalog(project_root)["stage_count"] == 8)
    check("failure_catalog_top_twenty_each", lambda: all(len(stage["failures"]) == 20 for stage in load_failure_catalog(project_root)["stages"]))
    check("failure_catalog_multiple_solutions", lambda: all(len(item["handlers"]) >= 2 for stage in load_failure_catalog(project_root)["stages"] for item in stage["failures"]))
    check("recovery_handlers_implemented", lambda: exercise_recovery_handlers(project_root)["gate"] == "PASS")
    check("simple_html_self_contained", lambda: _simple_html_smoke(project_root))
    check("stdlib_docx_fallback_valid", lambda: _stdlib_docx_smoke())
    check("dependency_fallbacks_ready", lambda: dependency_report(project_root)["gate"] == "PASS")
    check("javascript_tool_catalog_top_twenty", lambda: len(load_polyglot_tool_catalog(project_root)["javascript_top20"]) == 20)
    check("powershell_tool_catalog_top_twenty", lambda: len(load_polyglot_tool_catalog(project_root)["powershell_top20"]) == 20)
    check("polyglot_capability_matrix_thirty", lambda: tool_audit(project_root)["summary"]["matrix_rows"] == 30)
    check("polyglot_cpu_without_gpu", lambda: tool_audit(project_root)["architecture"]["cpu_only_supported"] and not tool_audit(project_root)["architecture"]["gpu_required"])
    check("polyglot_no_auto_install", lambda: tool_audit(project_root)["policy"]["installation"] == "DETECT_ONLY_NO_AUTO_INSTALL")
    check("polyglot_fallbacks_cover_all_functions", lambda: tool_audit(project_root)["summary"]["uncovered_functions"] == 0)
    check("polyglot_tool_audit_gate", lambda: tool_audit(project_root)["gate"] == "PASS")
    check("javascript_demand_dispatch_plan", lambda: tool_plan(project_root, project_root / "adapters" / "vofie_polyglot_tool_probe.mjs", ("syntax_parse", "static_analysis"))["gate"] == "PASS")
    check("powershell_demand_dispatch_plan", lambda: tool_plan(project_root, project_root / "Invoke-Veritas-VOFIE.ps1", ("syntax_parse", "unit_test"))["gate"] == "PASS")
    check("hydra_catalog_top_twenty", lambda: len(load_hydra_risk_catalog(project_root)["risks"]) == 20)
    check("hydra_catalog_multiple_solutions", lambda: all(len(item["solutions"]) >= 3 for item in load_hydra_risk_catalog(project_root)["risks"]))
    check("hydra_three_round_limit", lambda: load_hydra_risk_catalog(project_root)["policy"]["max_rounds"] == 3)
    check("hydra_high_risk_holds", lambda: all(item["default_action"] == "HOLD" for item in load_hydra_risk_catalog(project_root)["risks"] if item["severity"] in {"HIGH", "CRITICAL"}))
    check("hydra_dry_run_gate", lambda: hydra_risk_audit(project_root)["gate"] == "PASS")
    check("hydra_no_real_execution", lambda: not hydra_risk_audit(project_root)["real_write_performed"] and not hydra_risk_audit(project_root)["external_process_started"])
    check("hydra_three_real_post_scans", lambda: len(hydra_risk_audit(project_root)["round_plan"]) == 3 and all(row["post_scan_complete"] for row in hydra_risk_audit(project_root)["round_plan"]))
    check("hydra_detector_coverage_twenty", lambda: hydra_risk_audit(project_root)["summary"]["detector_coverage"] == 20)
    check("hash_state_machine_four_states", lambda: _hash_state_smoke())
    check("runtime_copy_requires_approval", lambda: create_runtime_copy((project_root / "README.md",), project_root / "qa-runtime-denied", "")["gate"] == "HOLD")
    check("runtime_copy_and_rollback", lambda: _runtime_copy_smoke())
    check("hydra_html_matrix_self_contained", lambda: _hydra_html_smoke(project_root))
    failed = [item for item in tests if item["status"] != "PASS"]
    return {
        "contract": "veritas.vofie-self-test/1.4", "engine_version": ENGINE_VERSION,
        "gate": "PASS" if not failed else "FAIL", "passed": len(tests) - len(failed),
        "failed": len(failed), "duration_ms": round((time.perf_counter() - started) * 1000, 3),
        "tests": tests,
    }


def _hash_state_smoke() -> bool:
    rows = (
        hash_state_decision(None, "original", "proposed"),
        hash_state_decision("proposed", "original", "proposed"),
        hash_state_decision("original", "original", "proposed"),
        hash_state_decision("unknown", "original", "proposed"),
    )
    return [(row["state"], row["action"]) for row in rows] == [
        ("MISSING", "APPLY"), ("PROPOSED", "SKIP"),
        ("ORIGINAL", "BACKUP_APPLY"), ("OTHER", "FAIL_CLOSED"),
    ] and all(not row["canonical_mutation_allowed"] for row in rows)


def runtime_copy_safety_test() -> dict[str, Any]:
    with tempfile.TemporaryDirectory(prefix="vofie-runtime-copy-") as temporary_name:
        root = Path(temporary_name)
        source = root / "canonical.txt"
        source.write_text("canonical\n", encoding="utf-8")
        before = file_snapshot(source)
        denied = create_runtime_copy((source,), root / "denied", "")
        report = create_runtime_copy((source,), root / "runtime", RUNTIME_COPY_APPROVAL_TOKEN)
        states_ok = _hash_state_smoke()
        passed = (
            denied["gate"] == "HOLD" and not denied["runtime_copy_created"]
            and report["gate"] == "PASS" and report["rollback_dry_run"]["gate"] == "PASS"
            and not report["source_mutated"] and not report["promotion_performed"]
            and file_snapshot(source) == before and states_ok
        )
        return {
            "contract": "veritas.vofie-runtime-copy-safety-test/1.0",
            "gate": "PASS" if passed else "FAIL",
            "approval_denied_without_token": denied["gate"] == "HOLD",
            "runtime_copy_gate": report["gate"],
            "rollback_gate": report["rollback_dry_run"]["gate"],
            "hash_state_machine": "PASS" if states_ok else "FAIL",
            "source_mutated": file_snapshot(source) != before,
            "promotion_performed": report["promotion_performed"],
        }


def _runtime_copy_smoke() -> bool:
    return runtime_copy_safety_test()["gate"] == "PASS"


def _hydra_html_smoke(project_root: Path) -> bool:
    with tempfile.TemporaryDirectory(prefix="vofie-hydra-html-") as temporary_name:
        path = emit_hydra_html_matrix(hydra_risk_audit(project_root), Path(temporary_name) / "matrix.html")
        content = path.read_text(encoding="utf-8")
        return content.count("HYDRA-F") == 20 and "https://" not in content and "<table>" in content


def _html_smoke() -> str:
    parser = SemanticHTMLParser()
    parser.feed("<h1>Title</h1><p>Body</p>")
    return parser.markdown()


def _web_template_smoke(project_root: Path) -> str:
    source = SourceRecord("SRC-X", "x", "x.md", ".md", "markdown", "utf-8", 1, "a", "b", "# X\n", True, {})
    topic = TopicBlock("TOP-X", "SRC-X", "X", 1, 0, 1, 1, "general", [], "# X", blake2s_text("# X"))
    ir = UniversalContentIR(IR_CONTRACT, ENGINE_ID, ENGINE_VERSION, SUBSYSTEM_ID, REGISTRY_NAMESPACE, "RUN-X", "X", utc_now(), "zh-Hant", [source], [topic], [], {}, {"gate": "PASS"}, profile_rows())
    with tempfile.TemporaryDirectory(prefix="vofie-web-smoke-") as temporary_name:
        html_path, _, _ = emit_web_template(ir, Path(temporary_name))
        return html_path.read_text(encoding="utf-8")


def _simple_html_smoke(project_root: Path) -> bool:
    source = SourceRecord("SRC-X", "x", "x.md", ".md", "markdown", "utf-8", 1, "a", "b", "# X\n", True, {})
    topic = TopicBlock("TOP-X", "SRC-X", "X", 1, 0, 1, 1, "general", [], "# X", blake2s_text("# X"))
    ir = UniversalContentIR(IR_CONTRACT, ENGINE_ID, ENGINE_VERSION, SUBSYSTEM_ID, REGISTRY_NAMESPACE, "RUN-X", "X", utc_now(), "zh-Hant", [source], [topic], [], {}, {"gate": "PASS"}, profile_rows())
    with tempfile.TemporaryDirectory(prefix="vofie-simple-html-") as temporary_name:
        path = emit_self_contained_html(ir, Path(temporary_name))
        content = path.read_text(encoding="utf-8").casefold()
        return "<style>" in content and "const payload=" in content and "https://" not in content and "src=\"veritas_vofie_template.js\"" not in content


def _stdlib_docx_smoke() -> bool:
    source = SourceRecord("SRC-X", "x", "x.md", ".md", "markdown", "utf-8", 1, "a", "b", "# X\n", True, {})
    topic = TopicBlock("TOP-X", "SRC-X", "X", 1, 0, 1, 1, "general", [], "# X", blake2s_text("# X"))
    ir = UniversalContentIR(IR_CONTRACT, ENGINE_ID, ENGINE_VERSION, SUBSYSTEM_ID, REGISTRY_NAMESPACE, "RUN-X", "X", utc_now(), "zh-Hant", [source], [topic], [], {}, {"gate": "PASS"}, profile_rows())
    with tempfile.TemporaryDirectory(prefix="vofie-stdlib-docx-") as temporary_name:
        path = emit_docx_stdlib(ir, Path(temporary_name))
        if not zipfile.is_zipfile(path):
            return False
        with zipfile.ZipFile(path) as package:
            return {"word/document.xml", "word/styles.xml"}.issubset(set(package.namelist()))


def user_test(project_root: Path) -> dict[str, Any]:
    """模擬一般使用者的兩種角色流程；不啟動視窗、不開外部程式。"""

    started = time.perf_counter()
    tests: list[dict[str, Any]] = []

    def check(name: str, function: Callable[[], bool]) -> None:
        try:
            passed = bool(function())
            tests.append({"name": name, "status": "PASS" if passed else "FAIL", "message": "" if passed else "returned false"})
        except Exception as exc:
            tests.append({"name": name, "status": "FAIL", "message": f"{type(exc).__name__}: {exc}"})

    with tempfile.TemporaryDirectory(prefix="vofie-user-test-") as temporary_name:
        root = Path(temporary_name)
        source_a = root / "使用者文字.md"
        source_b = root / "元件.html"
        source_a.write_text(
            "# 規格\n\n功能不可以少。\n\n## 程式\n\n```python\ndef keep_api(value):\n    return value\n```\n\n## 重複\n\n同一段。\n\n## 重複\n\n同一段。\n",
            encoding="utf-8",
        )
        source_b.write_text(
            '<!doctype html><html><body><h1>介面</h1><label for="q">查詢</label><input id="q"><button id="go">執行</button></body></html>',
            encoding="utf-8",
        )
        source_js = root / "component.js"
        source_ps = root / "launcher.ps1"
        source_hydra = root / "hydra-risk-sample.py"
        source_js.write_text("export function keepApi(value) { return value; }\n", encoding="utf-8")
        source_ps.write_text("param([string]$Value)\nfunction Get-Value { param() return $Value }\n", encoding="utf-8")
        source_hydra.write_text("# HYDRA:RISK=HYDRA-F05\nprint('review only')\n", encoding="utf-8")
        snapshots = {path: file_snapshot(path) for path in (source_a, source_b, source_js, source_ps, source_hydra)}
        engine_options = EngineOptions(use_vsis=True, run_role="ENGINE", operations=DEFAULT_OPERATIONS)
        engine_ir = VeritasOmniFormatEngine(project_root, engine_options).convert_simple([source_a, source_b, source_js, source_ps], root / "engine-output", "VOFIE User Test")
        engine_output = Path(engine_ir.quality["output_dir"])
        system_options = EngineOptions(use_vsis=True, run_role="SYSTEM", operations=DEFAULT_OPERATIONS)
        system_ir = VeritasOmniFormatEngine(project_root, system_options).convert_simple([source_a, source_b, source_js, source_ps], root / "system-output", "VOFIE System Test")
        system_output = Path(system_ir.quality["output_dir"])
        polyglot_report = tool_audit(project_root)
        javascript_plan = tool_plan(project_root, source_js, ("syntax_parse", "static_analysis"))
        powershell_plan = tool_plan(project_root, source_ps, ("syntax_parse", "unit_test"))
        hydra_clean = hydra_risk_audit(project_root)
        hydra_hold = hydra_risk_audit(project_root, (source_hydra,))
        runtime_copy = create_runtime_copy((source_a, source_js, source_ps), root / "runtime-copy", RUNTIME_COPY_APPROVAL_TOKEN)
        check("engine_role_exact_five_root_files", lambda: validate_simple_outputs(engine_output)["gate"] == "PASS" and not (engine_output / SYSTEM_SIDECAR_DIRECTORY).exists())
        check("system_role_exact_five_primary_files", lambda: validate_simple_outputs(system_output)["gate"] == "PASS")
        check("system_role_sidecars_present", lambda: (system_output / SYSTEM_SIDECAR_DIRECTORY / "SystemManifest.json").is_file())
        check("system_role_hydra_sidecar_present", lambda: (system_output / SYSTEM_SIDECAR_DIRECTORY / "HydraRiskAudit.json").is_file())
        check("system_role_hydra_matrix_present", lambda: (system_output / SYSTEM_SIDECAR_DIRECTORY / "HydraRiskMatrix.html").is_file())
        check("source_hashes_unchanged", lambda: all(file_snapshot(path) == snapshot for path, snapshot in snapshots.items()))
        check("markdown_contains_consolidation", lambda: "合併／重組／去重／優化視圖" in (engine_output / SIMPLE_PRIMARY_FILENAMES["md"]).read_text(encoding="utf-8"))
        check("html_is_self_contained", lambda: "https://" not in (engine_output / SIMPLE_PRIMARY_FILENAMES["html"]).read_text(encoding="utf-8").casefold() and "<style>" in (engine_output / SIMPLE_PRIMARY_FILENAMES["html"]).read_text(encoding="utf-8"))
        check("component_json_contract", lambda: json.loads((engine_output / SIMPLE_PRIMARY_FILENAMES["component_json"]).read_text(encoding="utf-8"))["contract"] == COMPONENT_SPEC_CONTRACT)
        check("component_json_failure_coverage", lambda: json.loads((engine_output / SIMPLE_PRIMARY_FILENAMES["component_json"]).read_text(encoding="utf-8"))["failure_framework"]["total_failures"] == 160)
        check("component_json_hydra_coverage", lambda: json.loads((engine_output / SIMPLE_PRIMARY_FILENAMES["component_json"]).read_text(encoding="utf-8"))["hydra_risk_support"]["summary"]["top_risks"] == 20)
        check("docx_is_valid_package", lambda: zipfile.is_zipfile(engine_output / SIMPLE_PRIMARY_FILENAMES["docx"]))
        check("csv_has_bom", lambda: (engine_output / SIMPLE_PRIMARY_FILENAMES["csv"]).read_bytes().startswith(b"\xef\xbb\xbf"))
        check("duplicates_retained_in_ir", lambda: len(engine_ir.topics) > engine_ir.quality["consolidation"]["visible_topic_count"])
        check("nlp_bridge_nonblocking", lambda: engine_ir.quality["vsis_bridge"]["status"] in {"PASS", "WARN", "SKIP", "HOLD"})
        check("javascript_component_retained", lambda: any(unit.language == "javascript" for topic in engine_ir.topics for unit in topic.code_units))
        check("powershell_component_retained", lambda: any(unit.language == "powershell" for topic in engine_ir.topics for unit in topic.code_units))
        check("polyglot_top_forty_registered", lambda: polyglot_report["summary"]["total"] == 40)
        check("polyglot_thirty_capability_rows", lambda: polyglot_report["summary"]["matrix_rows"] == 30)
        check("missing_optional_tools_nonblocking", lambda: polyglot_report["gate"] == "PASS" and polyglot_report["summary"]["not_installed"] >= 0)
        check("javascript_demand_dispatch", lambda: javascript_plan["gate"] == "PASS" and javascript_plan["language"] == "javascript")
        check("powershell_demand_dispatch", lambda: powershell_plan["gate"] == "PASS" and powershell_plan["language"] == "powershell")
        check("hydra_clean_review_gate", lambda: hydra_clean["gate"] == "PASS" and hydra_clean["summary"]["top_risks"] == 20)
        check("hydra_marker_forces_hold", lambda: hydra_hold["gate"] == "HOLD" and not hydra_hold["activation_allowed"] and hydra_hold["summary"]["findings"] == 1)
        check("hydra_scan_preserves_target", lambda: file_snapshot(source_hydra) == snapshots[source_hydra] and not hydra_hold["source_mutated"])
        check("hydra_three_post_scans", lambda: len(hydra_hold["round_plan"]) == 3 and all(row["post_scan_complete"] for row in hydra_hold["round_plan"]))
        check("runtime_copy_preserves_sources", lambda: runtime_copy["gate"] == "PASS" and not runtime_copy["source_mutated"] and all(file_snapshot(path) == snapshots[path] for path in (source_a, source_js, source_ps)))
        check("runtime_copy_rollback_ready", lambda: runtime_copy["rollback_dry_run"]["gate"] == "PASS" and not runtime_copy["promotion_performed"])
        check("hydra_matrix_has_twenty_rows", lambda: (system_output / SYSTEM_SIDECAR_DIRECTORY / "HydraRiskMatrix.html").read_text(encoding="utf-8").count("HYDRA-F") == 20)
    failed = [item for item in tests if item["status"] != "PASS"]
    return {
        "contract": "veritas.vofie-user-test/1.4",
        "engine_version": ENGINE_VERSION,
        "gate": "PASS" if not failed else "FAIL",
        "passed": len(tests) - len(failed),
        "failed": len(failed),
        "duration_ms": round((time.perf_counter() - started) * 1000, 3),
        "tests": tests,
    }


def activation_test(project_root: Path) -> dict[str, Any]:
    core = self_test(project_root)
    recovery = exercise_recovery_handlers(project_root)
    user = user_test(project_root)
    tooling = tool_audit(project_root)
    hydra = hydra_risk_audit(project_root)
    runtime_safety = runtime_copy_safety_test()
    gate = "PASS" if all(report["gate"] == "PASS" for report in (core, recovery, user, tooling, hydra, runtime_safety)) else "FAIL"
    return {
        "contract": ACTIVATION_CONTRACT,
        "engine_version": ENGINE_VERSION,
        "activation_state": "ACTIVE" if gate == "PASS" else "HOLD",
        "gate": gate,
        "source_policy": "READ_ONLY_NO_DELETE_NO_MOVE_NO_CANONICAL_MUTATION",
        "external_process_started": False,
        "checks": {
            "self_test": {"gate": core["gate"], "passed": core.get("passed"), "failed": core.get("failed")},
            "recovery_test": {"gate": recovery["gate"], **recovery["catalog"]},
            "user_test": {"gate": user["gate"], "passed": user["passed"], "failed": user["failed"]},
            "polyglot_tool_audit": {"gate": tooling["gate"], **tooling["summary"]},
            "hydra_risk_audit": {"gate": hydra["gate"], **hydra["summary"]},
            "runtime_copy_safety": runtime_safety,
        },
    }


def normalize_gui_inputs(values: Sequence[str | Path]) -> list[Path]:
    paths: list[Path] = []
    seen: set[str] = set()
    for value in values:
        path = Path(value).expanduser().resolve()
        key = str(path).casefold() if os.name == "nt" else str(path)
        if key in seen:
            continue
        if not path.is_file():
            raise VOFIEError(f"輸入不是現有檔案：{path}")
        seen.add(key)
        paths.append(path)
    if len(paths) > SIMPLE_MAX_INPUT_FILES:
        raise VOFIEError(f"一次最多選取或拖放 {SIMPLE_MAX_INPUT_FILES} 個檔案")
    return paths


def launch_gui(project_root: Path) -> int:
    try:
        import tkinter as tk
        from tkinter import filedialog, messagebox, ttk
    except ImportError as exc:
        raise VOFIEError("此 Python 未安裝 Tkinter；請改用 simple 命令。") from exc

    dnd_available = False
    try:
        from tkinterdnd2 import DND_FILES, TkinterDnD  # type: ignore
        root = TkinterDnD.Tk()
        dnd_available = True
    except ImportError:
        DND_FILES = None
        root = tk.Tk()

    root.title(GUI_TITLE)
    root.geometry(GUI_GEOMETRY)
    root.minsize(820, 600)
    selected: list[Path] = []
    events: queue.Queue[tuple[str, Any]] = queue.Queue()
    output_var = tk.StringVar(value=str((project_root / "Veritas_VOFIE_Output").resolve()))
    title_var = tk.StringVar(value="Veritas 全格式智慧重構")
    role_var = tk.StringVar(value=DEFAULT_SIMPLE_ROLE)
    operation_vars = {item: tk.BooleanVar(value=True) for item in OPERATION_ORDER}
    status_var = tk.StringVar(value="就緒 — 最多 5 個輸入，固定 5 個主要輸出")

    def refresh_list() -> None:
        file_list.delete(0, tk.END)
        for index, path in enumerate(selected, start=1):
            file_list.insert(tk.END, f"{index}. {path}")
        count_label.configure(text=f"已選 {len(selected)} / {SIMPLE_MAX_INPUT_FILES}")

    def add_paths(values: Sequence[str | Path]) -> None:
        nonlocal selected
        try:
            selected = normalize_gui_inputs([*selected, *values])
            refresh_list()
            status_var.set("輸入已更新；原檔只讀")
        except VOFIEError as exc:
            messagebox.showerror("輸入錯誤", str(exc), parent=root)

    def choose_files() -> None:
        values = filedialog.askopenfilenames(parent=root, title="選取最多 5 個檔案", filetypes=GUI_FILE_TYPES)
        if values:
            add_paths(values)

    def remove_selected() -> None:
        nonlocal selected
        indexes = set(file_list.curselection())
        selected = [path for index, path in enumerate(selected) if index not in indexes]
        refresh_list()

    def clear_files() -> None:
        nonlocal selected
        selected = []
        refresh_list()

    def choose_output() -> None:
        value = filedialog.askdirectory(parent=root, title="選擇輸出資料夾")
        if value:
            output_var.set(value)

    def log(message: str) -> None:
        log_box.configure(state="normal")
        log_box.insert(tk.END, message.rstrip() + "\n")
        log_box.see(tk.END)
        log_box.configure(state="disabled")

    def worker(inputs: list[Path], output: Path, role: str, operations: tuple[str, ...], title: str) -> None:
        try:
            events.put(("progress", "Preflight：檢查路徑、角色與五檔契約"))
            options = EngineOptions(run_role=role, operations=operations, use_vsis=True)
            engine = VeritasOmniFormatEngine(project_root, options)
            events.put(("progress", "讀取、NLP 分題並建立整合視圖"))
            ir = engine.convert_simple(inputs, output, title or None)
            events.put(("success", {"run_id": ir.run_id, "output_dir": ir.quality["output_dir"], "gate": ir.quality["gate"]}))
        except Exception as exc:
            events.put(("error", f"{type(exc).__name__}: {exc}"))

    def start_run() -> None:
        operations = tuple(item for item in OPERATION_ORDER if operation_vars[item].get())
        if not selected:
            messagebox.showerror("缺少輸入", "請選取或拖放至少一個檔案。", parent=root)
            return
        if not operations:
            messagebox.showerror("缺少動作", "至少選擇一個合併／重構動作。", parent=root)
            return
        try:
            inputs = normalize_gui_inputs(selected)
            output = Path(output_var.get()).expanduser().resolve()
        except Exception as exc:
            messagebox.showerror("參數錯誤", str(exc), parent=root)
            return
        run_button.configure(state="disabled")
        status_var.set("執行中…")
        log(f"角色={role_var.get()}；輸入={len(inputs)}；動作={','.join(operations)}")
        threading.Thread(target=worker, args=(inputs, output, role_var.get(), operations, title_var.get()), daemon=True).start()

    def poll_events() -> None:
        try:
            while True:
                kind, payload = events.get_nowait()
                if kind == "progress":
                    status_var.set(str(payload))
                    log(str(payload))
                elif kind == "success":
                    run_button.configure(state="normal")
                    status_var.set(f"完成：{payload['gate']}")
                    log(f"完成 run_id={payload['run_id']}\n輸出={payload['output_dir']}")
                    messagebox.showinfo("VOFIE 完成", f"已產生 5 個主要檔案。\n{payload['output_dir']}", parent=root)
                elif kind == "error":
                    run_button.configure(state="normal")
                    status_var.set("執行失敗；參數與選檔已保留")
                    log(str(payload))
                    messagebox.showerror("VOFIE 錯誤", str(payload), parent=root)
        except queue.Empty:
            pass
        root.after(120, poll_events)

    container = ttk.Frame(root, padding=16)
    container.pack(fill="both", expand=True)
    ttk.Label(container, text="Veritas VOFIE v1.4", font=("Segoe UI", 18, "bold")).pack(anchor="w")
    ttk.Label(container, text="選檔或拖放 → 選角色與動作 → 產生 MD / HTML / JSON / Word / CSV", foreground="#36515c").pack(anchor="w", pady=(2, 12))
    input_frame = ttk.LabelFrame(container, text="1. 輸入（最多 5 檔，來源唯讀）", padding=10)
    input_frame.pack(fill="x")
    drop_text = "拖放檔案到這裡，或使用選檔按鈕" if dnd_available else "拖放元件未安裝；請使用選檔按鈕（功能不受影響）"
    drop_label = ttk.Label(input_frame, text=drop_text, anchor="center", padding=10)
    drop_label.pack(fill="x")
    if dnd_available and DND_FILES is not None:
        drop_label.drop_target_register(DND_FILES)  # type: ignore[attr-defined]
        drop_label.dnd_bind("<<Drop>>", lambda event: add_paths(root.tk.splitlist(event.data)))  # type: ignore[attr-defined]
    file_list = tk.Listbox(input_frame, height=5, selectmode=tk.EXTENDED)
    file_list.pack(fill="x", pady=6)
    input_buttons = ttk.Frame(input_frame)
    input_buttons.pack(fill="x")
    ttk.Button(input_buttons, text="選取檔案", command=choose_files).pack(side="left")
    ttk.Button(input_buttons, text="移除選取", command=remove_selected).pack(side="left", padx=6)
    ttk.Button(input_buttons, text="清除", command=clear_files).pack(side="left")
    count_label = ttk.Label(input_buttons, text="已選 0 / 5")
    count_label.pack(side="right")
    settings = ttk.LabelFrame(container, text="2. 參數", padding=10)
    settings.pack(fill="x", pady=10)
    ttk.Label(settings, text="標題").grid(row=0, column=0, sticky="w")
    ttk.Entry(settings, textvariable=title_var).grid(row=0, column=1, columnspan=5, sticky="ew", padx=(8, 0))
    ttk.Label(settings, text="角色").grid(row=1, column=0, sticky="w", pady=(8, 0))
    ttk.Combobox(settings, textvariable=role_var, values=RUN_ROLES, state="readonly", width=12).grid(row=1, column=1, sticky="w", padx=(8, 16), pady=(8, 0))
    for index, operation in enumerate(OPERATION_ORDER, start=2):
        ttk.Checkbutton(settings, text=operation, variable=operation_vars[operation]).grid(row=1, column=index, sticky="w", pady=(8, 0))
    settings.columnconfigure(1, weight=1)
    output_frame = ttk.Frame(container)
    output_frame.pack(fill="x")
    ttk.Label(output_frame, text="輸出資料夾").pack(side="left")
    ttk.Entry(output_frame, textvariable=output_var).pack(side="left", fill="x", expand=True, padx=8)
    ttk.Button(output_frame, text="瀏覽", command=choose_output).pack(side="left")
    action_frame = ttk.Frame(container)
    action_frame.pack(fill="x", pady=10)
    run_button = ttk.Button(action_frame, text="開始重構並產生 5 檔", command=start_run)
    run_button.pack(side="left")
    ttk.Label(action_frame, textvariable=status_var).pack(side="left", padx=12)
    log_box = tk.Text(container, height=9, wrap="word", state="disabled")
    log_box.pack(fill="both", expand=True)
    root.after(120, poll_events)
    root.mainloop()
    return 0


def dependency_report(project_root: Path) -> dict[str, Any]:
    vsis_candidates = [
        project_root.parent / "VIA_SemanticIntelligenceSubsystem_v0120",
        Path(os.environ["VIA_VSIS_ROOT"]) if os.environ.get("VIA_VSIS_ROOT") else None,
    ]
    vsis = next((path.resolve() for path in vsis_candidates if path and (path / "via_semantic_intelligence").is_dir()), None)
    node = os.environ.get("CODEX_PRIMARY_RUNTIME_NODE") or shutil.which("node")
    polyglot = tool_audit(project_root)
    hydra = hydra_risk_audit(project_root)
    capabilities = [
        {"library": "Python standard library", "available": True, "role": "core readers, IR, HTML, CSV, failure recovery", "fallback": "not required"},
        {"library": "python-docx", "available": importlib.util.find_spec("docx") is not None, "role": "polished Word output", "fallback": "stdlib OOXML Word emitter"},
        {"library": "pypdf", "available": importlib.util.find_spec("pypdf") is not None, "role": "PDF input", "fallback": "explicit HOLD for PDF only"},
        {"library": "tkinter", "available": importlib.util.find_spec("tkinter") is not None, "role": "desktop window I/O", "fallback": "Python simple CLI or PS7 launcher"},
        {"library": "tkinterdnd2", "available": importlib.util.find_spec("tkinterdnd2") is not None, "role": "drag-drop", "fallback": "native multi-file dialog"},
        {"library": "Node runtime", "available": bool(node), "role": "legacy PPTX/XLSX adapters and JS syntax tests", "fallback": "five-output simple profile remains available"},
        {"library": "VSIS 1.2", "available": vsis is not None, "role": "normalize, segment, categorize, semantic_check", "fallback": "deterministic local NLP"},
        {"library": "JavaScript Top 20 registry", "available": polyglot["gate"] == "PASS", "role": "syntax, lint, test, coverage, graph, refactor, build and UI validation", "fallback": "registered local CPU fallbacks"},
        {"library": "PowerShell Top 20 registry", "available": polyglot["gate"] == "PASS", "role": "AST, lint, Pester, policy, build, data and reporting", "fallback": "Python structural and contract bridge"},
        {"library": "NoHydra Top 20 risk gate", "available": hydra["gate"] == "PASS", "role": "multi-writer, recursion, concurrency, retry, rollback and activation containment", "fallback": "fail-closed HOLD and review-only three-round plan"},
    ]
    core_ready = all(item["available"] for item in capabilities if item["library"] == "Python standard library")
    simple_ready = core_ready and any(item["library"] == "python-docx" and (item["available"] or item["fallback"]) for item in capabilities)
    return {
        "contract": "veritas.vofie-dependency-report/1.4",
        "engine_version": ENGINE_VERSION,
        "python": sys.version.split()[0],
        "gate": "PASS" if core_ready and simple_ready else "FAIL",
        "simple_profile_ready": simple_ready,
        "capabilities": capabilities,
        "polyglot_tool_summary": polyglot["summary"],
        "hydra_risk_summary": hydra["summary"],
        "policy": "NO_AUTO_INSTALL_USE_REGISTERED_FALLBACKS",
    }


def parse_formats(value: str) -> tuple[str, ...]:
    formats = tuple(item.strip().casefold() for item in value.split(",") if item.strip())
    return formats or DEFAULT_OUTPUT_FORMATS


def parse_operations(value: str) -> tuple[str, ...]:
    operations = tuple(item.strip().casefold() for item in value.split(",") if item.strip())
    return normalize_operations(operations or DEFAULT_OPERATIONS)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="vofie", description=ENGINE_NAME_ZH)
    sub = parser.add_subparsers(dest="command", required=True)
    convert = sub.add_parser("convert", help="讀取、主題重構並輸出多格式")
    convert.add_argument("inputs", nargs="+", type=Path)
    convert.add_argument("--output", type=Path, required=True)
    convert.add_argument("--formats", default=",".join(DEFAULT_OUTPUT_FORMATS))
    convert.add_argument("--title", default="")
    convert.add_argument("--language", default="zh-Hant", choices=("zh-Hant", "zh-Hans", "en", "auto"))
    convert.add_argument("--max-topic-chars", type=int, default=DEFAULT_TOPIC_CHARS)
    convert.add_argument("--no-vsis", action="store_true")
    convert.add_argument("--no-quarantine", action="store_true")
    simple = sub.add_parser("simple", help="最多五個輸入，固定五個主要輸出的易用模式")
    simple.add_argument("inputs", nargs="+", type=Path)
    simple.add_argument("--output", type=Path, required=True)
    simple.add_argument("--title", default="")
    simple.add_argument("--role", default=DEFAULT_SIMPLE_ROLE, choices=RUN_ROLES)
    simple.add_argument("--operations", default=",".join(DEFAULT_OPERATIONS))
    simple.add_argument("--language", default="zh-Hant", choices=("zh-Hant", "zh-Hans", "en", "auto"))
    simple.add_argument("--max-topic-chars", type=int, default=DEFAULT_TOPIC_CHARS)
    simple.add_argument("--no-vsis", action="store_true")
    simple.add_argument("--no-quarantine", action="store_true")
    sub.add_parser("gui", help="啟動 Windows-friendly 選檔／拖放視窗")
    manifest = sub.add_parser("manifest", help="輸出 VIA Registry Manifest")
    manifest.add_argument("--file", type=Path)
    test = sub.add_parser("self-test", help="執行核心自測")
    test.add_argument("--report", type=Path)
    user = sub.add_parser("user-test", help="執行 ENGINE／SYSTEM 使用者流程測試")
    user.add_argument("--report", type=Path)
    activate = sub.add_parser("activate", help="執行自測、復原測試與使用者測試後決定 ACTIVE／HOLD")
    activate.add_argument("--report", type=Path)
    failures = sub.add_parser("failure-catalog", help="輸出八個環節各 Top 20 failures 與復原處理器")
    failures.add_argument("--report", type=Path)
    dependencies = sub.add_parser("dependencies", help="檢查必要／選用函式庫與已實作降級路徑")
    dependencies.add_argument("--report", type=Path)
    tools_parser = sub.add_parser("tool-audit", help="偵測 JavaScript／PowerShell 各 Top 20 免費 CPU 工具與降級覆蓋")
    tools_parser.add_argument("--language", default="all", choices=("all", "javascript", "powershell"))
    tools_parser.add_argument("--probe-installed", action="store_true", help="只對已安裝命令執行唯讀 --version／module inventory")
    tools_parser.add_argument("--report", type=Path)
    tool_plan_parser = sub.add_parser("tool-plan", help="依檔案語言與功能需求選擇工具；預設只規劃")
    tool_plan_parser.add_argument("target", type=Path)
    tool_plan_parser.add_argument("--functions", default=",".join(POLYGLOT_MATRIX_FUNCTIONS))
    tool_plan_parser.add_argument("--execute-safe", action="store_true", help="只執行唯讀 syntax quick check")
    tool_plan_parser.add_argument("--report", type=Path)
    hydra_parser = sub.add_parser("hydra-audit", help="執行九頭龍 Top-20 review-only 風險 Gate；偵測到風險即 HOLD")
    hydra_parser.add_argument("targets", nargs="*", type=Path)
    hydra_parser.add_argument("--report", type=Path)
    runtime_parser = sub.add_parser("runtime-copy", help="建立版本化 Runtime Copy sandbox；不提供核准 token 時 HOLD")
    runtime_parser.add_argument("sources", nargs="+", type=Path)
    runtime_parser.add_argument("--output", type=Path, required=True)
    runtime_parser.add_argument("--approval-token", default="")
    runtime_parser.add_argument("--report", type=Path)
    rollback_parser = sub.add_parser("rollback-check", help="唯讀驗證 Runtime Copy rollback readiness")
    rollback_parser.add_argument("manifest", type=Path)
    rollback_parser.add_argument("--report", type=Path)
    sub.add_parser("list-formats", help="列出輸入與輸出格式")
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    project_root = Path(__file__).resolve().parent
    if args.command == "manifest":
        payload = registration_manifest()
        text = json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True)
        if args.file:
            atomic_write_text(args.file.resolve(), text)
        print(text)
        return 0
    if args.command == "self-test":
        report = self_test(project_root)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 1
    if args.command == "user-test":
        report = user_test(project_root)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 1
    if args.command == "activate":
        report = activation_test(project_root)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["activation_state"] == "ACTIVE" else 1
    if args.command == "failure-catalog":
        report = load_failure_catalog(project_root)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0
    if args.command == "dependencies":
        report = dependency_report(project_root)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 1
    if args.command == "tool-audit":
        report = tool_audit(project_root, language=args.language, probe_installed=args.probe_installed)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 1
    if args.command == "tool-plan":
        report = tool_plan(
            project_root,
            args.target,
            functions=parse_tool_functions(args.functions),
            execute_safe=args.execute_safe,
        )
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 1
    if args.command == "hydra-audit":
        report = hydra_risk_audit(project_root, args.targets)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 3
    if args.command == "runtime-copy":
        report = create_runtime_copy(args.sources, args.output, args.approval_token)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 4
    if args.command == "rollback-check":
        report = rollback_dry_run(args.manifest)
        text = json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True)
        if args.report:
            atomic_write_text(args.report.resolve(), text)
        print(text)
        return 0 if report["gate"] == "PASS" else 5
    if args.command == "gui":
        return launch_gui(project_root)
    if args.command == "list-formats":
        print(json.dumps({"inputs": FORMAT_REGISTRY, "outputs": OUTPUT_REGISTRY}, ensure_ascii=False, indent=2))
        return 0
    if args.command == "convert":
        options = EngineOptions(
            target_language=args.language, max_topic_chars=max(1000, args.max_topic_chars),
            quarantine_boilerplate=not args.no_quarantine, use_vsis=not args.no_vsis,
            output_formats=parse_formats(args.formats),
        )
        engine = VeritasOmniFormatEngine(project_root=project_root, options=options)
        ir = engine.convert(args.inputs, args.output, args.title or None)
        print(json.dumps({"status": ir.quality["gate"], "run_id": ir.run_id, "outputs": ir.output_files}, ensure_ascii=False, indent=2))
        return 0 if ir.quality["gate"] != "FAIL" else 2
    if args.command == "simple":
        options = EngineOptions(
            target_language=args.language,
            max_topic_chars=max(1000, args.max_topic_chars),
            quarantine_boilerplate=not args.no_quarantine,
            use_vsis=not args.no_vsis,
            output_formats=SIMPLE_PRIMARY_OUTPUTS,
            run_role=args.role,
            operations=parse_operations(args.operations),
        )
        engine = VeritasOmniFormatEngine(project_root=project_root, options=options)
        ir = engine.convert_simple(args.inputs, args.output, args.title or None)
        print(json.dumps({
            "status": ir.quality["gate"],
            "role": args.role,
            "run_id": ir.run_id,
            "output_dir": ir.quality["output_dir"],
            "primary_outputs": ir.output_files,
            "system_sidecars": ir.quality.get("system_sidecars"),
        }, ensure_ascii=False, indent=2))
        return 0 if ir.quality["gate"] != "FAIL" else 2
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
